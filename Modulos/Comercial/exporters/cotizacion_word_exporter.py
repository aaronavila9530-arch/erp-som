from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import os

from resource_utils import resource_path


# ============================================================
# ASSETS PATH — BLINDADO TOTAL (DEV / EXE / PROGRAM FILES)
# ============================================================
def get_assets_path() -> str:
    """
    Devuelve la ruta correcta a /assets usando resource_path.
    Funciona en:
    - Desarrollo
    - PyInstaller ONEDIR
    - Instalado en Program Files
    """

    assets_path = resource_path("assets")

    if not os.path.isdir(assets_path):
        raise RuntimeError(
            f"Assets folder not found via resource_path: {assets_path}"
        )

    return assets_path


# ============================================================
# EXPORT COTIZACIÓN WORD
# ============================================================
def export_cotizacion_word(data: dict, output_path: str):
    """
    Genera cotización en formato Word (.docx)

    Data esperado:
    {
        quotation_number: str,
        cliente: str,
        servicio: str,
        idioma: 'ES' | 'EN',
        texto: str
    }
    """

    ASSETS_PATH = get_assets_path()

    doc = Document()

    # ==================================================
    # HEADER (LOGO)
    # ==================================================
    section = doc.sections[0]
    header = section.header

    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run()

    header_img_path = os.path.join(ASSETS_PATH, "header.png")
    if os.path.isfile(header_img_path):
        hr.add_picture(
            header_img_path,
            width=Inches(2.5)
        )

    # ==================================================
    # BODY
    # ==================================================
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --------------------------------------------------
    # QUOTATION NUMBER
    # --------------------------------------------------
    quotation_number = data.get("quotation_number")
    if quotation_number:
        body.add_run(f"{quotation_number}\n").bold = True

    # --------------------------------------------------
    # FECHA / CIUDAD
    # --------------------------------------------------
    body.add_run(
        f"Alajuela, Costa Rica\n{date.today()}\n\n"
    )

    # --------------------------------------------------
    # SUBJECT / ASUNTO
    # --------------------------------------------------
    idioma = data.get("idioma", "ES")

    if idioma == "EN":
        subject = f"Quotation – {data.get('servicio', '')}"
        body.add_run(f"Subject: {subject}\n\n").bold = True
    else:
        subject = f"Cotización – {data.get('servicio', '')}"
        body.add_run(f"Asunto: {subject}\n\n").bold = True

    # --------------------------------------------------
    # TEXTO PRINCIPAL
    # --------------------------------------------------
    body.add_run(data.get("texto", ""))

    # ==================================================
    # SIGNATURE
    # ==================================================
    doc.add_paragraph("\n")

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if idioma == "EN":
        sig.add_run("Sincerely,\n\n")
    else:
        sig.add_run("Atentamente,\n\n")

    signature_path = os.path.join(ASSETS_PATH, "FIRMA DIANA.png")
    if os.path.isfile(signature_path):
        sig.add_run().add_picture(
            signature_path,
            width=Inches(1.8)
        )

    sig.add_run("\nDiana Quirós Benambourg\n")
    sig.add_run("Business Manager\n")
    sig.add_run("Marine Surveyors & Logistics Group SRL")

    # ==================================================
    # FOOTER
    # ==================================================
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = (
        "Head Office – Costa Rica, Alajuela, Plaza Aeropuerto G-14\n"
        "Phone (506) 8814-07-84 – (506) 4052-8382"
    )

    # ==================================================
    # SAVE
    # ==================================================
    doc.save(output_path)

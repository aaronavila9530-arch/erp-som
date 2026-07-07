import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText
from datetime import datetime, timedelta
from pathlib import Path

import api_client
from tkcalendar import DateEntry, Calendar
from Modulos.Informes.logra_questionnaires_data import ONG_QUESTIONNAIRES
from Modulos.Informes.popup.popup_ai_compare import PopupAICompare
from session_context import get_user


class LograQuestionnairesForm(ttk.Frame):
    ITEMS_PER_PAGE = 5
    MAX_BULLETS = 20
    MAX_ATTACHMENTS_PER_QUESTION = 10

    SECTION_LABELS = {
        "critical_questions": "Preguntas de apertura",
        "detailed_questions": "Preguntas por tema",
    }
    PRIORITY_COLORS = {
        "Alta": "#F8D7DA",
        "Media": "#FFF3CD",
        "Baja": "#D1E7DD",
    }
    STATUS_COLORS = {
        "Pendiente": "#F8D7DA",
        "En proceso": "#FFF3CD",
        "Completado": "#D1E7DD",
    }

    def __init__(self, parent, usuario=None, rol=None, on_back=None, review_mode=False):
        super().__init__(parent)
        self.parent = parent
        self.usuario = usuario
        self.rol = rol
        self.on_back = on_back
        self.review_mode = review_mode

        self.report_id = None
        self.form_var = tk.StringVar()
        self.search_var = tk.StringVar()
        self.section_var = tk.StringVar(value="critical_questions")
        self.page_index = 0
        self.answers = {}
        self.agenda_items = []
        self.agenda_notes = ""
        self.text_widgets = {}
        self._search_trace = None
        self._agenda_alerted = set()

        self.pack(fill="both", expand=True)
        self._build_ui()
        self._start_agenda_alert_monitor()

    # =========================================================
    # UI
    # =========================================================
    def _build_ui(self):
        self._build_topbar()
        self._build_filters()
        self._build_content()
        self.form_var.set(ONG_QUESTIONNAIRES[0]["title"])
        self._render_current_page()

    def _build_topbar(self):
        bar = ttk.Frame(self)
        bar.pack(fill="x", padx=12, pady=(10, 6))
        bar.columnconfigure(0, weight=1)

        title_box = ttk.Frame(bar)
        title_box.grid(row=0, column=0, sticky="ew")
        ttk.Label(
            title_box,
            text="ONG - Cuestionarios",
            font=("Segoe UI", 14, "bold")
        ).pack(anchor="w")
        ttk.Label(
            title_box,
            text="Agenda, preguntas del documento, hasta 20 bullet points por pregunta y adjuntos guardados en backend.",
            foreground="#555555"
        ).pack(anchor="w", pady=(2, 0))

        actions = ttk.Frame(bar)
        actions.grid(row=0, column=1, sticky="e")
        if not self.review_mode:
            ttk.Button(actions, text="Agenda", command=self._open_agenda).pack(side="left", padx=4)
            ttk.Button(actions, text="Mejorar con PORTIA", command=self._open_portia).pack(side="left", padx=4)
            ttk.Button(actions, text="Guardar", command=self._save_report).pack(side="left", padx=4)
        else:
            ttk.Button(actions, text="Ver agenda", command=self._open_agenda).pack(side="left", padx=4)
            ttk.Button(actions, text="Actualizar", command=self._save_report).pack(side="left", padx=4)
        ttk.Button(actions, text="Home", command=self._go_home).pack(side="left", padx=4)

    def _build_filters(self):
        filters = ttk.Frame(self)
        filters.pack(fill="x", padx=12, pady=(0, 8))
        filters.columnconfigure(1, weight=1)
        filters.columnconfigure(3, weight=1)

        ttk.Label(filters, text="Formulario").grid(row=0, column=0, sticky="w", padx=(0, 6))
        cb = ttk.Combobox(
            filters,
            textvariable=self.form_var,
            state="readonly",
            values=[item["title"] for item in ONG_QUESTIONNAIRES],
        )
        cb.grid(row=0, column=1, sticky="ew", padx=(0, 12))
        cb.bind("<<ComboboxSelected>>", self._on_context_changed)

        ttk.Label(filters, text="Buscar").grid(row=0, column=2, sticky="w", padx=(0, 6))
        search = ttk.Entry(filters, textvariable=self.search_var)
        search.grid(row=0, column=3, sticky="ew")
        search.bind("<KeyRelease>", self._on_context_changed)
        self._search_trace = self.search_var.trace_add("write", lambda *_: self._on_context_changed())

    def _build_content(self):
        shell = ttk.Frame(self)
        shell.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        shell.columnconfigure(1, weight=1)
        shell.rowconfigure(0, weight=1)

        nav = ttk.LabelFrame(shell, text="Secciones", padding=10)
        nav.grid(row=0, column=0, sticky="ns", padx=(0, 10))

        for key, label in self.SECTION_LABELS.items():
            ttk.Radiobutton(
                nav,
                text=label,
                value=key,
                variable=self.section_var,
                command=self._on_context_changed
            ).pack(anchor="w", pady=4)

        self.count_label = ttk.Label(nav, text="", justify="left")
        self.count_label.pack(anchor="w", pady=(18, 0))

        content = ttk.Frame(shell)
        content.grid(row=0, column=1, sticky="nsew")
        content.rowconfigure(1, weight=1)
        content.columnconfigure(0, weight=1)

        pager = ttk.Frame(content)
        pager.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        pager.columnconfigure(1, weight=1)
        ttk.Button(pager, text="Anterior", command=self._prev_page).grid(row=0, column=0, sticky="w")
        self.page_label = ttk.Label(pager, text="", anchor="center")
        self.page_label.grid(row=0, column=1, sticky="ew")
        ttk.Button(pager, text="Siguiente", command=self._next_page).grid(row=0, column=2, sticky="e")

        canvas_wrap = ttk.Frame(content)
        canvas_wrap.grid(row=1, column=0, sticky="nsew")
        canvas_wrap.rowconfigure(0, weight=1)
        canvas_wrap.columnconfigure(0, weight=1)

        self.canvas = tk.Canvas(canvas_wrap, highlightthickness=0)
        self.canvas.grid(row=0, column=0, sticky="nsew")
        scroll = ttk.Scrollbar(canvas_wrap, orient="vertical", command=self.canvas.yview)
        scroll.grid(row=0, column=1, sticky="ns")
        self.canvas.configure(yscrollcommand=scroll.set)

        self.scroll_frame = ttk.Frame(self.canvas)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scroll_frame, anchor="nw")
        self.scroll_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.bind("<Configure>", lambda e: self.canvas.itemconfig(self.canvas_window, width=e.width))
        self.scroll_frame.bind("<Enter>", self._bind_mousewheel)
        self.scroll_frame.bind("<Leave>", self._unbind_mousewheel)

    def _time_values(self):
        return [f"{hour:02d}:{minute:02d}" for hour in range(24) for minute in range(60)]

    # =========================================================
    # Scroll wheel
    # =========================================================
    def _bind_mousewheel(self, event=None):
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Button-4>", lambda e: self.canvas.yview_scroll(-3, "units"))
        self.canvas.bind_all("<Button-5>", lambda e: self.canvas.yview_scroll(3, "units"))

    def _unbind_mousewheel(self, event=None):
        self.canvas.unbind_all("<MouseWheel>")
        self.canvas.unbind_all("<Button-4>")
        self.canvas.unbind_all("<Button-5>")

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)) * 3, "units")

    # =========================================================
    # Data helpers
    # =========================================================
    def _current_form(self):
        title = self.form_var.get()
        return next((item for item in ONG_QUESTIONNAIRES if item["title"] == title), ONG_QUESTIONNAIRES[0])

    def _all_questions(self, form=None, section=None):
        form = form or self._current_form()
        section = section or self.section_var.get()
        return list(form.get(section, []))

    def _items(self):
        items = self._all_questions()
        query = self.search_var.get().strip().lower()
        if query:
            filtered = []
            for item in items:
                text = " ".join([
                    str(item.get("id") or item.get("number") or ""),
                    str(item.get("block") or ""),
                    str(item.get("question") or ""),
                ]).lower()
                if query in text:
                    filtered.append(item)
            return filtered
        return items

    def _item_key(self, item):
        return str(item.get("id") or item.get("number") or "").strip()

    def _answer_key(self, form, section, item):
        return f"{form['slug']}|{section}|{self._item_key(item)}"

    def _get_bullets(self, form, section, item):
        key = self._answer_key(form, section, item)
        self.answers.setdefault(key, [""])
        return self.answers[key]

    def _collect_visible_text(self):
        for key, widgets in list(self.text_widgets.items()):
            self.answers[key] = [
                widget.get("1.0", "end-1c").strip()
                for widget in widgets
            ][:self.MAX_BULLETS]

    def _on_context_changed(self, event=None):
        self._collect_visible_text()
        self.page_index = 0
        self._render_current_page()

    def _prev_page(self):
        self._collect_visible_text()
        if self.page_index > 0:
            self.page_index -= 1
            self._render_current_page()

    def _next_page(self):
        self._collect_visible_text()
        items = self._items()
        max_page = max(0, (len(items) - 1) // self.ITEMS_PER_PAGE)
        if self.page_index < max_page:
            self.page_index += 1
            self._render_current_page()

    # =========================================================
    # Render
    # =========================================================
    def _render_current_page(self):
        for widget in self.scroll_frame.winfo_children():
            widget.destroy()
        self.text_widgets.clear()

        form = self._current_form()
        section = self.section_var.get()
        items = self._items()
        total_pages = max(1, (len(items) + self.ITEMS_PER_PAGE - 1) // self.ITEMS_PER_PAGE)
        self.page_index = min(self.page_index, total_pages - 1)
        start = self.page_index * self.ITEMS_PER_PAGE
        visible = items[start:start + self.ITEMS_PER_PAGE]

        self.count_label.configure(
            text=(
                f"Apertura: {len(form.get('critical_questions', []))}\n"
                f"Por tema: {len(form.get('detailed_questions', []))}\n"
                f"Total: {len(form.get('critical_questions', [])) + len(form.get('detailed_questions', []))}"
            )
        )
        self.page_label.configure(text=f"Pagina {self.page_index + 1} de {total_pages} - {len(items)} preguntas")

        if not visible:
            ttk.Label(self.scroll_frame, text="No hay resultados.").pack(anchor="w", padx=8, pady=8)
            return

        for item in visible:
            self._build_question_card(form, section, item)

    def _build_question_card(self, form, section, item):
        item_key = self._item_key(item)
        title_bits = [item_key]
        if section == "detailed_questions":
            title_bits.append(item.get("block") or "Pregunta")

        card = ttk.LabelFrame(self.scroll_frame, text=" - ".join(title_bits), padding=10)
        card.pack(fill="x", padx=4, pady=7)
        card.columnconfigure(0, weight=1)

        if section == "detailed_questions" and item.get("priority"):
            ttk.Label(card, text=f"Referencia: {item.get('priority')}", foreground="#555555").grid(
                row=0, column=0, sticky="w"
            )

        ttk.Label(
            card,
            text=item.get("question", ""),
            wraplength=1080,
            justify="left",
            font=("Segoe UI", 9, "bold")
        ).grid(row=1, column=0, sticky="ew", pady=(6, 8))

        key = self._answer_key(form, section, item)
        bullets = self._get_bullets(form, section, item)
        self.text_widgets[key] = []

        bullets_box = ttk.Frame(card)
        bullets_box.grid(row=2, column=0, sticky="ew")
        bullets_box.columnconfigure(1, weight=1)

        for idx, value in enumerate(bullets, start=1):
            ttk.Label(bullets_box, text=f"{idx}.").grid(row=idx - 1, column=0, sticky="nw", padx=(0, 6), pady=3)
            text = ScrolledText(bullets_box, height=2, wrap="word", font=("Segoe UI", 9))
            text.insert("1.0", value)
            text.grid(row=idx - 1, column=1, sticky="ew", pady=3)
            self.text_widgets[key].append(text)

        actions = ttk.Frame(card)
        actions.grid(row=3, column=0, sticky="ew", pady=(8, 0))

        ttk.Button(actions, text="+ Bullet", command=lambda f=form, s=section, i=item: self._add_bullet(f, s, i)).pack(
            side="left", padx=(0, 4)
        )
        ttk.Button(actions, text="- Bullet", command=lambda f=form, s=section, i=item: self._remove_bullet(f, s, i)).pack(
            side="left", padx=(0, 12)
        )
        if not self.review_mode:
            ttk.Button(actions, text="Adjuntar", command=lambda f=form, s=section, i=item: self._attach_file(f, s, i)).pack(
                side="left"
            )
        if self.report_id:
            ttk.Button(
                actions,
                text="Actualizar pregunta",
                command=lambda f=form, s=section, i=item: self._update_question(f, s, i)
            ).pack(side="left", padx=(8, 0))

        self._build_attachments(card, form, section, item)

    # =========================================================
    # Agenda
    # =========================================================
    def _render_agenda(self):
        if not hasattr(self, "agenda_tree"):
            return
        for row in self.agenda_tree.get_children():
            self.agenda_tree.delete(row)
        for idx, item in enumerate(self.agenda_items):
            priority = item.get("priority") or "Media"
            status = item.get("status") or "Pendiente"
            tag = status if status in self.STATUS_COLORS else priority
            self.agenda_tree.insert(
                "",
                "end",
                iid=str(idx),
                values=(
                    item.get("date") or "",
                    item.get("start_time") or "",
                    item.get("end_time") or "",
                    item.get("place") or "",
                    item.get("person") or "",
                    item.get("phone") or item.get("telefono") or "",
                    item.get("company") or "",
                    item.get("topic") or "",
                    priority,
                    status,
                ),
                tags=(tag,)
            )

    def _add_agenda_item(self):
        item = {
            "date": self._format_agenda_date(),
            "start_time": self.agenda_start.get(),
            "end_time": self.agenda_end.get(),
            "place": self.agenda_place.get().strip(),
            "person": self.agenda_person.get().strip(),
            "phone": getattr(self, "agenda_phone", tk.StringVar()).get().strip(),
            "company": self.agenda_company.get().strip(),
            "topic": self.agenda_topic.get().strip(),
            "priority": self.agenda_priority.get(),
            "status": self.agenda_status.get(),
        }
        if not item["person"] and not item["topic"] and not item["place"]:
            messagebox.showwarning("Agenda ONG", "Agrega al menos persona, tema o lugar.")
            return
        self.agenda_items.append(item)
        self._render_agenda()

    def _remove_agenda_item(self):
        selected = self.agenda_tree.selection()
        if not selected:
            messagebox.showwarning("Agenda ONG", "Selecciona una linea de agenda.")
            return
        indexes = sorted([int(value) for value in selected], reverse=True)
        for index in indexes:
            if 0 <= index < len(self.agenda_items):
                self.agenda_items.pop(index)
        self._render_agenda()

    def _build_attachments(self, parent, form, section, item):
        if not self.report_id:
            if self.review_mode:
                return
            ttk.Label(parent, text="Guarda primero para habilitar adjuntos persistentes.", foreground="#777777").grid(
                row=4, column=0, sticky="w", pady=(8, 0)
            )
            return

        item_key = self._item_key(item)
        resp = api_client.list_logra_attachments_api(self.report_id, form["slug"], section, item_key)
        attachments = resp.get("data") or []
        if not attachments:
            return

        box = ttk.Frame(parent)
        box.grid(row=4, column=0, sticky="ew", pady=(8, 0))
        ttk.Label(box, text="Adjuntos:", foreground="#555555").pack(side="left", padx=(0, 6))
        for att in attachments[:self.MAX_ATTACHMENTS_PER_QUESTION]:
            item_box = ttk.Frame(box)
            item_box.pack(side="left", padx=3)
            ttk.Button(
                item_box,
                text=att.get("original_filename") or f"Adjunto {att.get('id')}",
                command=lambda a=att: self._open_attachment(a)
            ).pack(side="left")
            if not self.review_mode:
                ttk.Button(
                    item_box,
                    text="-",
                    width=2,
                    command=lambda a=att: self._delete_attachment(a)
                ).pack(side="left", padx=(2, 0))

    # =========================================================
    # Actions
    # =========================================================
    def _add_bullet(self, form, section, item):
        self._collect_visible_text()
        bullets = self._get_bullets(form, section, item)
        if len(bullets) >= self.MAX_BULLETS:
            messagebox.showwarning("ONG", "Cada pregunta permite maximo 20 bullet points.")
            return
        bullets.append("")
        self._render_current_page()

    def _remove_bullet(self, form, section, item):
        self._collect_visible_text()
        bullets = self._get_bullets(form, section, item)
        if len(bullets) <= 1:
            bullets[0] = ""
        else:
            bullets.pop()
        self._render_current_page()

    def _update_question(self, form, section, item):
        if not self.report_id:
            messagebox.showwarning("ONG", "Guarda o carga primero el reporte ONG.")
            return
        self._collect_visible_text()
        key = self._answer_key(form, section, item)
        payload = {
            "form_slug": form["slug"],
            "form_title": form["title"],
            "section": section,
            "item_key": self._item_key(item),
            "question_text": item.get("question", ""),
            "bullets": [value.strip() for value in self.answers.get(key, []) if value.strip()],
        }
        resp = api_client.update_logra_answer_api(self.report_id, payload)
        if not resp.get("success"):
            messagebox.showerror("ONG", f"No se pudo actualizar la pregunta:\n{resp.get('error') or resp}")
            return
        messagebox.showinfo("ONG", "Pregunta actualizada correctamente.")

    def _answers_payload(self):
        self._collect_visible_text()
        payload = []
        for form in ONG_QUESTIONNAIRES:
            for section in self.SECTION_LABELS:
                for item in form.get(section, []):
                    key = self._answer_key(form, section, item)
                    bullets = [value.strip() for value in self.answers.get(key, []) if value.strip()]
                    if not bullets:
                        continue
                    payload.append({
                        "form_slug": form["slug"],
                        "form_title": form["title"],
                        "section": section,
                        "item_key": self._item_key(item),
                        "question_text": item.get("question", ""),
                        "bullets": bullets,
                    })
        return payload

    def _save_report(self, silent=False):
        title = f"ONG - {self.form_var.get() or 'Cuestionarios'}"
        payload = {
            "id": self.report_id,
            "title": title,
            "created_by": self.usuario or get_user(),
            "agenda_items": self.agenda_items,
            "agenda_notes": self.agenda_notes,
            "answers": self._answers_payload(),
        }
        resp = api_client.save_logra_report_api(payload)
        if not resp.get("success"):
            if not silent:
                messagebox.showerror("ONG", f"No se pudo guardar:\n{resp.get('error') or resp}")
            return False
        self.report_id = (resp.get("report") or {}).get("id") or self.report_id
        if not silent:
            messagebox.showinfo("ONG", "Guardado correctamente.")
            self._render_current_page()
        return True

    def _attach_file(self, form, section, item):
        if not self.report_id:
            if not self._save_report(silent=False):
                messagebox.showerror(
                    "ONG",
                    "No se pudo guardar el reporte antes de adjuntar. Revisa que el backend este corriendo."
                )
                return

        existing = api_client.list_logra_attachments_api(
            self.report_id,
            form["slug"],
            section,
            self._item_key(item)
        ).get("data") or []
        if len(existing) >= self.MAX_ATTACHMENTS_PER_QUESTION:
            messagebox.showwarning(
                "ONG",
                f"Cada pregunta permite maximo {self.MAX_ATTACHMENTS_PER_QUESTION} adjuntos."
            )
            return

        path = filedialog.askopenfilename(title="Seleccionar adjunto")
        if not path:
            return

        resp = api_client.upload_logra_attachment_api(
            self.report_id,
            form["slug"],
            section,
            self._item_key(item),
            path
        )
        if not resp.get("success"):
            messagebox.showerror("ONG", f"No se pudo subir el adjunto:\n{resp.get('error') or resp}")
            return

        messagebox.showinfo("ONG", "Adjunto guardado correctamente.")
        self._render_current_page()

    def _open_attachment(self, attachment):
        resp = api_client.open_logra_attachment_api(attachment.get("id"))
        if not resp.get("success"):
            messagebox.showerror("ONG", f"No se pudo abrir el adjunto:\n{resp.get('error') or resp}")

    def _delete_attachment(self, attachment):
        name = attachment.get("original_filename") or f"Adjunto {attachment.get('id')}"
        if not messagebox.askyesno("ONG", f"Eliminar adjunto?\n\n{name}"):
            return
        resp = api_client.delete_logra_attachment_api(attachment.get("id"))
        if not resp.get("success"):
            messagebox.showerror("ONG", f"No se pudo eliminar el adjunto:\n{resp.get('error') or resp}")
            return
        messagebox.showinfo("ONG", "Adjunto eliminado correctamente.")
        self._render_current_page()

    def _open_agenda(self):
        PopupLograAgenda(self, self)

    def _open_portia(self):
        self._collect_visible_text()
        PopupLograPortia(self, self)

    def _open_saved_report(self):
        PopupLograOpen(self, self)

    def load_report(self, report_id):
        resp = api_client.get_logra_report_api(report_id)
        if resp.get("success") is False:
            messagebox.showerror("ONG", f"No se pudo abrir el reporte:\n{resp.get('error') or resp}")
            return

        self.report_id = (resp.get("report") or {}).get("id")
        report = resp.get("report") or {}
        self.agenda_items = report.get("agenda_items") or []
        self.agenda_notes = report.get("agenda_notes") or ""
        self.answers.clear()
        for item in resp.get("answers") or []:
            key = f"{item.get('form_slug')}|{item.get('section')}|{item.get('item_key')}"
            bullets = item.get("bullets") or []
            self.answers[key] = bullets if bullets else [""]

        messagebox.showinfo("ONG", f"Reporte ONG #{self.report_id} cargado.")
        self._render_agenda()
        self._render_current_page()

    def _go_home(self):
        for widget in self.parent.winfo_children():
            widget.destroy()
        from Modulos.Informes.informes_home_ui import InformesHomeUI

        InformesHomeUI(self.parent, usuario=self.usuario, rol=self.rol)

    # =========================================================
    # Agenda alerts
    # =========================================================
    def _start_agenda_alert_monitor(self):
        self.after(60000, self._check_agenda_alerts)

    def _parse_agenda_datetime(self, item, key):
        date_value = item.get("date_iso") or item.get("date") or ""
        time_value = item.get(key) or ""
        for date_format in ("%Y-%m-%d", "%B %d, %Y"):
            try:
                parsed_date = datetime.strptime(date_value, date_format).date()
                parsed_time = datetime.strptime(time_value, "%H:%M").time()
                return datetime.combine(parsed_date, parsed_time)
            except Exception:
                continue
        return None

    def _check_agenda_alerts(self):
        try:
            now = datetime.now()
            for index, item in enumerate(self.agenda_items or []):
                status = (item.get("status") or "").strip().lower()
                if "complet" in status:
                    continue
                start = self._parse_agenda_datetime(item, "start_time")
                end = self._parse_agenda_datetime(item, "end_time")
                if not start:
                    continue
                try:
                    reminder = int(item.get("reminder_minutes") or 0)
                except Exception:
                    reminder = 0

                label = item.get("topic") or item.get("person") or "Reunion ONG"
                if reminder > 0 and start - timedelta(minutes=reminder) <= now < start:
                    key = (index, item.get("date_iso") or item.get("date"), item.get("start_time"), "before")
                    if key not in self._agenda_alerted:
                        self._agenda_alerted.add(key)
                        messagebox.showinfo("Agenda ONG", f"La reunion '{label}' inicia en menos de {reminder} minutos.")

                if start <= now and (not end or now <= end):
                    key = (index, item.get("date_iso") or item.get("date"), item.get("start_time"), "current")
                    if key not in self._agenda_alerted:
                        self._agenda_alerted.add(key)
                        messagebox.showinfo("Agenda ONG", f"La reunion '{label}' esta en curso.")

                if end and now > end:
                    key = (index, item.get("date_iso") or item.get("date"), item.get("end_time"), "late")
                    if key not in self._agenda_alerted:
                        self._agenda_alerted.add(key)
                        messagebox.showwarning("Agenda ONG", f"La reunion '{label}' ya paso y no esta marcada como completada.")
        finally:
            if self.winfo_exists():
                self.after(60000, self._check_agenda_alerts)


class PopupLograOpen(tk.Toplevel):
    def __init__(self, parent, form_instance):
        super().__init__(parent)
        self.form_instance = form_instance
        self.title("Abrir ONG guardado")
        self.geometry("760x420")
        self.transient(parent)
        self.grab_set()
        self.rows = []
        self._build_ui()
        self._load()

    def _build_ui(self):
        container = ttk.Frame(self, padding=12)
        container.pack(fill="both", expand=True)
        container.rowconfigure(0, weight=1)
        container.columnconfigure(0, weight=1)

        columns = ("id", "title", "status", "updated_at")
        self.tree = ttk.Treeview(container, columns=columns, show="headings", height=12)
        for col, title, width in [
            ("id", "ID", 70),
            ("title", "Titulo", 360),
            ("status", "Estado", 110),
            ("updated_at", "Actualizado", 180),
        ]:
            self.tree.heading(col, text=title)
            self.tree.column(col, width=width, anchor="w")
        self.tree.grid(row=0, column=0, sticky="nsew")

        scroll = ttk.Scrollbar(container, orient="vertical", command=self.tree.yview)
        scroll.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=scroll.set)

        actions = ttk.Frame(container)
        actions.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(10, 0))
        ttk.Button(actions, text="Cancelar", command=self.destroy).pack(side="right")
        ttk.Button(actions, text="Abrir seleccionado", command=self._open_selected).pack(side="right", padx=6)

    def _load(self):
        resp = api_client.list_logra_reports_api()
        self.rows = resp.get("data") or []
        for row in self.rows:
            self.tree.insert(
                "",
                "end",
                iid=str(row.get("id")),
                values=(
                    row.get("id"),
                    row.get("title") or "",
                    row.get("status") or "",
                    str(row.get("updated_at") or ""),
                )
            )

    def _open_selected(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("ONG", "Selecciona un reporte.")
            return
        self.form_instance.load_report(int(selected[0]))
        self.destroy()


class PopupLograAgenda(tk.Toplevel):
    MAX_ITEMS = 150
    COLUMNS = (
        "report_title",
        "date",
        "start_time",
        "end_time",
        "place",
        "person",
        "phone",
        "company",
        "topic",
        "priority",
        "status",
        "reminder_minutes",
    )
    HEADERS = {
        "report_title": "Report",
        "date": "Date",
        "start_time": "Start",
        "end_time": "End",
        "place": "Place",
        "person": "Person",
        "phone": "Phone",
        "company": "Company/Role",
        "topic": "Topic",
        "priority": "Priority",
        "status": "Status",
        "reminder_minutes": "Reminder min",
    }

    def __init__(self, parent, form_instance):
        super().__init__(parent)
        self.form_instance = form_instance
        self.title("Agenda ONG")
        self.geometry("1180x620")
        self.transient(parent)
        self.grab_set()

        self.items = [dict(item) for item in form_instance.agenda_items]
        self.selected_date = datetime.now().date()
        self.date_long_var = tk.StringVar()
        self.start_hour_var = tk.StringVar(value="09")
        self.start_minute_var = tk.StringVar(value="00")
        self.end_hour_var = tk.StringVar(value="10")
        self.end_minute_var = tk.StringVar(value="00")
        self.place_var = tk.StringVar()
        self.person_var = tk.StringVar()
        self.phone_var = tk.StringVar()
        self.company_var = tk.StringVar()
        self.topic_var = tk.StringVar()
        self.priority_var = tk.StringVar(value="Media")
        self.status_var = tk.StringVar(value="Pendiente")
        self.reminder_var = tk.StringVar(value="30")
        self.agenda_action_var = tk.StringVar(value="Nueva")
        self.export_var = tk.StringVar(value="PDF")
        self.view_mode = tk.StringVar(value="list")
        self.calendar_month = datetime.now().date().replace(day=1)
        self._build_ui()
        self._sync_long_date()
        self._render()

    def _build_ui(self):
        root = ttk.Frame(self, padding=14)
        root.pack(fill="both", expand=True)
        root.columnconfigure(0, weight=1)
        root.rowconfigure(2, weight=1)

        top = ttk.LabelFrame(root, text="Meeting details", padding=10)
        top.grid(row=0, column=0, sticky="ew")
        ttk.Label(top, text="Date").grid(row=0, column=0, sticky="w")
        ttk.Entry(top, textvariable=self.date_long_var, state="readonly", width=22).grid(
            row=0, column=1, sticky="w", padx=(4, 4), pady=3
        )
        ttk.Button(top, text="Seleccionar fecha", command=self._open_date_picker).grid(
            row=0, column=2, sticky="w", padx=(0, 14), pady=3
        )

        time_box = ttk.Frame(top)
        time_box.grid(row=0, column=3, columnspan=4, sticky="w", padx=(0, 14), pady=3)
        ttk.Label(time_box, text="Start").pack(side="left", padx=(0, 4))
        ttk.Entry(time_box, textvariable=self.start_hour_var, width=3, justify="center").pack(side="left")
        ttk.Label(time_box, text=":").pack(side="left")
        ttk.Entry(time_box, textvariable=self.start_minute_var, width=3, justify="center").pack(side="left", padx=(0, 10))
        ttk.Label(time_box, text="End").pack(side="left", padx=(0, 4))
        ttk.Entry(time_box, textvariable=self.end_hour_var, width=3, justify="center").pack(side="left")
        ttk.Label(time_box, text=":").pack(side="left")
        ttk.Entry(time_box, textvariable=self.end_minute_var, width=3, justify="center").pack(side="left")
        self._field(top, "Place", ttk.Entry(top, textvariable=self.place_var, width=24), 0, 7)
        self._field(top, "Person", ttk.Entry(top, textvariable=self.person_var), 1, 0)
        self._field(top, "Phone", ttk.Entry(top, textvariable=self.phone_var), 1, 2)
        self._field(top, "Company/Role", ttk.Entry(top, textvariable=self.company_var), 1, 4)
        self._field(top, "Topic", ttk.Entry(top, textvariable=self.topic_var), 1, 6)
        self._field(
            top,
            "Priority",
            ttk.Combobox(top, textvariable=self.priority_var, state="readonly", values=["Alta", "Media", "Baja"], width=10),
            1,
            8,
        )
        self._field(
            top,
            "Status",
            ttk.Combobox(top, textvariable=self.status_var, state="readonly", values=["Pendiente", "En proceso", "Completado"], width=12),
            2,
            0,
        )
        self._field(top, "Reminder min", ttk.Entry(top, textvariable=self.reminder_var, width=8), 2, 2)

        actions = ttk.Frame(root)
        actions.grid(row=1, column=0, sticky="ew", pady=(10, 8))
        actions.columnconfigure(0, weight=1)
        actions.columnconfigure(1, weight=1)
        actions.columnconfigure(2, weight=1)

        agenda_actions = ttk.LabelFrame(actions, text="Agenda", padding=(8, 6))
        agenda_actions.grid(row=0, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(agenda_actions, text="Buscar", command=self._search_backend).pack(side="left", padx=(0, 4))
        self.view_button = ttk.Button(agenda_actions, text="Calendario", command=self._toggle_view)
        self.view_button.pack(side="left", padx=(0, 4))
        ttk.Button(agenda_actions, text="Guardar", command=self._save).pack(side="left", padx=(8, 0))

        meeting_actions = ttk.LabelFrame(actions, text="Reunion seleccionada", padding=(8, 6))
        meeting_actions.grid(row=0, column=1, sticky="ew", padx=(0, 8))
        ttk.Combobox(
            meeting_actions,
            textvariable=self.agenda_action_var,
            state="readonly",
            width=22,
            values=[
                "Nueva",
                "Cargar seleccion",
                "Actualizar seleccion",
                "Cambiar status",
                "Notas",
                "Eliminar",
            ],
        ).pack(side="left", padx=(0, 4))
        ttk.Button(meeting_actions, text="Ejecutar", command=self._run_agenda_action).pack(side="left")

        export_actions = ttk.LabelFrame(actions, text="Exportacion", padding=(8, 6))
        export_actions.grid(row=0, column=2, sticky="e")
        ttk.Combobox(export_actions, textvariable=self.export_var, state="readonly", width=8, values=["PDF", "Excel", "Word"]).pack(side="left")
        ttk.Button(export_actions, text="Exportar", command=self._export_selected).pack(side="left", padx=4)
        ttk.Button(export_actions, text="Cerrar", command=self.destroy).pack(side="left", padx=(8, 0))

        self.table_box = ttk.Frame(root)
        self.table_box.grid(row=2, column=0, sticky="nsew")
        self.table_box.rowconfigure(0, weight=1)
        self.table_box.columnconfigure(0, weight=1)

        self.tree = ttk.Treeview(self.table_box, columns=self.COLUMNS, show="headings", height=16)
        widths = {
            "report_title": 240,
            "date": 150,
            "start_time": 70,
            "end_time": 70,
            "place": 140,
            "person": 150,
            "phone": 120,
            "company": 140,
            "topic": 300,
            "priority": 90,
            "status": 110,
            "reminder_minutes": 100,
        }
        for col in self.COLUMNS:
            self.tree.heading(col, text=self.HEADERS[col])
            self.tree.column(col, width=widths[col], anchor="w")
        self.tree.grid(row=0, column=0, sticky="nsew")
        yscroll = ttk.Scrollbar(self.table_box, orient="vertical", command=self.tree.yview)
        yscroll.grid(row=0, column=1, sticky="ns")
        xscroll = ttk.Scrollbar(self.table_box, orient="horizontal", command=self.tree.xview)
        xscroll.grid(row=1, column=0, sticky="ew")
        self.tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)
        self.tree.bind("<<TreeviewSelect>>", lambda _event: self._load_selected_into_form(silent=True))
        self.tree.bind("<Double-1>", lambda _event: self._load_selected_into_form(silent=True))
        self.tree.tag_configure("Alta", background="#F8D7DA")
        self.tree.tag_configure("Media", background="#FFF3CD")
        self.tree.tag_configure("Baja", background="#D1E7DD")
        self.tree.tag_configure("Pendiente", background="#F8D7DA")
        self.tree.tag_configure("En proceso", background="#FFF3CD")
        self.tree.tag_configure("Completado", background="#D1E7DD")

        self.calendar_box = ttk.Frame(root)
        self.calendar_box.grid(row=2, column=0, sticky="nsew")
        self.calendar_box.columnconfigure(0, weight=1)
        self.calendar_box.rowconfigure(1, weight=1)
        self._build_calendar_shell()
        self.table_box.tkraise()

    def _field(self, parent, label, widget, row, col):
        ttk.Label(parent, text=label).grid(row=row, column=col, sticky="w", padx=(0, 4), pady=3)
        widget.grid(row=row, column=col + 1, sticky="ew", padx=(0, 12), pady=3)

    def _build_calendar_shell(self):
        header = ttk.Frame(self.calendar_box)
        header.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        header.columnconfigure(1, weight=1)
        ttk.Button(header, text="Anterior", command=lambda: self._move_calendar_month(-1)).grid(row=0, column=0, sticky="w")
        self.calendar_title = ttk.Label(header, text="", anchor="center", font=("Segoe UI", 13, "bold"))
        self.calendar_title.grid(row=0, column=1, sticky="ew")
        ttk.Button(header, text="Siguiente", command=lambda: self._move_calendar_month(1)).grid(row=0, column=2, sticky="e")

        self.calendar_grid = ttk.Frame(self.calendar_box)
        self.calendar_grid.grid(row=1, column=0, sticky="nsew")
        for col in range(7):
            self.calendar_grid.columnconfigure(col, weight=1, uniform="calendar")
        for row in range(7):
            self.calendar_grid.rowconfigure(row, weight=1)

    def _toggle_view(self):
        if self.view_mode.get() == "list":
            self.view_mode.set("calendar")
            self.view_button.configure(text="Lista")
            self.calendar_box.tkraise()
            self._render_calendar()
        else:
            self.view_mode.set("list")
            self.view_button.configure(text="Calendario")
            self.table_box.tkraise()

    def _move_calendar_month(self, delta):
        year = self.calendar_month.year
        month = self.calendar_month.month + delta
        while month < 1:
            month += 12
            year -= 1
        while month > 12:
            month -= 12
            year += 1
        self.calendar_month = self.calendar_month.replace(year=year, month=month, day=1)
        self._render_calendar()

    def _sync_long_date(self):
        try:
            value = self.selected_date.strftime("%B %d, %Y")
            value = value.replace(" 0", " ")
            self.date_long_var.set(value)
        except Exception:
            self.date_long_var.set("")

    def _open_date_picker(self):
        popup = tk.Toplevel(self)
        popup.title("Seleccionar fecha")
        popup.geometry("310x290")
        popup.transient(self)
        popup.grab_set()

        cal = Calendar(
            popup,
            selectmode="day",
            locale="en_US",
            year=self.selected_date.year,
            month=self.selected_date.month,
            day=self.selected_date.day,
        )
        cal.pack(fill="both", expand=True, padx=10, pady=10)

        actions = ttk.Frame(popup)
        actions.pack(fill="x", padx=10, pady=(0, 10))

        def accept():
            self.selected_date = cal.selection_get()
            self._sync_long_date()
            popup.destroy()

        ttk.Button(actions, text="Cancelar", command=popup.destroy).pack(side="right")
        ttk.Button(actions, text="Seleccionar", command=accept).pack(side="right", padx=6)

    def _valid_time(self, value):
        try:
            datetime.strptime(value.strip(), "%H:%M")
            return True
        except Exception:
            return False

    def _two_digit(self, value, max_value):
        text = str(value or "").strip()
        if not text.isdigit():
            return None
        number = int(text)
        if number < 0 or number > max_value:
            return None
        return f"{number:02d}"

    def _compose_time(self, hour_var, minute_var):
        hour = self._two_digit(hour_var.get(), 23)
        minute = self._two_digit(minute_var.get(), 59)
        if hour is None or minute is None:
            return None
        return f"{hour}:{minute}"

    def _agenda_item_date(self, item):
        date_value = str(item.get("date_iso") or item.get("date") or "").strip()
        for date_format in ("%Y-%m-%d", "%B %d, %Y"):
            try:
                return datetime.strptime(date_value, date_format).date()
            except Exception:
                continue
        return None

    def _calendar_color(self, item):
        return {
            "Pendiente": "#F8D7DA",
            "En proceso": "#FFF3CD",
            "Completado": "#D1E7DD",
            "Alta": "#F8D7DA",
            "Media": "#FFF3CD",
            "Baja": "#D1E7DD",
        }.get(item.get("status") or item.get("priority"), "#FFFFFF")

    def _select_calendar_item(self, index):
        if 0 <= index < len(self.items):
            self.tree.selection_set(str(index))
            self.tree.see(str(index))
            item_date = self._agenda_item_date(self.items[index])
            if item_date:
                self.selected_date = item_date
                self._sync_long_date()

    def _render_calendar(self):
        if not hasattr(self, "calendar_grid"):
            return
        for widget in self.calendar_grid.winfo_children():
            widget.destroy()

        self.calendar_title.configure(text=self.calendar_month.strftime("%B %Y"))
        weekdays = ["Sun", "Mon", "Tue", "Wed", "Thu", "Fri", "Sat"]
        for col, day in enumerate(weekdays):
            ttk.Label(self.calendar_grid, text=day, anchor="center", font=("Segoe UI", 9, "bold")).grid(
                row=0, column=col, sticky="ew", padx=1, pady=(0, 3)
            )

        by_date = {}
        for index, item in enumerate(self.items):
            item_date = self._agenda_item_date(item)
            if not item_date:
                continue
            by_date.setdefault(item_date, []).append((index, item))

        first = self.calendar_month
        cursor = first - timedelta(days=(first.weekday() + 1) % 7)
        for cell in range(42):
            day = cursor + timedelta(days=cell)
            row = (cell // 7) + 1
            col = cell % 7
            frame = tk.Frame(
                self.calendar_grid,
                bd=1,
                relief="solid",
                background="#FFFFFF" if day.month == self.calendar_month.month else "#F2F4F7",
                padx=4,
                pady=3,
            )
            frame.grid(row=row, column=col, sticky="nsew", padx=1, pady=1)
            frame.columnconfigure(0, weight=1)
            tk.Label(
                frame,
                text=str(day.day),
                background=frame.cget("background"),
                foreground="#101828" if day.month == self.calendar_month.month else "#98A2B3",
                font=("Segoe UI", 9, "bold"),
                anchor="w",
            ).grid(row=0, column=0, sticky="ew")

            for meeting_row, (index, item) in enumerate(by_date.get(day, [])[:3], start=1):
                text = f"{item.get('start_time') or '--:--'} {item.get('topic') or item.get('person') or 'Meeting'}"
                tk.Button(
                    frame,
                    text=text,
                    anchor="w",
                    relief="flat",
                    background=self._calendar_color(item),
                    command=lambda i=index: self._select_calendar_item(i),
                    font=("Segoe UI", 8),
                ).grid(row=meeting_row, column=0, sticky="ew", pady=(2, 0))
            extra = len(by_date.get(day, [])) - 3
            if extra > 0:
                tk.Label(
                    frame,
                    text=f"+{extra} mas",
                    background=frame.cget("background"),
                    foreground="#003A75",
                    font=("Segoe UI", 8, "bold"),
                ).grid(row=4, column=0, sticky="w", pady=(2, 0))

    def _render(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        for index, item in enumerate(self.items):
            tag = item.get("status") or item.get("priority") or ""
            self.tree.insert(
                "",
                "end",
                iid=str(index),
                values=tuple(item.get(col, "") for col in self.COLUMNS),
                tags=(tag,)
            )
        if getattr(self, "view_mode", None) and self.view_mode.get() == "calendar":
            self._render_calendar()

    def _agenda_item_key(self, item):
        return "|".join([
            str(item.get("date_iso") or item.get("date") or "").strip().lower(),
            str(item.get("start_time") or "").strip().lower(),
            str(item.get("end_time") or "").strip().lower(),
            str(item.get("place") or "").strip().lower(),
            str(item.get("person") or "").strip().lower(),
            str(item.get("phone") or item.get("telefono") or "").strip().lower(),
            str(item.get("company") or item.get("company_role") or "").strip().lower(),
            str(item.get("topic") or "").strip().lower(),
        ])

    def _run_agenda_action(self):
        action = self.agenda_action_var.get()
        handlers = {
            "Nueva": self._add,
            "Cargar seleccion": self._load_selected_into_form,
            "Actualizar seleccion": self._update_selected,
            "Cambiar status": self._change_selected_status,
            "Notas": self._open_notes,
            "Eliminar": self._remove,
        }
        handlers.get(action, self._add)()

    def _search_backend(self):
        listing = api_client.list_logra_reports_api()
        if listing.get("success") is False:
            messagebox.showerror("Agenda ONG", f"No se pudo buscar la agenda:\n{listing.get('error') or listing}")
            return

        rows = listing.get("data") or []
        if not rows:
            self.items = []
            self.form_instance.agenda_items = []
            self._render()
            messagebox.showinfo("Agenda ONG", "No hay agendas ONG guardadas en backend.")
            return

        all_items = []
        seen = set()
        skipped = 0
        for report in rows:
            report_title = report.get("title") or f"ONG #{report.get('id')}"
            for agenda_index, item in enumerate(report.get("agenda_items") or []):
                if not isinstance(item, dict):
                    continue
                row = dict(item)
                key = self._agenda_item_key(row)
                if key in seen:
                    skipped += 1
                    continue
                seen.add(key)
                row["report_id"] = report.get("id")
                row["agenda_index"] = agenda_index
                row["report_title"] = report_title
                all_items.append(row)

        self.items = all_items
        self.form_instance.agenda_items = [dict(item) for item in self.items]
        self._render()
        extra = f" Duplicadas omitidas: {skipped}." if skipped else ""
        messagebox.showinfo("Agenda ONG", f"Agendas cargadas desde backend. Lineas: {len(self.items)}.{extra}")

    def _add(self):
        if len(self.items) >= self.MAX_ITEMS:
            messagebox.showwarning("Agenda ONG", "La agenda permite maximo 150 lineas.")
            return
        self._sync_long_date()
        start = self._compose_time(self.start_hour_var, self.start_minute_var)
        end = self._compose_time(self.end_hour_var, self.end_minute_var)
        if not start or not end:
            messagebox.showwarning("Agenda ONG", "Usa hora 00-23 y minutos 00-59.")
            return
        try:
            reminder = max(0, int(self.reminder_var.get() or 0))
        except Exception:
            messagebox.showwarning("Agenda ONG", "Reminder min debe ser un numero entero.")
            return
        item = {
            "report_title": f"ONG - {self.form_instance.form_var.get() or 'Cuestionarios'}",
            "date": self.date_long_var.get(),
            "date_iso": self.selected_date.isoformat(),
            "start_time": start,
            "end_time": end,
            "place": self.place_var.get().strip(),
            "person": self.person_var.get().strip(),
            "phone": self.phone_var.get().strip(),
            "company": self.company_var.get().strip(),
            "topic": self.topic_var.get().strip(),
            "priority": self.priority_var.get(),
            "status": self.status_var.get(),
            "reminder_minutes": reminder,
        }
        if not item["person"] and not item["topic"] and not item["place"]:
            messagebox.showwarning("Agenda ONG", "Agrega al menos persona, tema o lugar.")
            return
        self.items.append(item)
        self._render()

    def _selected_index(self, silent=False):
        selected = self.tree.selection()
        if not selected:
            if not silent:
                messagebox.showwarning("Agenda ONG", "Selecciona una linea.")
            return None
        try:
            index = int(selected[0])
        except Exception:
            if not silent:
                messagebox.showwarning("Agenda ONG", "Seleccion invalida.")
            return None
        if index < 0 or index >= len(self.items):
            if not silent:
                messagebox.showwarning("Agenda ONG", "Seleccion invalida.")
            return None
        return index

    def _load_selected_into_form(self, silent=False):
        index = self._selected_index(silent=silent)
        if index is None:
            return
        item = self.items[index]
        item_date = self._agenda_item_date(item)
        if item_date:
            self.selected_date = item_date
            self._sync_long_date()
        start = str(item.get("start_time") or "09:00")
        end = str(item.get("end_time") or "10:00")
        self.start_hour_var.set(start[:2] if len(start) >= 2 else "09")
        self.start_minute_var.set(start[3:5] if len(start) >= 5 else "00")
        self.end_hour_var.set(end[:2] if len(end) >= 2 else "10")
        self.end_minute_var.set(end[3:5] if len(end) >= 5 else "00")
        self.place_var.set(item.get("place") or "")
        self.person_var.set(item.get("person") or "")
        self.phone_var.set(item.get("phone") or item.get("telefono") or "")
        self.company_var.set(item.get("company") or item.get("company_role") or "")
        self.topic_var.set(item.get("topic") or "")
        self.priority_var.set(item.get("priority") or "Media")
        self.status_var.set(item.get("status") or "Pendiente")
        self.reminder_var.set(str(item.get("reminder_minutes") or 30))

    def _agenda_payload_from_fields(self, current_item=None):
        self._sync_long_date()
        start = self._compose_time(self.start_hour_var, self.start_minute_var)
        end = self._compose_time(self.end_hour_var, self.end_minute_var)
        if not start or not end:
            messagebox.showwarning("Agenda ONG", "Usa hora 00-23 y minutos 00-59.")
            return None
        try:
            reminder = max(0, int(self.reminder_var.get() or 0))
        except Exception:
            messagebox.showwarning("Agenda ONG", "Reminder min debe ser un numero entero.")
            return None
        payload = dict(current_item or {})
        payload.update({
            "date": self.date_long_var.get(),
            "date_iso": self.selected_date.isoformat(),
            "start_time": start,
            "end_time": end,
            "place": self.place_var.get().strip(),
            "person": self.person_var.get().strip(),
            "phone": self.phone_var.get().strip(),
            "company": self.company_var.get().strip(),
            "topic": self.topic_var.get().strip(),
            "priority": self.priority_var.get(),
            "status": self.status_var.get(),
            "reminder_minutes": reminder,
        })
        if not payload["person"] and not payload["topic"] and not payload["place"]:
            messagebox.showwarning("Agenda ONG", "Agrega al menos persona, tema o lugar.")
            return None
        return payload

    def _update_selected(self):
        index = self._selected_index()
        if index is None:
            return
        payload = self._agenda_payload_from_fields(self.items[index])
        if payload is None:
            return
        report_id = payload.get("report_id") or self.form_instance.report_id
        agenda_index = payload.get("agenda_index", index)
        if not report_id:
            self.items[index] = payload
            self.form_instance.agenda_items = [dict(item) for item in self.items]
            self._render()
            messagebox.showinfo("Agenda ONG", "Linea actualizada localmente. Guarda la agenda para persistirla.")
            return
        resp = api_client.update_logra_agenda_item_api(report_id, agenda_index, payload)
        if not resp.get("success"):
            messagebox.showerror("Agenda ONG", f"No se pudo actualizar la linea:\n{resp.get('error') or resp}")
            return
        self.items[index] = payload
        self.form_instance.agenda_items = [dict(item) for item in self.items]
        self._render()
        messagebox.showinfo("Agenda ONG", "Linea actualizada correctamente.")

    def _change_selected_status(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Agenda ONG", "Selecciona una linea.")
            return
        for iid in selected:
            index = int(iid)
            if 0 <= index < len(self.items):
                self.items[index]["status"] = self.status_var.get()
                report_id = self.items[index].get("report_id") or self.form_instance.report_id
                agenda_index = self.items[index].get("agenda_index", index)
                if report_id:
                    resp = api_client.update_logra_agenda_item_api(report_id, agenda_index, self.items[index])
                    if not resp.get("success"):
                        messagebox.showerror("Agenda ONG", f"No se pudo actualizar el status:\n{resp.get('error') or resp}")
                        return
        self.form_instance.agenda_items = [dict(item) for item in self.items]
        self._render()

    def _open_notes(self):
        popup = tk.Toplevel(self)
        popup.title("Anotaciones generales ONG")
        popup.geometry("720x420")
        popup.transient(self)
        popup.grab_set()

        root = ttk.Frame(popup, padding=12)
        root.pack(fill="both", expand=True)
        root.rowconfigure(0, weight=1)
        root.columnconfigure(0, weight=1)
        text = ScrolledText(root, wrap="word", font=("Segoe UI", 10))
        text.grid(row=0, column=0, sticky="nsew")
        text.insert("1.0", self.form_instance.agenda_notes or "")

        actions = ttk.Frame(root)
        actions.grid(row=1, column=0, sticky="ew", pady=(10, 0))

        def save_notes():
            self.form_instance.agenda_notes = text.get("1.0", "end-1c").strip()
            popup.destroy()

        ttk.Button(actions, text="Cancelar", command=popup.destroy).pack(side="right")
        ttk.Button(actions, text="Guardar anotaciones", command=save_notes).pack(side="right", padx=6)

    def _remove(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Agenda ONG", "Selecciona una linea.")
            return
        if not messagebox.askyesno("Agenda ONG", "Deseas eliminar la linea seleccionada de la agenda?"):
            return

        for index in sorted((int(i) for i in selected), reverse=True):
            if 0 <= index < len(self.items):
                item = self.items[index]
                report_id = item.get("report_id")
                agenda_index = item.get("agenda_index")
                if report_id is not None and agenda_index is not None:
                    resp = api_client.delete_logra_agenda_item_api(report_id, agenda_index)
                    if not resp.get("success"):
                        messagebox.showerror(
                            "Agenda ONG",
                            f"No se pudo eliminar en backend:\n{resp.get('error') or resp}"
                        )
                        continue
                self.items.pop(index)
        self._render()
        messagebox.showinfo("Agenda ONG", "Linea eliminada correctamente.")

    def _save(self):
        self.form_instance.agenda_items = [dict(item) for item in self.items]
        if len(self.form_instance.agenda_items) > self.MAX_ITEMS:
            messagebox.showwarning("Agenda ONG", "La agenda permite maximo 150 lineas.")
            return
        if self.form_instance._save_report(silent=False):
            messagebox.showinfo("Agenda ONG", "Agenda guardada correctamente.")

    def _rows_for_export(self):
        return [[item.get(col, "") for col in self.COLUMNS] for item in self.items]

    def _export_selected(self):
        selected = self.export_var.get()
        if selected == "Excel":
            self._export_excel()
        elif selected == "Word":
            self._export_word()
        else:
            self._export_pdf()

    def _export_excel(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            initialfile="ONG_agenda.xlsx",
            filetypes=[("Excel", "*.xlsx")]
        )
        if not path:
            return
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = "ONG Agenda"
        if self.form_instance.agenda_notes:
            ws.append(["General notes", self.form_instance.agenda_notes])
            ws.append([])
        ws.append([self.HEADERS[col] for col in self.COLUMNS])
        for row in self._rows_for_export():
            ws.append(row)
        header_fill = PatternFill("solid", fgColor="003B71")
        header_row = 3 if self.form_instance.agenda_notes else 1
        for cell in ws[header_row]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")
        colors = {"Alta": "F8D7DA", "Media": "FFF3CD", "Baja": "D1E7DD", "Pendiente": "F8D7DA", "En proceso": "FFF3CD", "Completado": "D1E7DD"}
        priority_index = self.COLUMNS.index("priority")
        status_index = self.COLUMNS.index("status")
        for row in ws.iter_rows(min_row=header_row + 1):
            fill = colors.get(row[status_index].value) or colors.get(row[priority_index].value)
            if fill:
                for cell in row:
                    cell.fill = PatternFill("solid", fgColor=fill)
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        widths = [28, 18, 10, 10, 22, 22, 18, 22, 42, 12, 14, 14]
        for idx, width in enumerate(widths, start=1):
            ws.column_dimensions[chr(64 + idx)].width = width
        wb.save(path)
        messagebox.showinfo("Agenda ONG", "Excel generado correctamente.")

    def _export_word(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile="ONG_agenda.docx",
            filetypes=[("Word", "*.docx")]
        )
        if not path:
            return
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.section import WD_ORIENT

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        title = doc.add_paragraph()
        run = title.add_run("ONG Meeting Agenda")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 59, 113)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M').replace(' 0', ' ')}")
        if self.form_instance.agenda_notes:
            notes = doc.add_paragraph()
            notes.add_run("General notes: ").bold = True
            notes.add_run(self.form_instance.agenda_notes)
        table = doc.add_table(rows=1, cols=len(self.COLUMNS))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for idx, col in enumerate(self.COLUMNS):
            hdr[idx].text = self.HEADERS[col]
        for item in self.items:
            cells = table.add_row().cells
            for idx, col in enumerate(self.COLUMNS):
                cells[idx].text = str(item.get(col, ""))
        doc.save(path)
        messagebox.showinfo("Agenda ONG", "Word generado correctamente.")

    def _export_pdf(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            initialfile="ONG_agenda.pdf",
            filetypes=[("PDF", "*.pdf")]
        )
        if not path:
            return
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(path, pagesize=landscape(letter), leftMargin=28, rightMargin=28, topMargin=28, bottomMargin=28)
        story = [
            Paragraph("ONG Meeting Agenda", styles["Title"]),
            Paragraph(datetime.now().strftime("Generated: %B %d, %Y %H:%M").replace(" 0", " "), styles["Normal"]),
            Spacer(1, 10),
        ]
        if self.form_instance.agenda_notes:
            story.extend([
                Paragraph(f"<b>General notes:</b> {self.form_instance.agenda_notes}", styles["Normal"]),
                Spacer(1, 10),
            ])
        data = [[self.HEADERS[col] for col in self.COLUMNS]] + self._rows_for_export()
        table = Table(data, repeatRows=1)
        style = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003B71")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#A6A6A6")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
        ])
        for row_idx, item in enumerate(self.items, start=1):
            fill = {"Alta": "#F8D7DA", "Media": "#FFF3CD", "Baja": "#D1E7DD", "Pendiente": "#F8D7DA", "En proceso": "#FFF3CD", "Completado": "#D1E7DD"}.get(item.get("status")) or "#FFFFFF"
            style.add("BACKGROUND", (0, row_idx), (-1, row_idx), colors.HexColor(fill))
        table.setStyle(style)
        story.append(table)
        doc.build(story)
        messagebox.showinfo("Agenda ONG", "PDF generado correctamente.")


class PopupLograAgendaSearch(tk.Toplevel):
    def __init__(self, parent, agenda_popup):
        super().__init__(parent)
        self.agenda_popup = agenda_popup
        self.title("Buscar agenda ONG")
        self.geometry("760x380")
        self.transient(parent)
        self.grab_set()
        self.rows = []
        self._build_ui()
        self._load()

    def _build_ui(self):
        root = ttk.Frame(self, padding=12)
        root.pack(fill="both", expand=True)
        root.rowconfigure(0, weight=1)
        root.columnconfigure(0, weight=1)

        self.tree = ttk.Treeview(root, columns=("id", "title", "agenda", "updated_at"), show="headings", height=10)
        for col, title, width in [
            ("id", "ID", 70),
            ("title", "Titulo", 380),
            ("agenda", "Lineas agenda", 120),
            ("updated_at", "Actualizado", 170),
        ]:
            self.tree.heading(col, text=title)
            self.tree.column(col, width=width, anchor="w")
        self.tree.grid(row=0, column=0, sticky="nsew")
        self.tree.bind("<Double-1>", lambda e: self._select())

        actions = ttk.Frame(root)
        actions.grid(row=1, column=0, sticky="ew", pady=(10, 0))
        ttk.Button(actions, text="Cancelar", command=self.destroy).pack(side="right")
        ttk.Button(actions, text="Cargar agenda", command=self._select).pack(side="right", padx=6)

    def _load(self):
        resp = api_client.list_logra_reports_api()
        self.rows = resp.get("data") or []
        for row in self.rows:
            self.tree.insert(
                "",
                "end",
                iid=str(row.get("id")),
                values=(
                    row.get("id"),
                    row.get("title") or "",
                    len(row.get("agenda_items") or []),
                    str(row.get("updated_at") or ""),
                )
            )

    def _select(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Agenda ONG", "Selecciona una agenda.")
            return
        report_id = int(selected[0])
        resp = api_client.get_logra_report_api(report_id)
        if resp.get("success") is False:
            messagebox.showerror("Agenda ONG", f"No se pudo cargar:\n{resp.get('error') or resp}")
            return
        report = resp.get("report") or {}
        self.agenda_popup.form_instance.report_id = report_id
        self.agenda_popup.items = [dict(item) for item in (report.get("agenda_items") or [])]
        self.agenda_popup.form_instance.agenda_items = [dict(item) for item in self.agenda_popup.items]
        self.agenda_popup.form_instance.agenda_notes = report.get("agenda_notes") or ""
        self.agenda_popup._render()
        self.destroy()


class PopupLograPortia(tk.Toplevel):
    def __init__(self, parent, form_instance):
        super().__init__(parent)
        self.form_instance = form_instance
        self.title("Mejorar con PORTIA - ONG")
        self.geometry("720x520")
        self.transient(parent)
        self.grab_set()

        self.form_var = tk.StringVar(value=form_instance.form_var.get())
        self.section_var = tk.StringVar(value=form_instance.section_var.get())
        self.question_var = tk.StringVar()
        self.bullet_var = tk.StringVar()
        self.language_var = tk.StringVar(value="ES")
        self.question_items = []
        self._build_ui()
        self._load_questions()

    def _build_ui(self):
        root = ttk.Frame(self, padding=14)
        root.pack(fill="both", expand=True)
        root.columnconfigure(1, weight=1)
        root.rowconfigure(5, weight=1)

        ttk.Label(root, text="Formulario").grid(row=0, column=0, sticky="w", pady=4)
        form_cb = ttk.Combobox(
            root,
            textvariable=self.form_var,
            state="readonly",
            values=[item["title"] for item in ONG_QUESTIONNAIRES],
        )
        form_cb.grid(row=0, column=1, sticky="ew", pady=4)
        form_cb.bind("<<ComboboxSelected>>", lambda e: self._load_questions())

        ttk.Label(root, text="Tipo").grid(row=1, column=0, sticky="w", pady=4)
        section_cb = ttk.Combobox(
            root,
            textvariable=self.section_var,
            state="readonly",
            values=list(LograQuestionnairesForm.SECTION_LABELS.keys()),
        )
        section_cb.grid(row=1, column=1, sticky="ew", pady=4)
        section_cb.bind("<<ComboboxSelected>>", lambda e: self._load_questions())

        ttk.Label(root, text="Pregunta").grid(row=2, column=0, sticky="w", pady=4)
        self.question_cb = ttk.Combobox(root, textvariable=self.question_var, state="readonly")
        self.question_cb.grid(row=2, column=1, sticky="ew", pady=4)
        self.question_cb.bind("<<ComboboxSelected>>", lambda e: self._load_bullets())

        ttk.Label(root, text="Bullet").grid(row=3, column=0, sticky="w", pady=4)
        self.bullet_cb = ttk.Combobox(root, textvariable=self.bullet_var, state="readonly")
        self.bullet_cb.grid(row=3, column=1, sticky="ew", pady=4)

        lang = ttk.LabelFrame(root, text="Salida")
        lang.grid(row=4, column=1, sticky="w", pady=8)
        ttk.Radiobutton(lang, text="Espanol", value="ES", variable=self.language_var).pack(side="left", padx=8)
        ttk.Radiobutton(lang, text="Ingles", value="EN", variable=self.language_var).pack(side="left", padx=8)

        self.preview = ScrolledText(root, height=9, wrap="word", font=("Segoe UI", 9))
        self.preview.grid(row=5, column=0, columnspan=2, sticky="nsew", pady=(8, 10))

        actions = ttk.Frame(root)
        actions.grid(row=6, column=0, columnspan=2, sticky="ew")
        ttk.Button(actions, text="Cancelar", command=self.destroy).pack(side="right")
        ttk.Button(actions, text="Mejorar con PORTIA", command=self._execute).pack(side="right", padx=6)

    def _selected_form(self):
        return next((item for item in ONG_QUESTIONNAIRES if item["title"] == self.form_var.get()), ONG_QUESTIONNAIRES[0])

    def _load_questions(self):
        form = self._selected_form()
        section = self.section_var.get()
        self.question_items = form.get(section, [])
        values = []
        for item in self.question_items:
            label = f"{item.get('id') or item.get('number')} - {item.get('question', '')[:120]}"
            values.append(label)
        self.question_cb.configure(values=values)
        if values:
            self.question_var.set(values[0])
        else:
            self.question_var.set("")
        self._load_bullets()

    def _selected_question(self):
        value = self.question_var.get()
        values = list(self.question_cb.cget("values") or [])
        try:
            index = values.index(value)
            return self.question_items[index]
        except Exception:
            return None

    def _load_bullets(self):
        form = self._selected_form()
        section = self.section_var.get()
        item = self._selected_question()
        if not item:
            self.bullet_cb.configure(values=[])
            self.bullet_var.set("")
            return
        bullets = self.form_instance._get_bullets(form, section, item)
        values = [f"{idx}. {text[:100] if text else '(vacio)'}" for idx, text in enumerate(bullets, start=1)]
        self.bullet_cb.configure(values=values)
        self.bullet_var.set(values[0] if values else "")

    def _execute(self):
        form = self._selected_form()
        section = self.section_var.get()
        item = self._selected_question()
        if not item:
            messagebox.showwarning("PORTIA", "Selecciona una pregunta.")
            return

        values = list(self.bullet_cb.cget("values") or [])
        try:
            bullet_index = values.index(self.bullet_var.get())
        except ValueError:
            bullet_index = 0

        bullets = self.form_instance._get_bullets(form, section, item)
        current_text = bullets[bullet_index].strip() if bullet_index < len(bullets) else ""
        if not current_text:
            messagebox.showwarning("PORTIA", "El bullet seleccionado esta vacio.")
            return

        payload = {
            "text": current_text,
            "language": self.language_var.get(),
            "report_type": LograQuestionnairesForm.SECTION_LABELS.get(section, section),
            "form_title": form["title"],
            "question": item.get("question", ""),
        }
        resp = api_client.improve_logra_text_api(payload)
        if not resp.get("success"):
            messagebox.showerror("PORTIA", f"No se pudo mejorar el texto:\n{resp.get('error') or resp}")
            return

        improved = resp.get("text") or ""
        self.preview.delete("1.0", "end")
        self.preview.insert("1.0", improved)

        PopupAICompare(
            self,
            original_text=current_text,
            ai_text=improved,
            on_accept=lambda: self._accept_text(form, section, item, bullet_index, improved),
            on_retry=self._execute,
        )

    def _accept_text(self, form, section, item, bullet_index, text):
        bullets = self.form_instance._get_bullets(form, section, item)
        while len(bullets) <= bullet_index:
            bullets.append("")
        bullets[bullet_index] = text
        self.form_instance._render_current_page()
        self.destroy()

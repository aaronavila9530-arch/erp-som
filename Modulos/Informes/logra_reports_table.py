import tkinter as tk
from datetime import datetime
from html import escape
from tkinter import filedialog, ttk, messagebox

import api_client
from Modulos.Informes.logra_questionnaires_data import LOGRA_QUESTIONNAIRES


class LograReportsTable(ttk.Frame):
    PAGE_SIZE = 50

    def __init__(self, parent):
        super().__init__(parent)
        self.rows = []
        self.row_map = {}
        self._build_ui()
        self._load()

    def _build_ui(self):
        self.pack(fill="both", expand=True)
        top = ttk.Frame(self)
        top.pack(fill="x", padx=10, pady=10)

        ttk.Button(top, text="Buscar", command=self._load).pack(side="left")
        ttk.Label(top, text="Acciones").pack(side="left", padx=(10, 4))
        self.action_var = tk.StringVar(value="Revisar")
        ttk.Combobox(
            top,
            textvariable=self.action_var,
            state="readonly",
            width=18,
            values=["Revisar", "Exportar Word", "Exportar PDF"]
        ).pack(side="left")
        ttk.Button(top, text="Ejecutar", command=self._run_action).pack(side="left", padx=6)
        self.info_label = ttk.Label(top, text="")
        self.info_label.pack(side="right")

        columns = ("id", "title", "status", "agenda", "attachments", "updated_at")
        self.tree = ttk.Treeview(self, columns=columns, show="headings", height=12)
        headers = {
            "id": "ID",
            "title": "Titulo",
            "status": "Status",
            "agenda": "Agenda",
            "attachments": "Adjuntos",
            "updated_at": "Actualizado",
        }
        widths = {
            "id": 70,
            "title": 360,
            "status": 120,
            "agenda": 90,
            "attachments": 90,
            "updated_at": 180,
        }
        for col in columns:
            self.tree.heading(col, text=headers[col])
            self.tree.column(col, width=widths[col], anchor="w")
        self.tree.pack(fill="both", expand=True, padx=10, pady=(0, 8), side="top")
        self.tree.bind("<Double-1>", lambda e: self._review_selected())

    def _load(self):
        resp = api_client.list_logra_reports_api()
        self.rows = resp.get("data") or []
        self.row_map = {str(row.get("id")): row for row in self.rows}
        for item in self.tree.get_children():
            self.tree.delete(item)
        for row in self.rows:
            agenda_count = len(row.get("agenda_items") or [])
            self.tree.insert(
                "",
                "end",
                iid=str(row.get("id")),
                values=(
                    row.get("id"),
                    row.get("title") or "",
                    row.get("status") or "",
                    agenda_count,
                    row.get("attachment_count") or 0,
                    str(row.get("updated_at") or ""),
                ),
            )
        self.info_label.config(text=f"Resultados: {len(self.rows)}")

    def _selected_report_id(self):
        selected = self.tree.selection()
        return int(selected[0]) if selected else None

    def _run_action(self):
        action = self.action_var.get()
        if action == "Exportar Word":
            self._export_selected_word()
        elif action == "Exportar PDF":
            self._export_selected_pdf()
        else:
            self._review_selected()

    def _review_selected(self):
        report_id = self._selected_report_id()
        if not report_id:
            messagebox.showwarning("LOGRA", "Selecciona un reporte LOGRA.")
            return
        from Modulos.Informes.logra_questionnaires_form import LograQuestionnairesForm

        for widget in self.master.winfo_children():
            widget.destroy()
        form = LograQuestionnairesForm(self.master, review_mode=True)
        form.load_report(report_id)

    def _selected_payload(self):
        report_id = self._selected_report_id()
        if not report_id:
            messagebox.showwarning("LOGRA", "Selecciona un reporte LOGRA.")
            return None
        resp = api_client.get_logra_report_api(report_id)
        if resp.get("error") or not resp.get("report"):
            messagebox.showerror("LOGRA", f"No se pudo cargar el reporte:\n{resp.get('error') or resp}")
            return None
        return resp

    def _ordered_answers(self, payload):
        answers = {
            (
                item.get("form_slug"),
                item.get("section"),
                str(item.get("item_key") or "")
            ): item
            for item in payload.get("answers") or []
        }
        ordered = []
        for form in LOGRA_QUESTIONNAIRES:
            for section in ("critical_questions", "detailed_questions"):
                for question in form.get(section, []):
                    key = (
                        form.get("slug"),
                        section,
                        str(question.get("id") or question.get("number") or "").strip()
                    )
                    answer = answers.get(key)
                    if answer:
                        ordered.append((form, section, question, answer))
        return ordered

    def _attachments_by_question(self, payload):
        grouped = {}
        for att in payload.get("attachments") or []:
            key = (
                att.get("form_slug"),
                att.get("section"),
                str(att.get("item_key") or "")
            )
            grouped.setdefault(key, []).append(att)
        return grouped

    def _safe_report_filename(self, payload, extension):
        report = payload.get("report") or {}
        title = str(report.get("title") or f"LOGRA_{report.get('id') or 'report'}")
        safe = "".join(ch if ch.isalnum() or ch in (" ", "-", "_") else "_" for ch in title).strip()
        return f"{safe or 'LOGRA_report'}.{extension}"

    def _export_selected_word(self):
        payload = self._selected_payload()
        if not payload:
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=self._safe_report_filename(payload, "docx"),
            filetypes=[("Word", "*.docx")]
        )
        if not path:
            return
        try:
            self._build_word_report(payload, path)
            messagebox.showinfo("LOGRA", "Reporte Word generado correctamente.")
        except Exception as exc:
            messagebox.showerror("LOGRA", f"No se pudo generar Word:\n{exc}")

    def _export_selected_pdf(self):
        payload = self._selected_payload()
        if not payload:
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            initialfile=self._safe_report_filename(payload, "pdf"),
            filetypes=[("PDF", "*.pdf")]
        )
        if not path:
            return
        try:
            self._build_pdf_report(payload, path)
            messagebox.showinfo("LOGRA", "Reporte PDF generado correctamente.")
        except Exception as exc:
            messagebox.showerror("LOGRA", f"No se pudo generar PDF:\n{exc}")

    def _report_meta_rows(self, payload):
        report = payload.get("report") or {}
        agenda = report.get("agenda_items") or []
        return [
            ("Reporte", report.get("title") or ""),
            ("ID", report.get("id") or ""),
            ("Categoria", report.get("category") or "LOGRA"),
            ("Status", report.get("status") or ""),
            ("Agenda", f"{len(agenda)} reuniones"),
            ("Actualizado", str(report.get("updated_at") or "")),
            ("Generado", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ]

    def _build_word_report(self, payload, path):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        report = payload.get("report") or {}
        attachments = self._attachments_by_question(payload)
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        styles = doc.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(9)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(str(report.get("title") or "LOGRA Report"))
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 59, 113)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("Professional LOGRA Questionnaire Report").italic = True

        meta = doc.add_table(rows=0, cols=2)
        meta.style = "Table Grid"
        for label, value in self._report_meta_rows(payload):
            cells = meta.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].bold = True

        agenda = report.get("agenda_items") or []
        if agenda:
            doc.add_heading("Meeting Agenda", level=1)
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            headers = ["Date", "Start", "End", "Place", "Person", "Company/Role", "Topic"]
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
                table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
            for item in agenda:
                cells = table.add_row().cells
                values = [
                    item.get("date") or item.get("date_long") or item.get("date_iso") or "",
                    item.get("start_time") or "",
                    item.get("end_time") or "",
                    item.get("place") or "",
                    item.get("person") or "",
                    item.get("company_role") or "",
                    item.get("topic") or "",
                ]
                for idx, value in enumerate(values):
                    cells[idx].text = str(value)

        current_form = None
        current_section = None
        for form, section_key, question, answer in self._ordered_answers(payload):
            if current_form != form.get("slug"):
                doc.add_heading(form.get("title") or "", level=1)
                current_form = form.get("slug")
                current_section = None
            if current_section != section_key:
                doc.add_heading(LograReportsTable._section_label(section_key), level=2)
                current_section = section_key

            qid = str(question.get("id") or question.get("number") or answer.get("item_key") or "")
            q = doc.add_paragraph()
            q_run = q.add_run(f"{qid}. {answer.get('question_text') or question.get('question') or ''}")
            q_run.bold = True

            bullets = answer.get("bullets") or []
            if bullets:
                for bullet in bullets:
                    doc.add_paragraph(str(bullet), style="List Bullet")
            else:
                doc.add_paragraph("Sin respuesta registrada.", style="Intense Quote")

            key = (form.get("slug"), section_key, qid)
            attached = attachments.get(key) or []
            if attached:
                p = doc.add_paragraph()
                p.add_run("Adjuntos: ").bold = True
                p.add_run(", ".join(str(att.get("original_filename") or att.get("id")) for att in attached))

        doc.save(path)

    def _build_pdf_report(self, payload, path):
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        report = payload.get("report") or {}
        attachments = self._attachments_by_question(payload)
        doc = SimpleDocTemplate(
            path,
            pagesize=letter,
            rightMargin=0.55 * inch,
            leftMargin=0.55 * inch,
            topMargin=0.55 * inch,
            bottomMargin=0.55 * inch,
            title=str(report.get("title") or "LOGRA Report")
        )
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name="LograTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=16,
            textColor=colors.HexColor("#003B71"),
            alignment=1,
            spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            name="LograH1",
            parent=styles["Heading1"],
            fontSize=13,
            textColor=colors.HexColor("#003B71"),
            spaceBefore=12,
            spaceAfter=6,
        ))
        styles.add(ParagraphStyle(
            name="LograH2",
            parent=styles["Heading2"],
            fontSize=10.5,
            textColor=colors.HexColor("#4B6478"),
            spaceBefore=8,
            spaceAfter=4,
        ))
        styles.add(ParagraphStyle(
            name="Question",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8.5,
            leading=11,
            spaceBefore=6,
            spaceAfter=3,
        ))
        styles.add(ParagraphStyle(
            name="LograBullet",
            parent=styles["Normal"],
            leftIndent=12,
            firstLineIndent=-6,
            fontSize=8.2,
            leading=10.5,
            spaceAfter=2,
        ))

        story = [
            Paragraph(escape(str(report.get("title") or "LOGRA Report")), styles["LograTitle"]),
            Paragraph("Professional LOGRA Questionnaire Report", styles["Normal"]),
            Spacer(1, 8),
        ]
        meta_data = [[Paragraph(f"<b>{escape(str(label))}</b>", styles["Normal"]), Paragraph(escape(str(value)), styles["Normal"])]
                     for label, value in self._report_meta_rows(payload)]
        meta_table = Table(meta_data, colWidths=[1.35 * inch, 5.4 * inch])
        meta_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C1CA")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E9EEF3")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.extend([meta_table, Spacer(1, 10)])

        agenda = report.get("agenda_items") or []
        if agenda:
            story.append(Paragraph("Meeting Agenda", styles["LograH1"]))
            agenda_data = [["Date", "Start", "End", "Place", "Person", "Company/Role", "Topic"]]
            for item in agenda:
                agenda_data.append([
                    item.get("date") or item.get("date_long") or item.get("date_iso") or "",
                    item.get("start_time") or "",
                    item.get("end_time") or "",
                    item.get("place") or "",
                    item.get("person") or "",
                    item.get("company_role") or "",
                    item.get("topic") or "",
                ])
            table = Table([[Paragraph(escape(str(cell)), styles["Normal"]) for cell in row] for row in agenda_data],
                          repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003B71")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C1CA")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.extend([table, Spacer(1, 8)])

        current_form = None
        current_section = None
        for form, section_key, question, answer in self._ordered_answers(payload):
            if current_form != form.get("slug"):
                if current_form is not None:
                    story.append(PageBreak())
                story.append(Paragraph(escape(str(form.get("title") or "")), styles["LograH1"]))
                current_form = form.get("slug")
                current_section = None
            if current_section != section_key:
                story.append(Paragraph(escape(self._section_label(section_key)), styles["LograH2"]))
                current_section = section_key

            qid = str(question.get("id") or question.get("number") or answer.get("item_key") or "")
            question_text = answer.get("question_text") or question.get("question") or ""
            story.append(Paragraph(escape(f"{qid}. {question_text}"), styles["Question"]))
            bullets = answer.get("bullets") or []
            if bullets:
                for bullet in bullets:
                    story.append(Paragraph(f"&#8226; {escape(str(bullet))}", styles["LograBullet"]))
            else:
                story.append(Paragraph("Sin respuesta registrada.", styles["LograBullet"]))

            attached = attachments.get((form.get("slug"), section_key, qid)) or []
            if attached:
                names = ", ".join(str(att.get("original_filename") or att.get("id")) for att in attached)
                story.append(Paragraph(f"<b>Adjuntos:</b> {escape(names)}", styles["Normal"]))
            story.append(Spacer(1, 4))

        doc.build(story)

    @staticmethod
    def _section_label(section_key):
        return {
            "critical_questions": "Preguntas de apertura",
            "detailed_questions": "Preguntas por tema",
        }.get(section_key, section_key)

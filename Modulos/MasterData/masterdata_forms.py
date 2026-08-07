from __future__ import annotations

import os
import re
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any

from session_context import get_company_code, get_company_name


ASSETS_DIR = Path(__file__).resolve().parents[2] / "assets"
LOGO_PATH = ASSETS_DIR / "msl_logo.png"
WATERMARK_PATH = ASSETS_DIR / "watermark.png"


@dataclass(frozen=True)
class MasterDataFormSpec:
    key: str
    label: str
    endpoint: str
    code_field: str
    fields: tuple[str, ...]
    required_fields: tuple[str, ...]


FORM_SPECS: dict[str, MasterDataFormSpec] = {
    "cliente": MasterDataFormSpec(
        key="cliente",
        label="Cliente",
        endpoint="/clientes",
        code_field="Codigo",
        fields=(
            "Codigo", "NombreJuridico", "NombreComercial", "Pais", "Correo", "Telefono",
            "CedulaJuridicaVAT", "Comentarios", "Provincia", "Canton", "Distrito",
            "DireccionExacta", "FechaDePago", "Prefijo", "ContactoPrincipal", "ContactoSecundario",
        ),
        required_fields=("Codigo", "NombreJuridico", "Pais"),
    ),
    "proveedor": MasterDataFormSpec(
        key="proveedor",
        label="Proveedor",
        endpoint="/proveedores",
        code_field="Codigo",
        fields=(
            "Codigo", "Nombre", "Apellidos", "NombreComercial", "Cedula", "Pais",
            "Provincia", "Canton", "Distrito", "DireccionExacta", "Prefijo", "Telefono",
            "Correo", "TerminosPago", "Banco", "CuentaIBAN", "SwiftCode", "UID",
            "DireccionBanco", "TipoProveeduria", "Comentarios",
        ),
        required_fields=("Codigo", "Nombre", "Pais"),
    ),
    "empleado": MasterDataFormSpec(
        key="empleado",
        label="Empleado",
        endpoint="/empleados",
        code_field="codigo",
        fields=(
            "codigo", "nombre", "apellidos", "estado_civil", "genero", "nacionalidad",
            "prefijo", "telefono", "provincia", "canton", "distrito", "direccion",
            "jornada", "salario", "pago", "banco", "cuenta_iban", "moneda",
            "enfermedades", "contacto_emergencia", "telefono_emergencia",
            "activo1", "marca1", "serial1", "activo2", "marca2", "serial2",
            "activo3", "marca3", "serial3",
        ),
        required_fields=("nombre", "apellidos"),
    ),
    "surveyor": MasterDataFormSpec(
        key="surveyor",
        label="Surveyor",
        endpoint="/surveyores",
        code_field="codigo",
        fields=(
            "codigo", "nombre", "apellidos", "email", "estado_civil", "genero",
            "nacionalidad", "prefijo", "telefono", "provincia", "canton", "distrito",
            "direccion", "jornada", "operacion", "honorario", "pago", "banco",
            "direccion_banco", "cuenta_iban", "moneda", "swift", "uid",
            "enfermedades", "contacto_emergencia", "telefono_emergencia", "puerto",
        ),
        required_fields=("codigo", "nombre", "apellidos", "pais_o_nacionalidad"),
    ),
}


FORM_LABELS = {spec.label: key for key, spec in FORM_SPECS.items()}


def get_spec(entity_key_or_label: str) -> MasterDataFormSpec:
    value = str(entity_key_or_label or "").strip()
    key = FORM_LABELS.get(value, value.lower())
    if key not in FORM_SPECS:
        raise ValueError(f"Tipo de formulario no soportado: {entity_key_or_label}")
    return FORM_SPECS[key]


def export_masterdata_form(entity_key: str, fmt: str, output_path: str) -> str:
    spec = get_spec(entity_key)
    suffix = Path(output_path).suffix.lower()
    if fmt.lower() in {"word", "docx"} or suffix == ".docx":
        _export_docx(spec, output_path)
    elif fmt.lower() in {"excel", "xlsx"} or suffix == ".xlsx":
        _export_xlsx(spec, output_path)
    else:
        raise ValueError("Formato no soportado. Use Word (.docx) o Excel (.xlsx).")
    return output_path


def import_masterdata_files(paths: list[str]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for raw_path in paths:
        path = Path(raw_path)
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            records.extend(_read_xlsx(path))
        elif suffix == ".docx":
            records.extend(_read_docx(path))
        else:
            records.append({
                "file": str(path),
                "entity": "",
                "data": {},
                "error": "Formato no soportado",
            })
    return records


def clean_record(spec: MasterDataFormSpec, data: dict[str, Any]) -> dict[str, Any]:
    cleaned = {field: _clean_value(data.get(field)) for field in spec.fields}

    if spec.key == "cliente":
        cleaned["FechaDePago"] = _format_date(cleaned.get("FechaDePago"))
    elif spec.key in {"surveyor", "empleado"}:
        if cleaned.get("honorario"):
            cleaned["honorario"] = _clean_number(cleaned["honorario"])
        if cleaned.get("salario"):
            cleaned["salario"] = _clean_number(cleaned["salario"])

    return cleaned


def validate_record(spec: MasterDataFormSpec, data: dict[str, Any]) -> list[str]:
    missing: list[str] = []
    for field in spec.required_fields:
        if field == "pais_o_nacionalidad":
            if not (data.get("nacionalidad") or data.get("Pais") or data.get("pais")):
                missing.append("nacionalidad/pais")
            continue
        if not data.get(field):
            missing.append(field)
    return missing


def normalize_import_key(value: Any) -> str:
    text = str(value or "").strip()
    match = re.search(r"\(([^()]+)\)\s*$", text)
    if match:
        text = match.group(1)
    return text.strip()


def _company_header() -> tuple[str, str]:
    name = get_company_name()
    code = get_company_code()
    try:
        from api_client import get_current_company_profile_api

        profile = get_current_company_profile_api()
        name = profile.get("legal_name") or profile.get("company_name") or name
        code = profile.get("company_code") or code
    except Exception:
        pass
    return name, code


def _export_xlsx(spec: MasterDataFormSpec, output_path: str) -> None:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    company_name, company_code = _company_header()
    wb = Workbook()
    ws = wb.active
    ws.title = spec.key

    ws["A1"] = company_name
    ws["A2"] = f"Formulario Master Data - {spec.label}"
    ws["A3"] = f"Empresa: {company_code} | Llene una fila por registro. No cambie los encabezados."
    ws["A1"].font = Font(bold=True, size=14, color="003A75")
    ws["A2"].font = Font(bold=True, size=12)
    ws["A3"].font = Font(italic=True, color="666666")

    if LOGO_PATH.exists():
        try:
            logo = XLImage(str(LOGO_PATH))
            logo.width = 120
            logo.height = 70
            ws.add_image(logo, "H1")
        except Exception:
            pass

    if WATERMARK_PATH.exists():
        try:
            watermark = XLImage(str(WATERMARK_PATH))
            watermark.width = 330
            watermark.height = 210
            ws.add_image(watermark, "E8")
        except Exception:
            pass

    header_row = 5
    for idx, field in enumerate(spec.fields, start=1):
        cell = ws.cell(header_row, idx, field)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="003A75")
        cell.alignment = Alignment(horizontal="center")
        ws.column_dimensions[get_column_letter(idx)].width = max(15, min(28, len(field) + 4))

    for row in range(header_row + 1, header_row + 6):
        for idx in range(1, len(spec.fields) + 1):
            ws.cell(row, idx).alignment = Alignment(vertical="top", wrap_text=True)

    ws.freeze_panes = "A6"
    ws.sheet_view.showGridLines = True
    ws.oddHeader.center.text = company_name
    ws.oddFooter.center.text = f"{spec.label} | Master Data | &D"
    wb.save(output_path)


def _export_docx(spec: MasterDataFormSpec, output_path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    company_name, company_code = _company_header()
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        try:
            header_p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))
            header_p.add_run("\n")
        except Exception:
            pass
    run = header_p.add_run(company_name)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 58, 117)
    header_p.add_run(f"\n{company_code}")

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run(f"{spec.label} | Master Data | Generado {date.today().isoformat()}")

    watermark = doc.add_paragraph()
    watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wm_run = watermark.add_run(company_name)
    wm_run.font.size = Pt(24)
    wm_run.font.color.rgb = RGBColor(215, 215, 215)
    wm_run.bold = True

    title = doc.add_heading(f"Formulario Master Data - {spec.label}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Complete la columna Valor. No cambie los nombres de campo.")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"
    hdr[2].text = "Clave interna"

    for field in spec.fields:
        cells = table.add_row().cells
        cells[0].text = _pretty_label(field)
        cells[1].text = ""
        cells[2].text = field

    doc.save(output_path)


def _read_xlsx(path: Path) -> list[dict[str, Any]]:
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    records: list[dict[str, Any]] = []
    for ws in wb.worksheets:
        spec = _infer_spec(ws.title, [])
        header_row = None
        headers: list[str] = []
        for row_idx in range(1, min(ws.max_row, 20) + 1):
            row_values = [normalize_import_key(ws.cell(row_idx, col).value) for col in range(1, ws.max_column + 1)]
            row_keys = [v for v in row_values if v]
            inferred = _infer_spec(ws.title, row_keys)
            if inferred:
                spec = inferred
            if spec and sum(1 for h in row_values if h in spec.fields) >= 2:
                header_row = row_idx
                headers = row_values
                break
        if not spec or not header_row:
            records.append({"file": str(path), "entity": ws.title, "data": {}, "error": "No se pudo identificar el formulario"})
            continue
        for row_idx in range(header_row + 1, ws.max_row + 1):
            data = {}
            has_value = False
            for col_idx, field in enumerate(headers, start=1):
                if field not in spec.fields:
                    continue
                value = ws.cell(row_idx, col_idx).value
                if _clean_value(value) != "":
                    has_value = True
                data[field] = value
            if has_value:
                records.append({"file": str(path), "entity": spec.key, "data": clean_record(spec, data), "error": ""})
    return records


def _read_docx(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    title_text = "\n".join(p.text for p in doc.paragraphs[:8])
    table_fields: dict[str, Any] = {}
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) >= 3:
                field = normalize_import_key(cells[2].text)
                value = cells[1].text
            elif len(cells) >= 2:
                field = normalize_import_key(cells[0].text)
                value = cells[1].text
            else:
                continue
            if field:
                table_fields[field] = value

    spec = _infer_spec(title_text, list(table_fields))
    if not spec:
        return [{"file": str(path), "entity": "", "data": {}, "error": "No se pudo identificar el formulario"}]
    return [{"file": str(path), "entity": spec.key, "data": clean_record(spec, table_fields), "error": ""}]


def _infer_spec(text: str, fields: list[str]) -> MasterDataFormSpec | None:
    haystack = f"{text} {' '.join(fields)}".lower()
    for key, spec in FORM_SPECS.items():
        if key in haystack or spec.label.lower() in haystack:
            return spec
    for spec in FORM_SPECS.values():
        if sum(1 for field in fields if field in spec.fields) >= 3:
            return spec
    return None


def _clean_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    return str(value).strip()


def _clean_number(value: Any) -> str:
    text = _clean_value(value).replace(",", "")
    return text


def _format_date(value: Any) -> str:
    text = _clean_value(value)
    if not text:
        return ""
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(text, fmt).date().isoformat()
        except Exception:
            continue
    return text


def _pretty_label(field: str) -> str:
    label = re.sub(r"(?<!^)(?=[A-Z])", " ", field)
    label = label.replace("_", " ")
    return label[:1].upper() + label[1:]

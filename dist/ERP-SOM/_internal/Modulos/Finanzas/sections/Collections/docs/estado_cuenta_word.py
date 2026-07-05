from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from tkinter import filedialog, messagebox
import os
import sys
import tempfile
import shutil
from Modulos.Finanzas.date_utils import to_long_english_date


# ============================================================
# ASSETS PATH — BLINDADO (DEV / EXE PyInstaller)
# ============================================================

def get_assets_path() -> str:
    """
    Devuelve la ruta correcta a /assets en:
    - Desarrollo (repo)
    - EXE PyInstaller (sys._MEIPASS)
    """

    if getattr(sys, "frozen", False):
        base_path = getattr(sys, "_MEIPASS", os.path.dirname(sys.executable))
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))

    assets_path = os.path.join(base_path, "assets")

    if not os.path.isdir(assets_path):
        raise FileNotFoundError(
            f"Assets folder not found: {assets_path}"
        )

    return assets_path


# ============================================================
# CONSTANTES
# ============================================================

FOOTER_TEXT = (
    "Head Office – Costa Rica, Alajuela, Plaza Aeropuerto G-14\n"
    "Phone (506) 8814-07-84 – (506) 4052-8382"
)


# ============================================================
# GENERAR ESTADO DE CUENTA WORD (BLINDADO)
# ============================================================

def generar_estado_cuenta_word(
    idioma,
    cliente,
    resumen_kpis,
    facturas,
    datos_bancarios
):

    # ========================================================
    # SAVE AS — USUARIO DECIDE
    # ========================================================
    filename = f"Estado_de_Cuenta_{cliente}_{date.today().isoformat()}.docx"

    path = filedialog.asksaveasfilename(
        initialfile=filename,
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )

    if not path:
        return

    try:
        ASSETS_PATH = get_assets_path()
        HEADER_PATH = os.path.join(ASSETS_PATH, "header.png")
        WATERMARK_PATH = os.path.join(ASSETS_PATH, "watermark.png")

        if not os.path.isfile(HEADER_PATH):
            raise FileNotFoundError(f"Header not found: {HEADER_PATH}")

        if not os.path.isfile(WATERMARK_PATH):
            raise FileNotFoundError(f"Watermark not found: {WATERMARK_PATH}")

        # ====================================================
        # DOCUMENTO
        # ====================================================
        doc = Document()
        section = doc.sections[0]

        # ====================================================
        # HEADER + WATERMARK
        # ====================================================
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.add_run().add_picture(HEADER_PATH, width=Inches(2.2))

        wm_p = header.add_paragraph()
        wm_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wm_p.add_run().add_picture(WATERMARK_PATH, width=Inches(4.5))

        hoy = to_long_english_date(date.today())

        # ====================================================
        # RECÁLCULO DE BUCKETS
        # ====================================================
        total_ar = overdue = 0.0
        current = b1 = b2 = b3 = b4 = 0.0

        for f in facturas:

            aging = int(f.get("aging_dias") or 0)
            tipo_doc = f.get("tipo_documento")
            estado = f.get("estado_factura")

            if tipo_doc == "NOTA_CREDITO" and estado == "APLICADA":
                continue

            monto = float(
                f.get("saldo_pendiente")
                if tipo_doc == "FACTURA"
                else f.get("total") or 0
            )

            total_ar += monto

            if aging <= 0:
                current += monto
            elif aging <= 30:
                b1 += monto; overdue += monto
            elif aging <= 60:
                b2 += monto; overdue += monto
            elif aging <= 90:
                b3 += monto; overdue += monto
            else:
                b4 += monto; overdue += monto

        # ====================================================
        # TEXTO PRINCIPAL
        # ====================================================
        if idioma == "ES":
            body = (
                f"Estimado/a {cliente},\n\n"
                f"Asunto: Estado de cuenta MSL SRL – Cliente {cliente} – {hoy}\n\n"
                f"Por medio del presente le compartimos el estado de cuenta al día de hoy, "
                f"el cual refleja un balance adeudado total de {total_ar:,.2f}, "
                f"de los cuales {overdue:,.2f} se encuentran vencidos.\n\n"
                f"El balance se compone de la siguiente manera:\n"
            )
        else:
            body = (
                f"Dear {cliente},\n\n"
                f"Subject: Statement of Account MSL SRL – Client {cliente} – {hoy}\n\n"
                f"Please find below the statement of account as of today, reflecting a total "
                f"outstanding balance of {total_ar:,.2f}, of which {overdue:,.2f} are overdue.\n\n"
                f"Breakdown of outstanding balance:\n"
            )

        doc.add_paragraph(body)

        # ====================================================
        # BULLETS
        # ====================================================
        bullets = [
            ("Current", current),
            ("1–30 days", b1),
            ("31–60 days", b2),
            ("61–90 days", b3),
            ("90+ days", b4),
        ]

        for label, amount in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{label}: {amount:,.2f}")

        # ====================================================
        # DATOS BANCARIOS
        # ====================================================
        doc.add_paragraph("\nBank details:" if idioma == "EN" else "\nDatos bancarios:")

        for k, v in datos_bancarios.items():
            if k != "idioma" and v:
                doc.add_paragraph(f"{k}: {v}")

        # ====================================================
        # TABLA FACTURAS
        # ====================================================
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        table.autofit = False

        headers = [
            "Número Documento", "Fecha Emisión", "Fecha Vencimiento",
            "Aging", "Total", "Buque", "Operación", "Num. Informe"
        ]

        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        column_widths = [
            Inches(1.65), Inches(0.9), Inches(0.9), Inches(0.55),
            Inches(0.85), Inches(1.0), Inches(1.0), Inches(0.9)
        ]
        for row in table.rows:
            for idx, width in enumerate(column_widths):
                row.cells[idx].width = width

        for f in facturas:
            r = table.add_row().cells
            r[0].text = str(f.get("numero_documento", ""))
            r[1].text = to_long_english_date(f.get("fecha_emision", ""))
            r[2].text = to_long_english_date(f.get("fecha_vencimiento", ""))
            r[3].text = str(f.get("aging_dias", ""))
            r[4].text = f"{float(f.get('total') or 0):,.2f}"
            r[5].text = str(f.get("buque_contenedor", ""))
            r[6].text = str(f.get("operacion", ""))
            r[7].text = str(f.get("num_informe", ""))
            for idx, width in enumerate(column_widths):
                r[idx].width = width

        # ====================================================
        # FOOTER
        # ====================================================
        fp = section.footer.paragraphs[0]
        fp.text = FOOTER_TEXT
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ====================================================
        # GUARDADO ATÓMICO
        # ====================================================
        temp_dir = os.path.dirname(path)
        with tempfile.NamedTemporaryFile(
            delete=False,
            dir=temp_dir,
            suffix=".docx"
        ) as tmp:
            temp_path = tmp.name

        doc.save(temp_path)
        shutil.move(temp_path, path)

        messagebox.showinfo(
            "Estado de cuenta",
            "Estado de cuenta generado correctamente."
        )

    except PermissionError:
        messagebox.showerror(
            "Error",
            "No se pudo guardar el archivo.\n\n"
            "Cierre Word si el archivo está abierto\n"
            "o seleccione otra ubicación."
        )

    except Exception as e:
        messagebox.showerror(
            "Error",
            f"No se pudo generar el estado de cuenta.\n\n{e}"
        )

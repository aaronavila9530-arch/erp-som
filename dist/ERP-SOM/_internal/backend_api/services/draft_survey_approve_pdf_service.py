import os
import sys
import shutil
import tempfile
import subprocess
import time
from pathlib import Path

from docx import Document
try:
    from services.template_autofit import apply_docx_autofit
except ModuleNotFoundError:
    from backend_api.services.template_autofit import apply_docx_autofit

print("==== DraftSurveyApprovePdfService LOADED ====")
print("FILE:", __file__)
print("os module:", os)
print("sys module:", sys)

try:
    from docx2pdf import convert as docx2pdf_convert
    print("docx2pdf loaded OK")
except Exception as e:
    print("docx2pdf NOT available:", e)
    docx2pdf_convert = None

from services.draft_survey_excel_service import DraftSurveyExcelGenerator
from services.pdf_merge_service import merge_pdf_list


# =========================================================
# TEMPLATE PATH
# =========================================================
TEMPLATE_WORD_PATH = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        "..",
        "templates",
        "draft_word_template.docx"
    )
)

print("TEMPLATE_WORD_PATH:", TEMPLATE_WORD_PATH)


# =========================================================
# SERVICE
# =========================================================
class DraftSurveyApprovePdfService:

    # =====================================================
    # PUBLIC
    # =====================================================
    def generate_final_pdf(self, word_payload: dict, excel_payload: dict) -> str:

        print("\n[1] generate_final_pdf START")

        print("[2] Generating Word PDF...")
        word_pdf = self._generate_word_pdf_from_template(word_payload)
        print("[3] Word PDF OK:", word_pdf)

        print("[4] Generating Excel PDF...")
        excel_pdf = self._generate_excel_pdf_from_template(excel_payload)
        print("[5] Excel PDF OK:", excel_pdf)

        print("[6] Before merge - os exists?", os)
        print("[7] Calling merge_pdf_list...")

        final_pdf = merge_pdf_list([word_pdf, excel_pdf])

        print("[8] Merge returned:", final_pdf)

        if not final_pdf:
            raise RuntimeError("Final merged PDF is empty.")

        print("[9] Checking existence of final PDF...")
        print("os:", os)
        print("Exists?:", os.path.exists(final_pdf))

        if not os.path.exists(final_pdf):
            raise RuntimeError("Final merged PDF was not created.")

        print("[10] FINAL PDF SUCCESS:", final_pdf)

        return final_pdf

    # =====================================================
    # WORD PIPELINE
    # =====================================================
    def _generate_word_pdf_from_template(self, payload: dict) -> str:

        print("  -> [W1] Checking Word template...")

        if not os.path.exists(TEMPLATE_WORD_PATH):
            raise FileNotFoundError(
                f"Word template not found: {TEMPLATE_WORD_PATH}"
            )

        out_dir = tempfile.mkdtemp(prefix="draft_word_")
        tmp_docx = os.path.join(out_dir, "draft_word_filled.docx")
        tmp_pdf = os.path.join(out_dir, "draft_word_filled.pdf")

        print("  -> [W2] Loading Word template...")
        doc = Document(TEMPLATE_WORD_PATH)

        replacements = self._build_replacements(payload)

        print("  -> [W3] Replacing placeholders...")
        self._replace_in_document(doc, replacements)

        for section in doc.sections:
            self._replace_in_container(section.header, replacements)
            self._replace_in_container(section.footer, replacements)

        print("  -> [W4] Saving DOCX...")
        apply_docx_autofit(doc)
        doc.save(tmp_docx)

        print("  -> [W5] Converting DOCX to PDF...")
        self._convert_docx_to_pdf(tmp_docx, tmp_pdf)

        print("  -> [W6] Checking Word PDF exists:", tmp_pdf)
        print("  -> Exists?:", os.path.exists(tmp_pdf))

        if not os.path.exists(tmp_pdf):
            raise RuntimeError("Word PDF was not created.")

        return tmp_pdf

    def _build_replacements(self, payload: dict) -> dict:
        repl = {}
        for k, v in (payload or {}).items():
            repl["{" + str(k) + "}"] = "" if v is None else str(v)
        return repl

    def _replace_in_document(self, doc: Document, repl: dict):
        for p in doc.paragraphs:
            self._replace_in_paragraph_safe(p, repl)

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph_safe(p, repl)

    def _replace_in_container(self, container, repl: dict):
        for p in container.paragraphs:
            self._replace_in_paragraph_safe(p, repl)

        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph_safe(p, repl)

    def _replace_in_paragraph_safe(self, paragraph, repl: dict):
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        if "{" not in full_text:
            return

        for ph, val in repl.items():
            if ph in full_text:
                full_text = full_text.replace(ph, val)

        idx = 0
        for run in paragraph.runs:
            length = len(run.text)
            run.text = full_text[idx: idx + length]
            idx += length

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str):

        print("  -> [W7] Starting PDF conversion...")

        out_dir = os.path.dirname(pdf_path)
        os.makedirs(out_dir, exist_ok=True)

        if docx2pdf_convert:
            print("  -> Using docx2pdf")
            docx2pdf_convert(docx_path, out_dir)
            return

        print("  -> Using LibreOffice fallback")
        cmd = [
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", out_dir,
            docx_path
        ]

        r = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

        print("  -> LibreOffice return code:", r.returncode)

        if r.returncode != 0:
            raise RuntimeError(
                f"DOCX->PDF conversion failed:\n{r.stderr}"
            )

    # =====================================================
    # EXCEL -> PDF
    # =====================================================
    def _generate_excel_pdf_from_template(self, payload: dict) -> str:

        print("  -> [E1] Importing Excel COM...")

        import pythoncom
        import win32com.client

        print("  -> [E2] Generating Excel file...")
        generator = DraftSurveyExcelGenerator()
        excel_path = generator.generate(payload or {})
        print("  -> Excel generated:", excel_path)

        out_dir = tempfile.mkdtemp(prefix="draft_excel_pdf_")
        pdf_path = os.path.abspath(os.path.join(out_dir, "draft_survey.pdf"))

        print("  -> [E3] Exporting Excel to PDF...")

        pythoncom.CoInitialize()
        excel = None
        workbook = None

        try:
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False

            workbook = excel.Workbooks.Open(excel_path)
            workbook.ExportAsFixedFormat(Type=0, Filename=pdf_path)

        finally:
            if workbook:
                workbook.Close(False)
            if excel:
                excel.Quit()
            pythoncom.CoUninitialize()

        print("  -> [E4] Excel PDF path:", pdf_path)
        print("  -> Exists?:", os.path.exists(pdf_path))

        if not os.path.exists(pdf_path):
            raise RuntimeError("Excel PDF was not created.")

        return pdf_path

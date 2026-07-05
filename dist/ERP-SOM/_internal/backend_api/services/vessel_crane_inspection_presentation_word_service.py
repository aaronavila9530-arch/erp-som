import os
import tempfile
import subprocess
from datetime import datetime, date
from docx import Document
try:
    from services.template_autofit import apply_docx_autofit
except ModuleNotFoundError:
    from backend_api.services.template_autofit import apply_docx_autofit


class VesselCraneInspectionPresentationWordService:

    TEMPLATE_NAME = "presentation_crane_inspection.docx"

    # =========================================================
    # MAIN
    # =========================================================
    def generate_pdf_by_id(self, conn, record_id: int) -> str:

        data = self._get_data(conn, record_id)

        if not data:
            raise ValueError("Record not found")

        template_path = self._get_template_path()

        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f"Template not found: {template_path}"
            )

        doc = Document(template_path)

        # -----------------------------------------------------
        # FORMAT DATES
        # -----------------------------------------------------

        data = self._format_dates(data)

        # -----------------------------------------------------
        # REPLACE PLACEHOLDERS (ROBUST)
        # -----------------------------------------------------

        self._replace_everywhere(doc, data)

        # -----------------------------------------------------
        # TEMP FILES
        # -----------------------------------------------------

        temp_dir = tempfile.mkdtemp()

        docx_path = os.path.join(
            temp_dir,
            f"crane_inspection_presentation_{record_id}.docx"
        )

        pdf_path = os.path.join(
            temp_dir,
            f"crane_inspection_presentation_{record_id}.pdf"
        )

        apply_docx_autofit(doc)
        doc.save(docx_path)

        # -----------------------------------------------------
        # CONVERT DOCX → PDF
        # -----------------------------------------------------

        self._convert_to_pdf(docx_path, temp_dir)

        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF conversion failed")

        return pdf_path


    # =========================================================
    # CONVERT DOCX → PDF
    # =========================================================

    def _convert_to_pdf(self, docx_path: str, output_dir: str):

        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path
            ],
            check=True
        )


    # =========================================================
    # TEMPLATE PATH
    # =========================================================

    def _get_template_path(self):

        base_dir = os.path.dirname(os.path.abspath(__file__))

        return os.path.abspath(
            os.path.join(
                base_dir,
                "..",
                "templates",
                self.TEMPLATE_NAME
            )
        )


    # =========================================================
    # GET DATA
    # =========================================================

    def _get_data(self, conn, record_id):

        cur = conn.cursor()

        cur.execute(
            """
            SELECT *
            FROM vessel_crane_inspection_reports
            WHERE id = %s
            """,
            (record_id,)
        )

        row = cur.fetchone()

        if not row:
            return None

        columns = [desc[0] for desc in cur.description]

        return dict(zip(columns, row))


    # =========================================================
    # FORMAT DATES
    # =========================================================

    def _format_dates(self, data):

        formatted = {}

        for k, v in data.items():

            if isinstance(v, (datetime, date)):
                formatted[k] = v.strftime("%B %d %Y")
            else:
                formatted[k] = v

        return formatted


    # =========================================================
    # REPLACE EVERYWHERE
    # =========================================================

    def _replace_everywhere(self, doc, data):

        # paragraphs
        for p in doc.paragraphs:
            self._replace_paragraph(p, data)

        # tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_paragraph(p, data)

        # headers / footers
        for section in doc.sections:

            for p in section.header.paragraphs:
                self._replace_paragraph(p, data)

            for p in section.footer.paragraphs:
                self._replace_paragraph(p, data)


    # =========================================================
    # PARAGRAPH SAFE REPLACEMENT
    # =========================================================

    def _replace_paragraph(self, paragraph, data):

        full_text = "".join(run.text for run in paragraph.runs)

        replaced = full_text

        for key, value in data.items():

            placeholder = "{" + key + "}"

            replaced = replaced.replace(
                placeholder,
                "" if value is None else str(value)
            )

        if replaced != full_text:

            paragraph.runs[0].text = replaced

            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ""

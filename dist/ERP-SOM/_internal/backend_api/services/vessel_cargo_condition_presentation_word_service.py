import os
import tempfile
import subprocess
from datetime import datetime
from docx import Document
try:
    from services.template_autofit import apply_docx_autofit
except ModuleNotFoundError:
    from backend_api.services.template_autofit import apply_docx_autofit


class VesselCargoConditionPresentationWordService:

    TEMPLATE_NAME = "presentation_cargo_condition.docx"

    # =========================================================
    # MAIN
    # =========================================================
    def generate_pdf_by_id(self, conn, record_id: int) -> str:

        data = self._get_data(conn, record_id)

        if not data:
            raise ValueError("Record not found")

        template_path = self._get_template_path()

        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f"Template not found: {template_path}"
            )

        doc = Document(template_path)

        # Replace placeholders (split-run safe + date normalization)
        self._replace_placeholders(doc, data)

        temp_dir = tempfile.mkdtemp()

        docx_path = os.path.join(
            temp_dir,
            f"cargo_condition_presentation_{record_id}.docx"
        )

        pdf_path = os.path.join(
            temp_dir,
            f"cargo_condition_presentation_{record_id}.pdf"
        )

        apply_docx_autofit(doc)
        doc.save(docx_path)

        self._convert_to_pdf(docx_path, temp_dir)

        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF conversion failed")

        return pdf_path

    # =========================================================
    # CONVERT DOCX → PDF
    # =========================================================
    def _convert_to_pdf(self, docx_path: str, output_dir: str):

        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    output_dir,
                    docx_path
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
        except Exception as e:
            raise RuntimeError(f"LibreOffice conversion failed: {e}")

    # =========================================================
    # GET DATA FROM DB
    # =========================================================
    def _get_data(self, conn, record_id: int):

        cur = conn.cursor()

        cur.execute("""
            SELECT *
            FROM vessel_cargo_condition_surveys
            WHERE id = %s
        """, (record_id,))

        row = cur.fetchone()

        if not row:
            return None

        columns = [desc[0] for desc in cur.description]

        return dict(zip(columns, row))

    # =========================================================
    # TEMPLATE PATH
    # =========================================================
    def _get_template_path(self):

        base_dir = os.path.dirname(os.path.abspath(__file__))

        return os.path.abspath(
            os.path.join(
                base_dir,
                "..",
                "templates",
                self.TEMPLATE_NAME
            )
        )

    # =========================================================
    # DATE NORMALIZER (ISO → LONG ENGLISH)
    # =========================================================
    def _format_date_long_english(self, value):

        if not value:
            return ""

        value = str(value).strip()

        try:
            # Detect ISO yyyy-mm-dd
            dt = datetime.strptime(value, "%Y-%m-%d")
            return dt.strftime("%B %d %Y")
        except Exception:
            return value

    # =========================================================
    # SAFE VALUE FORMATTER
    # =========================================================
    def _safe_value(self, key, value):

        if value is None:
            return ""

        # Auto-detect date fields
        if key.endswith("_date") or key in ["service_start_date"]:
            return self._format_date_long_english(value)

        return str(value)

    # =========================================================
    # PLACEHOLDER REPLACER (RUN-SAFE + DATE SAFE)
    # =========================================================
    def _replace_placeholders(self, doc, data: dict):

        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data)

        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, data)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, data)

    def _replace_in_paragraph(self, paragraph, data):

        if not paragraph.text:
            return

        full_text = paragraph.text

        for key, value in data.items():
            placeholder = f"{{{key}}}"

            if placeholder in full_text:
                formatted_value = self._safe_value(key, value)
                full_text = full_text.replace(
                    placeholder,
                    formatted_value
                )

        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            paragraph.runs[0].text = full_text

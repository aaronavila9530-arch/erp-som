from copy import copy
import math

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.cell_range import CellRange


MIN_WORD_MARGIN_IN = 0.35
DEFAULT_CELL_MARGIN_TWIPS = 80
EXCEL_MIN_FONT_SIZE = 7
EXCEL_BASE_LINE_HEIGHT = 15


def _coerce_text(value):
    return "" if value is None else str(value)


def _safe_font_size(run):
    try:
        return run.font.size.pt if run.font.size else None
    except Exception:
        return None


def _set_cell_margins(cell, margin_twips=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def _allow_cell_wrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for no_wrap in list(tc_pr.findall(qn("w:noWrap"))):
        tc_pr.remove(no_wrap)


def _set_table_autofit(table):
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    for layout in list(tbl_pr.findall(qn("w:tblLayout"))):
        tbl_pr.remove(layout)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")


def _fit_paragraph_runs(paragraph, in_table_cell=False):
    text = (paragraph.text or "").strip()
    if not text:
        return

    paragraph.paragraph_format.widow_control = True

    if not in_table_cell:
        paragraph.paragraph_format.keep_together = len(text) <= 600
        paragraph.paragraph_format.keep_with_next = False
        return

    length = len(text)
    if length <= 45:
        return

    if length > 420:
        target_size = 6
    elif length > 260:
        target_size = 7
    elif length > 170:
        target_size = 8
    elif length > 100:
        target_size = 9
    else:
        target_size = 10

    for run in paragraph.runs:
        if not run.text:
            continue

        current_size = _safe_font_size(run)
        if current_size is None or current_size > target_size:
            run.font.size = Pt(target_size)

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def _normalize_section_margins(section):
    min_margin = Inches(MIN_WORD_MARGIN_IN)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        value = getattr(section, attr, None)
        if value is None or value < min_margin:
            setattr(section, attr, min_margin)


def _walk_docx_tables(tables):
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _walk_docx_tables(cell.tables)


def _fit_table(table):
    _set_table_autofit(table)
    for row in table.rows:
        for cell in row.cells:
            _set_cell_margins(cell, DEFAULT_CELL_MARGIN_TWIPS)
            _allow_cell_wrap(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                _fit_paragraph_runs(paragraph, in_table_cell=True)


def apply_docx_autofit(doc):
    """
    Apply layout-safe text fitting to generated Word files.

    Word cannot truly auto-fit text in every shape/textbox after python-docx
    replacement. This keeps existing template styling, allows table cells to
    wrap, adds small internal margins when missing, and only shrinks long text
    inside table cells so values do not get clipped.
    """
    for section in doc.sections:
        _normalize_section_margins(section)

    for paragraph in doc.paragraphs:
        _fit_paragraph_runs(paragraph, in_table_cell=False)

    for table in _walk_docx_tables(doc.tables):
        _fit_table(table)

    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _fit_paragraph_runs(paragraph, in_table_cell=False)
            for table in _walk_docx_tables(part.tables):
                _fit_table(table)


def _merged_range_for_cell(ws, row, col):
    coord = f"{get_column_letter(col)}{row}"
    for merged in ws.merged_cells.ranges:
        if coord in merged:
            return merged
    return None


def _range_width(ws, cell_range):
    if isinstance(cell_range, CellRange):
        columns = range(cell_range.min_col, cell_range.max_col + 1)
    else:
        columns = [cell_range.column]

    width = 0
    for col in columns:
        letter = get_column_letter(col)
        if ws.column_dimensions[letter].hidden:
            continue
        width += ws.column_dimensions[letter].width or 8.43
    return max(width, 8.43)


def _estimated_lines(value, width):
    text = _coerce_text(value)
    if not text:
        return 1

    chars_per_line = max(int(width * 1.15), 8)
    total = 0
    for line in text.splitlines() or [""]:
        if not line:
            total += 1
            continue
        words = line.split(" ")
        visual_len = 0
        for word in words:
            visual_len += max(len(word), chars_per_line) if len(word) > chars_per_line else len(word) + 1
        total += max(1, math.ceil(visual_len / chars_per_line))
    return max(total, 1)


def _is_probably_identifier(value):
    text = _coerce_text(value).strip()
    if len(text) < 14:
        return False
    return text.isdigit() or ("-" in text and not any(ch.isspace() for ch in text))


def _target_excel_font_size(value, width):
    text = _coerce_text(value).strip()
    if not text:
        return None

    lines = _estimated_lines(text, width)
    length = len(text)

    if _is_probably_identifier(text):
        return 8 if length > 22 else 9
    if lines >= 8 or length > 700:
        return 7
    if lines >= 5 or length > 420:
        return 8
    if lines >= 3 or length > 220:
        return 9
    return None


def _set_excel_page_defaults(ws):
    ws.page_margins.left = min(ws.page_margins.left or 0.25, 0.25)
    ws.page_margins.right = min(ws.page_margins.right or 0.25, 0.25)
    ws.page_margins.top = min(ws.page_margins.top or 0.35, 0.35)
    ws.page_margins.bottom = min(ws.page_margins.bottom or 0.35, 0.35)
    ws.page_margins.header = min(ws.page_margins.header or 0.15, 0.15)
    ws.page_margins.footer = min(ws.page_margins.footer or 0.15, 0.15)
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.print_options.horizontalCentered = True


def apply_excel_autofit(ws, min_row_height=15, max_row_height=135):
    """
    Apply pragmatic fitting to an Excel worksheet.

    Excel calculates perfect AutoFit only in the desktop app. For generated
    files we preserve template column widths, enable wrapping/shrink-to-fit,
    expand row heights heuristically, reduce long-text font size, and set
    print scaling so content does not spill outside page margins.
    """
    _set_excel_page_defaults(ws)
    row_heights = {}

    for row in ws.iter_rows():
        for cell in row:
            if cell.value in (None, ""):
                continue

            alignment = copy(cell.alignment)
            alignment.wrap_text = True
            alignment.shrink_to_fit = True
            alignment.vertical = alignment.vertical or "center"
            cell.alignment = alignment

            merged = _merged_range_for_cell(ws, cell.row, cell.column)
            effective_range = merged or cell
            width = _range_width(ws, effective_range)
            lines = _estimated_lines(cell.value, width)
            wanted_height = min(max_row_height, max(min_row_height, lines * EXCEL_BASE_LINE_HEIGHT))
            row_heights[cell.row] = max(row_heights.get(cell.row, min_row_height), wanted_height)

            target_size = _target_excel_font_size(cell.value, width)
            if target_size:
                font = copy(cell.font)
                current_size = font.sz or 11
                if current_size > target_size:
                    font.sz = max(EXCEL_MIN_FONT_SIZE, target_size)
                    cell.font = font

            if _is_probably_identifier(cell.value):
                cell.number_format = "@"

    for row_idx, height in row_heights.items():
        current = ws.row_dimensions[row_idx].height or min_row_height
        ws.row_dimensions[row_idx].height = max(current, height)


def apply_workbook_autofit(wb):
    for ws in wb.worksheets:
        apply_excel_autofit(ws)

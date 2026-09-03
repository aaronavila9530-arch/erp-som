import calendar
import json
import os
import tempfile
from datetime import date, timedelta
from decimal import Decimal
from html import escape

from psycopg2.extras import RealDictCursor


BLUE = "#123A63"
DARK = "#1F2933"
TEAL = "#2D9CDB"
GREEN = "#27AE60"
ORANGE = "#F2994A"
RED = "#C0392B"
GREY = "#F4F6F8"
LIGHT_BLUE = "#E8F0F8"
CHART_COLORS = [BLUE, TEAL, GREEN, ORANGE, RED, "#7F8C8D", "#8E44AD"]
MONTH_NAMES_ES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}


def _period_bounds(year: int, month: int):
    start = date(year, month, 1)
    last_day = calendar.monthrange(year, month)[1]
    end = date(year, month, last_day)
    prev_start = date(year - 1, 12, 1) if month == 1 else date(year, month - 1, 1)
    prev_end = start - timedelta(days=1)
    next_month = 1 if month == 12 else month + 1
    next_year = year + 1 if month == 12 else year
    next_start = date(next_year, next_month, 1)
    next_end = date(next_year, next_month, calendar.monthrange(next_year, next_month)[1])
    return start, end, prev_start, prev_end, next_start, next_end


def _f(value):
    if isinstance(value, Decimal):
        return float(value)
    try:
        return float(value or 0)
    except Exception:
        return 0.0


def _money(value):
    return f"USD {_f(value):,.2f}"


def _short_money(value):
    amount = _f(value)
    if abs(amount) >= 1000:
        return f"USD {amount / 1000:,.1f}K"
    return f"USD {amount:,.0f}"


def _pct(current, previous):
    previous = _f(previous)
    current = _f(current)
    if previous == 0:
        return None
    return ((current - previous) / previous) * 100.0


def _month_label(year: int, month: int):
    q = ((month - 1) // 3) + 1
    month_in_q = ((month - 1) % 3) + 1
    return f"Reporte financiero, Mes {month_in_q} Q{q} FY{str(year)[-2:]}"


def _fetch_one(cur, sql, params=()):
    cur.execute(sql, params)
    return cur.fetchone() or {}


def _fetch_all(cur, sql, params=()):
    cur.execute(sql, params)
    return cur.fetchall() or []


def _ensure_monthly_report_obligations(conn):
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS monthly_report_obligations (
            id BIGSERIAL PRIMARY KEY,
            period VARCHAR(7) NOT NULL,
            payee_name TEXT NOT NULL,
            concept TEXT,
            amount NUMERIC(18,2) NOT NULL DEFAULT 0,
            currency VARCHAR(3) NOT NULL DEFAULT 'USD',
            issue_date DATE,
            due_date DATE,
            source TEXT NOT NULL DEFAULT 'PREVIEW',
            accepted BOOLEAN NOT NULL DEFAULT TRUE,
            created_by TEXT,
            created_at TIMESTAMP NOT NULL DEFAULT now(),
            updated_at TIMESTAMP NOT NULL DEFAULT now()
        )
        """
    )
    cur.execute("ALTER TABLE monthly_report_obligations ADD COLUMN IF NOT EXISTS issue_date DATE")
    conn.commit()


def build_monthly_obligation_preview(conn, year: int, month: int):
    _ensure_monthly_report_obligations(conn)
    start, end, *_ = _period_bounds(year, month)
    period = f"{year}-{month:02d}"
    cur = conn.cursor(cursor_factory=RealDictCursor)

    saved = _fetch_all(cur, """
        SELECT
            id,
            payee_name,
            concept,
            amount,
            currency,
            COALESCE(issue_date, due_date) AS issue_date,
            due_date,
            source,
            accepted
        FROM monthly_report_obligations
        WHERE period = %s
        ORDER BY COALESCE(issue_date, due_date) NULLS LAST, payee_name
    """, (period,))
    if saved:
        return {"period": period, "source": "saved", "data": [dict(row) for row in saved]}

    rows = _fetch_all(cur, """
        SELECT
            payee_name,
            COALESCE(NULLIF(obligation_type, ''), NULLIF(payee_type, ''), 'GASTO MENSUAL') AS concept,
            currency,
            ROUND(AVG(total), 2) AS amount,
            MIN(EXTRACT(DAY FROM COALESCE(issue_date, due_date)))::int AS issue_day,
            COUNT(*) AS samples
        FROM payment_obligations
        WHERE COALESCE(issue_date, due_date) >= (%s::date - INTERVAL '12 months')
          AND COALESCE(issue_date, due_date) < %s::date
          AND COALESCE(total, 0) > 0
        GROUP BY payee_name, COALESCE(NULLIF(obligation_type, ''), NULLIF(payee_type, ''), 'GASTO MENSUAL'), currency
        HAVING COUNT(*) >= 1
        ORDER BY amount DESC, payee_name
        LIMIT 40
    """, (start, start))
    preview = []
    for row in rows:
        issue_day = min(max(int(row.get("issue_day") or 1), 1), calendar.monthrange(year, month)[1])
        preview.append({
            "payee_name": row.get("payee_name"),
            "concept": row.get("concept"),
            "amount": _f(row.get("amount")),
            "currency": row.get("currency") or "USD",
            "issue_date": date(year, month, issue_day).isoformat(),
            "due_date": None,
            "source": f"HISTORICO_{int(row.get('samples') or 0)}",
            "accepted": True,
        })

    salary_rows = _fetch_all(cur, """
        WITH payroll AS (
            SELECT
                pr.usuario,
                TRIM(CONCAT(COALESCE(e.nombre, ''), ' ', COALESCE(e.apellidos, ''))) AS employee_name,
                COALESCE(e.moneda, 'CRC') AS currency,
                COALESCE(pr.salario_bruto, pr.salario_neto, 0) AS amount,
                'PAYROLL_RUN' AS source
            FROM payroll_runs pr
            LEFT JOIN empleados e ON e.usuario = pr.usuario
            WHERE pr.year = %s
              AND pr.month = %s
              AND COALESCE(pr.salario_bruto, pr.salario_neto, 0) > 0
        ),
        employee_budget AS (
            SELECT
                e.usuario,
                TRIM(CONCAT(COALESCE(e.nombre, ''), ' ', COALESCE(e.apellidos, ''))) AS employee_name,
                COALESCE(e.moneda, 'CRC') AS currency,
                COALESCE(e.salario, 0) AS amount,
                'MASTER_DATA' AS source
            FROM empleados e
            WHERE COALESCE(e.estado, 'Activo') = 'Activo'
              AND COALESCE(e.salario, 0) > 0
              AND NOT EXISTS (
                  SELECT 1 FROM payroll p WHERE p.usuario = e.usuario
              )
        )
        SELECT * FROM payroll
        UNION ALL
        SELECT * FROM employee_budget
        ORDER BY employee_name
    """, (year, month))
    salary_date = date(year, month, calendar.monthrange(year, month)[1]).isoformat()
    for row in salary_rows:
        preview.append({
            "payee_name": row.get("employee_name") or row.get("usuario") or "Empleado",
            "concept": "SALARIO MENSUAL",
            "amount": _f(row.get("amount")),
            "currency": row.get("currency") or "CRC",
            "issue_date": salary_date,
            "due_date": None,
            "source": row.get("source") or "PAYROLL",
            "accepted": True,
        })
    return {"period": period, "source": "suggested", "data": preview}


def save_monthly_obligations(conn, year: int, month: int, rows, user=None):
    _ensure_monthly_report_obligations(conn)
    period = f"{year}-{month:02d}"
    cur = conn.cursor(cursor_factory=RealDictCursor)
    cur.execute("DELETE FROM monthly_report_obligations WHERE period = %s", (period,))
    inserted = []
    for row in rows or []:
        if not row.get("accepted", True):
            continue
        amount = _f(row.get("amount"))
        if amount <= 0:
            continue
        issue_date = row.get("issue_date") or row.get("due_date") or None
        due_date = row.get("due_date") or None
        cur.execute(
            """
            INSERT INTO monthly_report_obligations (
                period, payee_name, concept, amount, currency, issue_date, due_date, source, accepted, created_by, updated_at
            )
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,TRUE,%s,now())
            RETURNING id, payee_name, concept, amount, currency, issue_date, due_date, source, accepted
            """,
            (
                period,
                row.get("payee_name") or "N/A",
                row.get("concept") or "Obligacion mensual",
                amount,
                row.get("currency") or "USD",
                issue_date,
                due_date,
                row.get("source") or "PREVIEW",
                user,
            ),
        )
        inserted.append(dict(cur.fetchone()))
    conn.commit()
    return {"period": period, "saved": len(inserted), "data": inserted}


def _accounting_snapshot(cur, period: str):
    cur.execute("""
        SELECT
            COALESCE(SUM(CASE WHEN l.account_code LIKE '4%%' THEN l.credit - l.debit ELSE 0 END),0) AS accounting_revenue,
            COALESCE(SUM(CASE WHEN l.account_code LIKE '5%%' OR l.account_code LIKE '6%%' THEN l.debit - l.credit ELSE 0 END),0) AS accounting_expense,
            COALESCE(SUM(CASE WHEN l.account_code LIKE '1%%' THEN l.debit - l.credit ELSE 0 END),0) AS assets,
            COALESCE(SUM(CASE WHEN l.account_code LIKE '2%%' THEN l.credit - l.debit ELSE 0 END),0) AS liabilities,
            COALESCE(SUM(CASE WHEN l.account_code LIKE '3%%' THEN l.credit - l.debit ELSE 0 END),0) AS equity,
            COALESCE(SUM(CASE WHEN l.account_code LIKE '1.1.02%%' OR LOWER(l.account_name) LIKE '%%banco%%'
                              THEN l.debit - l.credit ELSE 0 END),0) AS bank_balance,
            COUNT(DISTINCT e.id) AS posted_entries
        FROM accounting_entries e
        JOIN accounting_lines l ON l.entry_id = e.id
        WHERE e.period = %s
          AND COALESCE(e.workflow_status, 'POSTED') = 'POSTED'
    """, (period,))
    row = cur.fetchone() or {}
    revenue = _f(row.get("accounting_revenue"))
    expense = _f(row.get("accounting_expense"))
    net_income = revenue - expense
    assets = _f(row.get("assets"))
    liabilities = _f(row.get("liabilities"))
    equity = _f(row.get("equity"))
    working_capital = assets - liabilities
    return {
        "accounting_revenue": revenue,
        "accounting_expense": expense,
        "net_income": net_income,
        "assets": assets,
        "liabilities": liabilities,
        "equity": equity,
        "bank_balance": _f(row.get("bank_balance")),
        "posted_entries": int(row.get("posted_entries") or 0),
        "working_capital": working_capital,
        "net_margin_pct": (net_income / revenue * 100.0) if revenue else None,
        "debt_to_assets_pct": (liabilities / assets * 100.0) if assets else None,
        "equity_ratio_pct": (equity / assets * 100.0) if assets else None,
    }


def _payments_cte():
    return """
        WITH payments AS (
            SELECT ca.codigo_cliente, ca.nombre_cliente, ca.fecha_pago, ca.monto_pagado AS amount, ca.numero_documento
            FROM cash_app ca
            WHERE ca.monto_pagado > 0 AND ca.tipo_aplicacion = 'PAGO'
            UNION ALL
            SELECT ip.codigo_cliente, ip.nombre_cliente, ip.fecha_pago, ip.monto AS amount, ip.documento AS numero_documento
            FROM incoming_payments ip
            WHERE ip.monto > 0
              AND NOT EXISTS (
                  SELECT 1 FROM cash_app ca2
                  WHERE ltrim(ca2.numero_documento, '0') = ltrim(ip.documento, '0')
                    AND ca2.codigo_cliente = ip.codigo_cliente
                    AND COALESCE(ca2.referencia, '') = COALESCE(ip.numero_referencia, '')
                    AND ca2.fecha_pago = ip.fecha_pago
                    AND ca2.monto_pagado = ip.monto
              )
        )
    """


def build_monthly_financial_data(conn, year: int, month: int):
    start, end, prev_start, prev_end, next_start, next_end = _period_bounds(year, month)
    cur = conn.cursor(cursor_factory=RealDictCursor)
    _ensure_monthly_report_obligations(conn)

    revenue = _fetch_one(cur, """
        SELECT COALESCE(SUM(total), 0) AS total, COUNT(*) AS count
        FROM invoicing
        WHERE tipo_documento = 'FACTURA'
          AND fecha_emision BETWEEN %s AND %s
    """, (start, end))

    prev_revenue = _fetch_one(cur, """
        SELECT COALESCE(SUM(total), 0) AS total, COUNT(*) AS count
        FROM invoicing
        WHERE tipo_documento = 'FACTURA'
          AND fecha_emision BETWEEN %s AND %s
    """, (prev_start, prev_end))

    collections = _fetch_one(cur, _payments_cte() + """
        SELECT COALESCE(SUM(amount), 0) AS total, COUNT(*) AS count
        FROM payments
        WHERE fecha_pago BETWEEN %s AND %s
    """, (start, end))

    prev_collections = _fetch_one(cur, _payments_cte() + """
        SELECT COALESCE(SUM(amount), 0) AS total, COUNT(*) AS count
        FROM payments
        WHERE fecha_pago BETWEEN %s AND %s
    """, (prev_start, prev_end))

    top_collections = _fetch_all(cur, _payments_cte() + """
        SELECT nombre_cliente, COALESCE(SUM(amount), 0) AS total
        FROM payments
        WHERE fecha_pago BETWEEN %s AND %s
        GROUP BY nombre_cliente
        ORDER BY total DESC
        LIMIT 7
    """, (start, end))

    ar = _fetch_one(cur, """
        SELECT COALESCE(SUM(saldo_pendiente), 0) AS total, COUNT(*) AS count
        FROM collections
        WHERE saldo_pendiente > 0
          AND tipo_documento = 'FACTURA'
          AND fecha_emision <= %s
    """, (end,))

    top_ar = _fetch_all(cur, """
        SELECT nombre_cliente, COALESCE(SUM(saldo_pendiente), 0) AS total
        FROM collections
        WHERE saldo_pendiente > 0
          AND tipo_documento = 'FACTURA'
          AND fecha_emision <= %s
        GROUP BY nombre_cliente
        ORDER BY total DESC
        LIMIT 7
    """, (end,))

    ar_aging = _fetch_all(cur, """
        SELECT bucket AS nombre_cliente, COALESCE(SUM(saldo_pendiente), 0) AS total
        FROM (
            SELECT saldo_pendiente,
                   CASE
                       WHEN fecha_vencimiento IS NULL THEN 'No due date'
                       WHEN (%s::date - fecha_vencimiento) <= 0 THEN 'Current'
                       WHEN (%s::date - fecha_vencimiento) BETWEEN 1 AND 30 THEN '1-30 days'
                       WHEN (%s::date - fecha_vencimiento) BETWEEN 31 AND 60 THEN '31-60 days'
                       WHEN (%s::date - fecha_vencimiento) BETWEEN 61 AND 90 THEN '61-90 days'
                       ELSE 'Over 90 days'
                   END AS bucket
            FROM collections
            WHERE saldo_pendiente > 0
              AND tipo_documento = 'FACTURA'
              AND fecha_emision <= %s
        ) x
        GROUP BY bucket
        ORDER BY MIN(CASE bucket
            WHEN 'Current' THEN 1
            WHEN '1-30 days' THEN 2
            WHEN '31-60 days' THEN 3
            WHEN '61-90 days' THEN 4
            WHEN 'Over 90 days' THEN 5
            ELSE 6
        END)
    """, (end, end, end, end, end))

    payment_trend = _fetch_all(cur, _payments_cte() + """
        SELECT bucket AS nombre_cliente, COUNT(*)::numeric AS total
        FROM (
            SELECT CASE
                       WHEN i.fecha_emision IS NULL THEN 'Unmatched'
                       WHEN (p.fecha_pago::date - i.fecha_emision::date) <= 30 THEN '0-30 days'
                       WHEN (p.fecha_pago::date - i.fecha_emision::date) <= 60 THEN '31-60 days'
                       WHEN (p.fecha_pago::date - i.fecha_emision::date) <= 90 THEN '61-90 days'
                       WHEN (p.fecha_pago::date - i.fecha_emision::date) <= 120 THEN '91-120 days'
                       ELSE 'Over 120 days'
                   END AS bucket
            FROM payments p
            LEFT JOIN invoicing i
              ON ltrim(i.numero_documento, '0') = ltrim(p.numero_documento, '0')
             AND i.codigo_cliente = p.codigo_cliente
             AND i.tipo_documento = 'FACTURA'
            WHERE p.fecha_pago <= %s
        ) x
        GROUP BY bucket
        ORDER BY MIN(CASE bucket
            WHEN '0-30 days' THEN 1
            WHEN '31-60 days' THEN 2
            WHEN '61-90 days' THEN 3
            WHEN '91-120 days' THEN 4
            WHEN 'Over 120 days' THEN 5
            ELSE 6
        END)
    """, (end,))

    billing_trend = _fetch_all(cur, """
        SELECT TO_CHAR(DATE_TRUNC('month', fecha_emision), 'Mon YYYY') AS nombre_cliente,
               COALESCE(SUM(total), 0) AS total,
               COUNT(*) AS count
        FROM invoicing
        WHERE tipo_documento = 'FACTURA'
          AND fecha_emision BETWEEN %s AND %s
        GROUP BY DATE_TRUNC('month', fecha_emision)
        ORDER BY DATE_TRUNC('month', fecha_emision)
    """, (date(year, 1, 1), end))

    prev_year_trend = _fetch_all(cur, """
        SELECT EXTRACT(MONTH FROM fecha_emision)::int AS month_num,
               TO_CHAR(DATE_TRUNC('month', fecha_emision), 'Mon') AS nombre_cliente,
               COALESCE(SUM(total), 0) AS total
        FROM invoicing
        WHERE tipo_documento = 'FACTURA'
          AND fecha_emision BETWEEN %s AND %s
        GROUP BY EXTRACT(MONTH FROM fecha_emision), DATE_TRUNC('month', fecha_emision)
        ORDER BY EXTRACT(MONTH FROM fecha_emision)
    """, (date(year - 1, 1, 1), date(year - 1, month, calendar.monthrange(year - 1, month)[1])))

    current_year_trend = _fetch_all(cur, """
        SELECT EXTRACT(MONTH FROM fecha_emision)::int AS month_num,
               TO_CHAR(DATE_TRUNC('month', fecha_emision), 'Mon') AS nombre_cliente,
               COALESCE(SUM(total), 0) AS total
        FROM invoicing
        WHERE tipo_documento = 'FACTURA'
          AND fecha_emision BETWEEN %s AND %s
        GROUP BY EXTRACT(MONTH FROM fecha_emision), DATE_TRUNC('month', fecha_emision)
        ORDER BY EXTRACT(MONTH FROM fecha_emision)
    """, (date(year, 1, 1), end))

    top_billing = _fetch_all(cur, """
        SELECT nombre_cliente, COALESCE(SUM(total), 0) AS total
        FROM invoicing
        WHERE tipo_documento = 'FACTURA'
          AND fecha_emision BETWEEN %s AND %s
        GROUP BY nombre_cliente
        ORDER BY total DESC
        LIMIT 7
    """, (start, end))

    itp_paid = _fetch_one(cur, """
        SELECT COALESCE(SUM(
            CASE WHEN currency = 'CRC' THEN (total - balance) / 500.0 ELSE (total - balance) END
        ), 0) AS total, COUNT(*) AS count
        FROM payment_obligations
        WHERE status IN ('PAID','PARTIAL')
          AND last_payment_date BETWEEN %s AND %s
    """, (start, end))

    payables_open = _fetch_one(cur, """
        SELECT COALESCE(SUM(
            CASE WHEN currency = 'CRC' THEN balance / 500.0 ELSE balance END
        ), 0) AS total, COUNT(*) AS count
        FROM payment_obligations
        WHERE status IN ('PENDING','PARTIAL')
          AND due_date <= %s
    """, (end,))

    top_payables_open = _fetch_all(cur, """
        SELECT payee_name AS nombre_cliente, COALESCE(SUM(
            CASE WHEN currency = 'CRC' THEN balance / 500.0 ELSE balance END
        ), 0) AS total
        FROM payment_obligations
        WHERE status IN ('PENDING','PARTIAL')
          AND due_date <= %s
        GROUP BY payee_name
        ORDER BY total DESC
        LIMIT 7
    """, (end,))

    next_receivables = _fetch_one(cur, """
        SELECT COALESCE(SUM(saldo_pendiente), 0) AS total, COUNT(*) AS count
        FROM collections
        WHERE saldo_pendiente > 0
          AND tipo_documento = 'FACTURA'
          AND fecha_vencimiento BETWEEN %s AND %s
    """, (next_start, next_end))

    top_next_receivables = _fetch_all(cur, """
        SELECT nombre_cliente, COALESCE(SUM(saldo_pendiente), 0) AS total
        FROM collections
        WHERE saldo_pendiente > 0
          AND tipo_documento = 'FACTURA'
          AND fecha_vencimiento BETWEEN %s AND %s
        GROUP BY nombre_cliente
        ORDER BY total DESC
        LIMIT 7
    """, (next_start, next_end))

    next_payables = _fetch_one(cur, """
        SELECT COALESCE(SUM(
            CASE WHEN currency = 'CRC' THEN balance / 500.0 ELSE balance END
        ), 0) AS total, COUNT(*) AS count
        FROM payment_obligations
        WHERE status IN ('PENDING','PARTIAL')
          AND due_date BETWEEN %s AND %s
    """, (next_start, next_end))

    top_next_payables = _fetch_all(cur, """
        SELECT payee_name AS nombre_cliente, COALESCE(SUM(
            CASE WHEN currency = 'CRC' THEN balance / 500.0 ELSE balance END
        ), 0) AS total
        FROM payment_obligations
        WHERE status IN ('PENDING','PARTIAL')
          AND due_date BETWEEN %s AND %s
        GROUP BY payee_name
        ORDER BY total DESC
        LIMIT 7
    """, (next_start, next_end))

    avg_days = _fetch_one(cur, _payments_cte() + """
        SELECT ROUND(AVG(p.fecha_pago::date - i.fecha_emision::date), 1) AS days
        FROM payments p
        JOIN invoicing i
          ON ltrim(i.numero_documento, '0') = ltrim(p.numero_documento, '0')
         AND i.codigo_cliente = p.codigo_cliente
         AND i.tipo_documento = 'FACTURA'
        WHERE p.fecha_pago BETWEEN %s AND %s
    """, (start, end))

    accounting = _accounting_snapshot(cur, f"{year}-{month:02d}")

    monthly_obligations = _fetch_all(cur, """
        SELECT
            payee_name AS nombre_cliente,
            concept,
            amount AS total,
            currency,
            COALESCE(issue_date, due_date) AS issue_date,
            due_date,
            source
        FROM monthly_report_obligations
        WHERE period = %s
          AND accepted = TRUE
        ORDER BY COALESCE(issue_date, due_date) NULLS LAST, payee_name
    """, (f"{year}-{month:02d}",))

    cur.close()

    revenue_total = _f(revenue.get("total"))
    prev_revenue_total = _f(prev_revenue.get("total"))
    collections_total = _f(collections.get("total"))
    prev_collections_total = _f(prev_collections.get("total"))
    ar_total = _f(ar.get("total"))
    itp_total = _f(itp_paid.get("total"))
    payables_open_total = _f(payables_open.get("total"))
    next_receivables_total = _f(next_receivables.get("total"))
    next_payables_total = _f(next_payables.get("total"))
    net_cash = collections_total - itp_total
    next_net_outlook = next_receivables_total - next_payables_total
    month_name = MONTH_NAMES_ES.get(month, calendar.month_name[month])
    prev_month_name = MONTH_NAMES_ES.get(prev_start.month, calendar.month_name[prev_start.month])
    next_month_name = MONTH_NAMES_ES.get(next_start.month, calendar.month_name[next_start.month])

    data = {
        "period": {
            "year": year,
            "month": month,
            "month_name": month_name,
            "label": f"{month_name} {year}",
            "report_label": _month_label(year, month),
            "start": start.isoformat(),
            "end": end.isoformat(),
            "previous_label": f"{prev_month_name} {prev_start.year}",
            "next_label": f"{next_month_name} {next_start.year}",
        },
        "metrics": {
            "revenue": revenue_total,
            "revenue_count": int(revenue.get("count") or 0),
            "previous_revenue": prev_revenue_total,
            "collections": collections_total,
            "collections_count": int(collections.get("count") or 0),
            "previous_collections": prev_collections_total,
            "ar_open": ar_total,
            "ar_count": int(ar.get("count") or 0),
            "itp_paid": itp_total,
            "itp_paid_count": int(itp_paid.get("count") or 0),
            "payables_open": payables_open_total,
            "payables_open_count": int(payables_open.get("count") or 0),
            "next_receivables": next_receivables_total,
            "next_receivables_count": int(next_receivables.get("count") or 0),
            "next_payables": next_payables_total,
            "next_payables_count": int(next_payables.get("count") or 0),
            "next_net_outlook": next_net_outlook,
            "net_cash": net_cash,
            "avg_days_to_pay": _f(avg_days.get("days")) if avg_days.get("days") is not None else None,
        },
        "accounting": accounting,
        "tables": {
            "top_collections": [dict(r) for r in top_collections],
            "top_ar": [dict(r) for r in top_ar],
            "ar_aging": [dict(r) for r in ar_aging],
            "payment_trend": [dict(r) for r in payment_trend],
            "billing_trend": [dict(r) for r in billing_trend],
            "top_billing": [dict(r) for r in top_billing],
            "top_payables_open": [dict(r) for r in top_payables_open],
            "top_next_receivables": [dict(r) for r in top_next_receivables],
            "top_next_payables": [dict(r) for r in top_next_payables],
            "prev_year_trend": [dict(r) for r in prev_year_trend],
            "current_year_trend": [dict(r) for r in current_year_trend],
            "monthly_obligations": [dict(r) for r in monthly_obligations],
        },
    }
    data["executive"] = _build_executive_dashboard(data)
    data["narrative"] = _build_financial_narrative(data)
    return data


def _ratio_label(value):
    return "N/A" if value is None else f"{_f(value):.1f}%"


def _build_executive_dashboard(data):
    m = data["metrics"]
    a = data["accounting"]
    collection_ratio = (m["collections"] / m["revenue"] * 100.0) if m["revenue"] else None
    ar_pressure = (m["ar_open"] / m["revenue"] * 100.0) if m["revenue"] else None
    payable_coverage = (m["collections"] / m["payables_open"] * 100.0) if m["payables_open"] else None

    alerts = []
    if abs(a["net_income"]) > 0 and a["net_income"] < 0:
        alerts.append("Resultado contable negativo en el periodo; revisar estructura de gastos y margen operativo.")
    if ar_pressure is not None and ar_pressure > 100:
        alerts.append("La cartera abierta supera la facturacion mensual; priorizar cobranza y seguimiento por cliente.")
    if payable_coverage is not None and payable_coverage < 100:
        alerts.append("La cobranza del mes no cubre completamente las cuentas por pagar abiertas.")
    if m["next_net_outlook"] < 0:
        alerts.append("El outlook del proximo mes muestra presion neta de caja negativa.")
    if not alerts:
        alerts.append("No se detectan alertas ejecutivas criticas con los datos disponibles del periodo.")

    return {
        "collection_ratio_pct": collection_ratio,
        "ar_pressure_pct": ar_pressure,
        "payable_coverage_pct": payable_coverage,
        "net_margin_pct": a["net_margin_pct"],
        "debt_to_assets_pct": a["debt_to_assets_pct"],
        "equity_ratio_pct": a["equity_ratio_pct"],
        "alerts": alerts,
        "decision_focus": [
            "Confirmar recuperacion de clientes con mayor cartera abierta.",
            "Calendarizar pagos de mayor impacto contra caja esperada.",
            "Validar variaciones entre facturacion operativa y resultado contable.",
            "Revisar cuentas bancarias y auxiliares antes del cierre mensual.",
        ],
    }


def _compact_rows(rows, limit=6):
    return [
        {"name": str(r.get("nombre_cliente") or "N/A"), "amount": round(_f(r.get("total")), 2)}
        for r in (rows or [])[:limit]
    ]


def _build_financial_narrative(data):
    ai = _ai_financial_narrative(data)
    if ai:
        return ai
    return _fallback_financial_narrative(data)


def _ai_financial_narrative(data):
    try:
        from ai.maritime_ai import _get_openai_client

        client = _get_openai_client()
        metrics = data["metrics"]
        period = data["period"]
        payload = {
            "period": period,
            "metrics": metrics,
            "accounting": data.get("accounting", {}),
            "executive": data.get("executive", {}),
            "top_collections": _compact_rows(data["tables"]["top_collections"]),
            "top_ar": _compact_rows(data["tables"]["top_ar"]),
            "ar_aging": _compact_rows(data["tables"]["ar_aging"]),
            "payment_trend": _compact_rows(data["tables"]["payment_trend"]),
            "billing_trend": _compact_rows(data["tables"]["billing_trend"]),
            "top_billing": _compact_rows(data["tables"]["top_billing"]),
            "top_payables_open": _compact_rows(data["tables"]["top_payables_open"]),
            "top_next_receivables": _compact_rows(data["tables"]["top_next_receivables"]),
            "top_next_payables": _compact_rows(data["tables"]["top_next_payables"]),
            "prev_year_trend": _compact_rows(data["tables"]["prev_year_trend"]),
            "current_year_trend": _compact_rows(data["tables"]["current_year_trend"]),
        }
        prompt = (
            "Eres CFO y analista financiero de Marine Surveyors & Logistics. "
            "Genera narrativa ejecutiva en espanol para un reporte financiero mensual, usando SOLO los datos entregados. "
            "No inventes numeros. Si un dato es cero, explicalo con prudencia como falta de registros o sin actividad registrada. "
            "Incluye lectura CFO de liquidez, rentabilidad, capital de trabajo, margen, cobertura de cuentas por pagar, "
            "presion de cartera, riesgos y decisiones recomendadas. "
            "Devuelve JSON valido con estas llaves exactas: introduction, collections, receivables, payment_trend, "
            "billing, payables, next_month_outlook, year_comparison, risk, conclusion. "
            "La introduction debe ser amplia: 5 a 7 parrafos ejecutivos, con contexto, alcance del analisis, "
            "lectura de liquidez, facturacion, cobranza, cuentas por cobrar, cuentas por pagar, riesgos y objetivo gerencial. "
            "La conclusion debe ser igualmente amplia: 5 a 7 parrafos ejecutivos, con cierre estrategico, riesgos, "
            "acciones recomendadas, prioridades de caja, seguimiento comercial y lectura del siguiente mes. "
            "Las demas secciones deben tener 2 a 4 parrafos cortos, estilo informe ejecutivo, con interpretacion "
            "y recomendaciones cuando aplique."
        )
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.25,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
        )
        text = response.choices[0].message.content.strip()
        if text.startswith("```"):
            text = text.strip("`")
            text = text.split("\n", 1)[-1]
        parsed = json.loads(text)
        required = [
            "introduction", "collections", "receivables", "payment_trend", "billing",
            "payables", "next_month_outlook", "year_comparison", "risk", "conclusion",
        ]
        return {key: str(parsed.get(key) or "").strip() for key in required}
    except Exception:
        return None


def _fallback_financial_narrative(data):
    p = data["period"]
    m = data["metrics"]
    billing_pct = _pct(m["revenue"], m["previous_revenue"])
    collections_pct = _pct(m["collections"], m["previous_collections"])
    concentration = _compact_rows(data["tables"]["top_ar"], 2)
    top_ar = ", ".join(f"{r['name']} ({_short_money(r['amount'])})" for r in concentration) or "sin concentración visible"
    top_collections = ", ".join(
        f"{r['name']} ({_short_money(r['amount'])})" for r in _compact_rows(data["tables"]["top_collections"], 3)
    ) or "sin recuperaciones registradas"

    billing_sentence = (
        "no cuenta con comparativo del mes anterior"
        if billing_pct is None else f"vario {billing_pct:+.1f}% contra el mes anterior"
    )
    collections_sentence = (
        "no cuenta con comparativo del mes anterior"
        if collections_pct is None else f"vario {collections_pct:+.1f}% contra el mes anterior"
    )
    risk_level = "controlado" if m["next_net_outlook"] >= 0 and m["ar_open"] >= m["payables_open"] else "sensible"

    return {
        "introduction": (
            f"El presente análisis evalúa la situación financiera, comercial y operativa de la compañía al cierre de "
            f"{p['label']}, considerando facturación, recuperación de cartera, cuentas por cobrar, cuentas por pagar, "
            f"liquidez operativa y riesgos asociados a la continuidad del negocio.\n\n"
            f"Durante el período se registraron {_money(m['revenue'])} en facturación y {_money(m['collections'])} "
            f"en cobranzas. El resultado debe leerse junto con una cartera abierta de {_money(m['ar_open'])} y "
            f"obligaciones pendientes por {_money(m['payables_open'])}, lo cual permite dimensionar la capacidad "
            f"de caja y la presión operativa de corto plazo.\n\n"
            f"El objetivo de este reporte es traducir los movimientos financieros del mes en una lectura ejecutiva "
            f"que permita tomar decisiones oportunas. Por ello, el análisis no se limita a presentar montos; también "
            f"interpreta la relación entre generación de ingresos, velocidad de cobranza, exposición de cartera y "
            f"compromisos de pago.\n\n"
            f"La posición de liquidez debe observarse con especial atención porque el flujo disponible depende de dos "
            f"fuentes principales: los cobros efectivamente recuperados durante el mes y la capacidad de convertir "
            f"las cuentas por cobrar abiertas en efectivo dentro de los plazos previstos. Cuando la cobranza se "
            f"concentra en pocos clientes, el seguimiento comercial y financiero se vuelve un control crítico.\n\n"
            f"Desde la perspectiva operativa, la lectura de {p['label']} permite identificar si la compañía cuenta con "
            f"una base suficiente para sostener sus obligaciones recurrentes, cubrir compromisos extraordinarios y "
            f"mantener capacidad de respuesta ante atrasos o cambios en el calendario de pagos.\n\n"
            f"El reporte también incorpora una vista prospectiva de {p['next_label']}, con el fin de anticipar presión "
            f"de caja, priorizar gestiones de cobro y calendarizar cuentas por pagar de forma disciplinada. Esta "
            f"visión busca que la administración actúe de manera preventiva y no únicamente reactiva."
        ),
        "collections": (
            f"Con base en los resultados de cobranza del período, se recuperaron {_money(m['collections'])} en "
            f"{m['collections_count']} registros de pago. Las principales contribuciones fueron {top_collections}.\n\n"
            f"La cobranza {collections_sentence}. Este comportamiento refleja la efectividad del seguimiento de cartera "
            f"y permite identificar si la liquidez depende de pocos clientes o de una base mas diversificada."
        ),
        "receivables": (
            f"Las cuentas por cobrar abiertas ascienden a {_money(m['ar_open'])} distribuidas en {m['ar_count']} facturas. "
            f"Los principales saldos se concentran en {top_ar}.\n\n"
            f"Esta composición permite priorizar gestiones de alto impacto. Mientras mayor sea la concentración en pocos "
            f"clientes, mayor debe ser la disciplina de seguimiento, confirmación de fechas de pago y control de riesgo crediticio."
        ),
        "payment_trend": (
            f"La tendencia de pago se evalúa con base en la antigüedad entre emisión de factura y recuperación. "
            f"El promedio observado para el período es "
            f"{m['avg_days_to_pay']:.1f} dias." if m["avg_days_to_pay"] is not None else
            "La tendencia de pago no cuenta con información suficiente para calcular un promedio confiable en este período."
        ),
        "billing": (
            f"La facturación de {p['label']} alcanzó {_money(m['revenue'])} en {m['revenue_count']} facturas y "
            f"{billing_sentence}. Este indicador muestra el pulso comercial del mes y permite evalúar si la compañía "
            f"mantiene suficiente generacion de ingresos para sostener su estructura operativa.\n\n"
            f"Cuando la facturación se desacelera, la compañía queda mas expuesta a la recuperación de cartera previa. "
            f"Por ello, el análisis debe observar simultaneamente ventas, cobros y cuentas por pagar."
        ),
        "payables": (
            f"Las cuentas por pagar abiertas al cierre ascienden a {_money(m['payables_open'])}, mientras que los pagos "
            f"registrados durante el período suman {_money(m['itp_paid'])}. Esta estructura evidencia la carga de "
            f"compromisos que debe administrarse frente a la caja recuperada.\n\n"
            f"Los rubros de mayor peso deben revisarse por recurrencia, necesidad operativa y posibilidad de negociacion "
            f"sin afectar la calidad del servicio."
        ),
        "next_month_outlook": (
            f"Para {p['next_label']}, la vista prospectiva muestra {_money(m['next_receivables'])} en cuentas por cobrar "
            f"esperadas por vencimiento y {_money(m['next_payables'])} en cuentas por pagar programadas. El resultado neto "
            f"proyectado es {_money(m['next_net_outlook'])} antes de nuevas facturas, cobros adicionales o ajustes posteriores.\n\n"
            f"Este outlook debe utilizarse como agenda de caja: priorizar cobros de mayor impacto, confirmar promesas de pago "
            f"y calendarizar obligaciones para evitar presión innecesaria al inicio del mes."
        ),
        "year_comparison": (
            f"El comparativo anual permite ubicar el desempeno de {p['label']} frente al comportamiento historico reciente. "
            f"La lectura principal es identificar si el mes responde a una tendencia sostenida, una recuperación puntual o "
            f"una desaceleracion que requiera acciones comerciales.\n\n"
            f"El seguimiento mensual debe enfocarse en volumen facturado, cantidad de operaciones, ticket promedio y "
            f"concentración por cliente."
        ),
        "risk": (
            f"El riesgo financiero actual se considera {risk_level}. La compañía debe observar tres factores: dependencia "
            f"de recuperación de cartera, concentración de ingresos en pocos clientes y compromisos fijos que consumen caja "
            f"de forma recurrente.\n\n"
            f"Cualquier atraso material en cobranza puede reducir el margen de maniobra. Por tanto, se recomienda mantener "
            f"contencion de gastos no esenciales y seguimiento semanal de cartera y obligaciones."
        ),
        "conclusion": (
            f"El cierre de {p['label']} muestra una posición que debe gestionarse con disciplina: la estabilidad depende de "
            f"convertir cuentas por cobrar en efectivo, sostener la facturación y calendarizar adecuadamente los pagos.\n\n"
            f"La recomendación ejecutiva es reforzar el seguimiento de clientes relevantes, proteger caja, evitar compromisos "
            f"no esenciales y utilizar el outlook de {p['next_label']} como base de planificación financiera.\n\n"
            f"En términos de gestión, la prioridad inmediata debe ser asegurar la recuperación de los saldos de mayor "
            f"impacto y confirmar fechas de pago con los clientes que concentran la cartera. Esta acción permite "
            f"reducir incertidumbre, mejorar previsibilidad de caja y sostener una operación más ordenada.\n\n"
            f"De forma paralela, las cuentas por pagar deben administrarse bajo un calendario realista, priorizando "
            f"obligaciones críticas para la continuidad del servicio y revisando aquellos compromisos que puedan "
            f"renegociarse, diferirse o ajustarse sin afectar la calidad operativa.\n\n"
            f"El análisis también confirma la importancia de monitorear la facturación mensual. Si la generación de "
            f"ingresos no mantiene un ritmo suficiente, la empresa dependerá cada vez más de cartera previa, lo cual "
            f"puede limitar el margen de maniobra ante gastos extraordinarios o atrasos de clientes.\n\n"
            f"Por lo tanto, la administración debería usar este reporte como tablero de control mensual: validar "
            f"cobranza, revisar exposición por cliente, medir comportamiento de pago, controlar obligaciones y "
            f"comparar el desempeño contra meses anteriores.\n\n"
            f"La conclusión estratégica es que la compañía puede mantener una posición operativa saludable si protege "
            f"la liquidez, ejecuta la cobranza con rigor y evita compromisos que no estén directamente alineados con "
            f"la generación de ingresos, continuidad del servicio o reducción de riesgo financiero."
        ),
    }


def _safe_filename(label, extension):
    return f"MSL_Financial_Report_{label.replace(' ', '_').replace(',', '')}.{extension}"


def _asset_path(name):
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    candidates = [
        os.path.join(root, "assets", name),
        os.path.join(os.path.dirname(__file__), "..", "assets", name),
    ]
    for candidate in candidates:
        candidate = os.path.abspath(candidate)
        if os.path.exists(candidate):
            return candidate
    return None


def _canvas_frame(canvas, doc):
    from reportlab.lib import colors

    canvas.saveState()
    width, height = doc.pagesize
    canvas.setFillColor(colors.HexColor("#F1F1F1"))
    canvas.rect(0, 0, width, height, stroke=0, fill=1)
    canvas.setStrokeColor(colors.black)
    canvas.setLineWidth(2)
    canvas.line(170, height - 86, width - 30, height - 86)
    canvas.line(170, 42, width - 30, 42)
    canvas.line(0, height - 86, 50, height - 86)
    canvas.line(0, 42, 50, 42)
    canvas.restoreState()


try:
    from reportlab.platypus import Flowable as _ReportFlowable
except Exception:
    _ReportFlowable = object


class _RotatedSideLabel(_ReportFlowable):
    def __init__(self, text):
        try:
            super().__init__()
        except Exception:
            pass
        self.text = text
        self.width = 42
        self.height = 520

    def wrap(self, avail_width, avail_height):
        return self.width, min(self.height, avail_height)

    def drawOn(self, canvas, x, y, _sW=0):
        canvas.saveState()
        canvas.translate(x + 24, y + 38)
        canvas.rotate(90)
        canvas.setFont("Helvetica-Bold", 17)
        canvas.setFillColorRGB(0, 0, 0)
        canvas.drawString(0, 0, self.text)
        canvas.restoreState()


def _chart_image(rows, title, path, kind="bar"):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None

    rows = [r for r in (rows or []) if _f(r.get("total")) > 0][:7]
    width, height = 900, 430
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 28)
        label_font = ImageFont.truetype("arial.ttf", 17)
        small_font = ImageFont.truetype("arial.ttf", 15)
    except Exception:
        title_font = label_font = small_font = None

    draw.text((30, 22), title, fill=DARK, font=title_font)
    draw.line((30, 64, width - 30, 64), fill="#D9E2EC", width=2)
    if not rows:
        draw.text((330, 205), "No records for this period", fill="#697386", font=label_font)
        try:
            img.save(path)
            return path
        except Exception:
            return None

    if kind == "pie":
        total = sum(_f(r.get("total")) for r in rows) or 1
        start_angle = 0
        box = (70, 105, 360, 395)
        for idx, row in enumerate(rows):
            amount = _f(row.get("total"))
            angle = amount / total * 360
            draw.pieslice(box, start=start_angle, end=start_angle + angle, fill=CHART_COLORS[idx % len(CHART_COLORS)])
            start_angle += angle
        legend_x = 430
        for idx, row in enumerate(rows):
            y = 115 + idx * 38
            draw.rectangle((legend_x, y, legend_x + 22, y + 22), fill=CHART_COLORS[idx % len(CHART_COLORS)])
            pct = _f(row.get("total")) / total * 100
            label = str(row.get("nombre_cliente") or "N/A")[:28]
            draw.text((legend_x + 32, y - 1), f"{label} - {pct:.1f}% - {_short_money(row.get('total'))}", fill=DARK, font=small_font)
    else:
        max_value = max(_f(r.get("total")) for r in rows) or 1
        chart_left, chart_top, chart_right, bar_h = 260, 96, 720, 30
        for idx, row in enumerate(rows):
            y = chart_top + idx * 45
            label = str(row.get("nombre_cliente") or "N/A")[:24]
            draw.text((35, y + 5), label, fill=DARK, font=small_font)
            bar_w = int((_f(row.get("total")) / max_value) * (chart_right - chart_left))
            color = CHART_COLORS[idx % len(CHART_COLORS)]
            draw.rounded_rectangle((chart_left, y, chart_left + bar_w, y + bar_h), radius=6, fill=color)
            draw.text((chart_left + bar_w + 10, y + 5), _short_money(row.get("total")), fill=DARK, font=small_font)
        draw.line((chart_left, chart_top - 10, chart_left, chart_top + len(rows) * 45), fill="#BCCCDC", width=1)
    try:
        img.save(path)
        return path
    except Exception:
        return None


def _build_charts(data, tmp_dir):
    tables = data["tables"]
    metrics = data["metrics"]
    accounting = data["accounting"]
    charts = {
        "collections": _chart_image(tables["top_collections"], "Distribución de cuentas por cobrar recuperadas", os.path.join(tmp_dir, "collections.png"), "pie"),
        "ar": _chart_image(tables["top_ar"], "Cuentas por cobrar a recuperar", os.path.join(tmp_dir, "ar.png"), "bar"),
        "ar_aging": _chart_image(tables["ar_aging"], "Antigüedad de cuentas por cobrar", os.path.join(tmp_dir, "ar_aging.png"), "bar"),
        "payment_trend": _chart_image(tables["payment_trend"], "Tendencia de pago por antigüedad", os.path.join(tmp_dir, "payment_trend.png"), "bar"),
        "billing": _chart_image(tables["billing_trend"], "Facturación mensual", os.path.join(tmp_dir, "billing.png"), "bar"),
        "payables": _chart_image(tables["top_payables_open"], "Cuentas por pagar", os.path.join(tmp_dir, "payables.png"), "pie"),
        "next_ar": _chart_image(tables["top_next_receivables"], "Cobros esperados próximo mes", os.path.join(tmp_dir, "next_ar.png"), "bar"),
        "next_payables": _chart_image(tables["top_next_payables"], "Pagos programados próximo mes", os.path.join(tmp_dir, "next_payables.png"), "bar"),
        "cash_bridge": _chart_image([
            {"nombre_cliente": "Cobranza", "total": metrics["collections"]},
            {"nombre_cliente": "Pagos ITP", "total": metrics["itp_paid"]},
            {"nombre_cliente": "Caja neta", "total": metrics["net_cash"]},
            {"nombre_cliente": "AR abierto", "total": metrics["ar_open"]},
        ], "Puente de caja del periodo", os.path.join(tmp_dir, "cash_bridge.png"), "bar"),
        "accounting_mix": _chart_image([
            {"nombre_cliente": "Ingresos contables", "total": accounting["accounting_revenue"]},
            {"nombre_cliente": "Gastos contables", "total": accounting["accounting_expense"]},
            {"nombre_cliente": "Activos", "total": accounting["assets"]},
            {"nombre_cliente": "Pasivos", "total": accounting["liabilities"]},
        ], "Composición contable POSTED", os.path.join(tmp_dir, "accounting_mix.png"), "bar"),
    }
    return charts


def generate_monthly_financial_pdf(conn, year: int, month: int):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import PageBreak, SimpleDocTemplate

    data = build_monthly_financial_data(conn, year, month)
    label = data["period"]["label"]
    tmp_dir = tempfile.mkdtemp(prefix="erp_som_financial_report_")
    charts = _build_charts(data, tmp_dir)
    path = os.path.join(tmp_dir, _safe_filename(label, "pdf"))

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverCompany", parent=styles["Title"], fontSize=20, textColor=colors.black, leading=22, alignment=0))
    styles.add(ParagraphStyle(name="CoverTitle", parent=styles["Title"], fontSize=38, textColor=colors.black, leading=44, alignment=0))
    styles.add(ParagraphStyle(name="SectionLead", parent=styles["Heading1"], textColor=colors.black, fontSize=15.5, leading=19, spaceBefore=0, spaceAfter=8))
    styles.add(ParagraphStyle(name="Body", parent=styles["BodyText"], fontSize=14, leading=20, alignment=4, spaceAfter=12))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=9, leading=11, textColor=colors.HexColor("#333333")))

    story = []
    story.extend(_pdf_cover(data, styles))
    story.append(PageBreak())
    _pdf_intro_page(story, styles, data)
    _pdf_section(story, styles, "Recupero cuentas por cobrar", f"A continuación la distribución de cuentas por cobrar recuperadas en {data['period']['month_name']}:", data["narrative"]["collections"], charts["collections"])
    _pdf_section(story, styles, "Cuentas por cobrar", f"A continuación la distribución de cuentas por cobrar para {data['period']['label']}:", data["narrative"]["receivables"], charts["ar"])
    _pdf_section(story, styles, "Tendencia de pago", "A continuación la tendencia de pagos por parte del cliente:", data["narrative"]["payment_trend"], charts["payment_trend"])
    _pdf_section(story, styles, "Facturación", "A continuación el desglose mensual acumulado de facturación para el año en curso:", data["narrative"]["billing"], charts["billing"])
    _pdf_section(story, styles, "Cuentas por pagar", "A continuación el detalle de cuentas por pagar:", data["narrative"]["payables"], charts["payables"])
    _pdf_section(story, styles, "Obligaciones del mes", "A continuación el preliminar editable de obligaciones mensuales aceptadas para este reporte:", "Incluye las obligaciones mensuales aprobadas antes de emitir el reporte, incluyendo salarios y compromisos recurrentes. Esta sección funciona como agenda de pagos y base de seguimiento administrativo.", charts["cash_bridge"], data["tables"]["monthly_obligations"])
    _pdf_section(story, styles, "Cronograma de pago", f"A continuación el cronograma de pago de {data['period']['next_label']}:", data["narrative"]["next_month_outlook"], charts["next_ar"], extra_chart=charts["next_payables"])
    _pdf_section(story, styles, f"Comparativo {year - 1} vs {year}", f"A continuación el comparativo entre {year - 1} y {year}:", data["narrative"]["year_comparison"], charts["accounting_mix"])
    _pdf_section(story, styles, "Análisis de riesgo financiero", "A continuación el análisis de riesgos financieros:", data["narrative"]["risk"], None, None)
    _pdf_section(story, styles, "Conclusión", "Conclusión y recomendaciones ejecutivas:", data["narrative"]["conclusion"], None, None)

    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=58, bottomMargin=52)
    doc.build(story, onFirstPage=_canvas_frame, onLaterPages=_canvas_frame)
    return path, _safe_filename(label, "pdf")


def _pdf_cover(data, styles):
    from reportlab.platypus import Image, Paragraph, Spacer, Table
    from reportlab.lib import colors

    p = data["period"]
    logo = Paragraph("● ●&nbsp;&nbsp;Marine&nbsp;&nbsp;&nbsp;&nbsp;Surveyors&nbsp;&nbsp;&nbsp;&nbsp;&amp;<br/>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;Logistics", styles["CoverCompany"])
    items = [
        Spacer(1, 24),
        logo,
        Spacer(1, 22),
        Paragraph("Alajuela, Costa Rica", styles["Body"]),
        Spacer(1, 72),
        Paragraph(escape(p["report_label"]), styles["CoverTitle"]),
        Spacer(1, 44),
        Paragraph("Aarón Ávila Vargas", styles["SectionLead"]),
        Spacer(1, 24),
    ]
    ship = _asset_path("barco.jpg")
    if ship:
        items.append(Image(ship, width=455, height=250))
    else:
        items.append(_pdf_kpi_table(data))
    return items


def _pdf_intro_page(story, styles, data):
    from reportlab.platypus import PageBreak, Paragraph, Spacer

    story.append(Paragraph("Introducción ejecutiva", styles["SectionLead"]))
    for paragraph in str(data["narrative"]["introduction"] or "").split("\n\n"):
        if paragraph.strip():
            story.append(Paragraph(escape(paragraph.strip()), styles["Body"]))
    story.append(Spacer(1, 8))
    story.append(_pdf_kpi_table(data))
    story.append(PageBreak())


def _pdf_kpi_table(data):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    m = data["metrics"]
    rows = [
        ["Facturación", "Cobranza", "Cartera", "CxP", "Outlook neto"],
        [_short_money(m["revenue"]), _short_money(m["collections"]), _short_money(m["ar_open"]), _short_money(m["payables_open"]), _short_money(m["next_net_outlook"])],
    ]
    table = Table(rows, colWidths=[92, 92, 92, 92, 92])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BLUE)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor(GREY)),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D7DEE8")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
    ]))
    return table


def _pdf_executive_dashboard(story, styles, data):
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    m = data["metrics"]
    a = data["accounting"]
    e = data["executive"]
    story.append(Paragraph("Executive Financial Dashboard", styles["Section"]))
    story.append(Paragraph(
        "Vista ejecutiva de liquidez, rentabilidad, capital de trabajo y riesgos principales del periodo. "
        "Los indicadores combinan facturacion, cobranza, cuentas por cobrar, cuentas por pagar y movimientos contables POSTED.",
        styles["Body"],
    ))
    rows = [
        ["Indicador", "Resultado", "Lectura ejecutiva"],
        ["Facturacion del mes", _money(m["revenue"]), "Capacidad comercial registrada en el periodo"],
        ["Cobranza del mes", _money(m["collections"]), f"Recuperacion equivalente a {_ratio_label(e['collection_ratio_pct'])} de la facturacion"],
        ["Resultado contable", _money(a["net_income"]), f"Margen neto contable {_ratio_label(e['net_margin_pct'])}"],
        ["Capital de trabajo", _money(a["working_capital"]), "Activos menos pasivos segun asientos POSTED"],
        ["Cartera abierta", _money(m["ar_open"]), f"Presion de cartera {_ratio_label(e['ar_pressure_pct'])} vs facturacion"],
        ["CxP abiertas", _money(m["payables_open"]), f"Cobertura por cobranza {_ratio_label(e['payable_coverage_pct'])}"],
        ["Outlook neto proximo mes", _money(m["next_net_outlook"]), "Cobros esperados menos pagos programados"],
    ]
    table = Table(rows, colWidths=[150, 120, 260])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(BLUE)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D7DEE8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#FBFCFD")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(table)
    story.append(Spacer(1, 14))
    story.append(Paragraph("Alertas ejecutivas", styles["Section"]))
    for alert in e["alerts"]:
        story.append(Paragraph(f"- {escape(alert)}", styles["Body"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Focos de decision", styles["Section"]))
    for item in e["decision_focus"]:
        story.append(Paragraph(f"- {escape(item)}", styles["Body"]))


def _pdf_section(story, styles, side_title, lead, text, chart_path, table_rows=None, extra_chart=None):
    from reportlab.platypus import Image, PageBreak, Paragraph, Spacer

    story.append(Paragraph(escape(side_title), styles["SectionLead"]))
    story.append(Paragraph(escape(lead), styles["SectionLead"]))
    if chart_path and os.path.exists(chart_path):
        story.append(Spacer(1, 4))
        story.append(Image(chart_path, width=468, height=224))
        story.append(Spacer(1, 8))
    if extra_chart and os.path.exists(extra_chart):
        story.append(Image(extra_chart, width=468, height=224))
        story.append(Spacer(1, 8))
    if table_rows is not None:
        story.append(_pdf_table(table_rows))
        story.append(Spacer(1, 8))
    for paragraph in str(text or "").split("\n\n"):
        if paragraph.strip():
            story.append(Paragraph(escape(paragraph.strip()), styles["Body"]))
    story.append(PageBreak())


def _pdf_table(rows):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    has_obligation_fields = any((row or {}).get("concept") or (row or {}).get("issue_date") or (row or {}).get("due_date") for row in rows or [])
    if has_obligation_fields:
        table_data = [["Proveedor", "Concepto", "Fecha factura", "Monto"]]
        for row in rows or []:
            table_data.append([
                str(row.get("nombre_cliente") or row.get("name") or "N/A"),
                str(row.get("concept") or "Obligacion mensual"),
                str(row.get("issue_date") or row.get("due_date") or ""),
                f"{row.get('currency') or 'USD'} {_f(row.get('total') or row.get('amount')):,.2f}",
            ])
        if len(table_data) == 1:
            table_data.append(["No records for this period", "", "", "USD 0.00"])
        table = Table(table_data, colWidths=[160, 145, 80, 95])
    else:
        table_data = [["Detalle", "Monto"]]
        for row in rows or []:
            table_data.append([str(row.get("nombre_cliente") or row.get("name") or "N/A"), _money(row.get("total") or row.get("amount"))])
        if len(table_data) == 1:
            table_data.append(["No records for this period", "USD 0.00"])
        table = Table(table_data, colWidths=[330, 130])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(LIGHT_BLUE)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(DARK)),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D7DEE8")),
        ("ALIGN", (1, 1), (1, -1), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def _comparison_rows(data):
    rows = []
    current = {int(r.get("month_num") or 0): _f(r.get("total")) for r in data["tables"]["current_year_trend"]}
    previous = {int(r.get("month_num") or 0): _f(r.get("total")) for r in data["tables"]["prev_year_trend"]}
    for month_num in sorted(set(current) | set(previous)):
        month_name = calendar.month_abbr[month_num]
        rows.append({
            "nombre_cliente": month_name,
            "total": current.get(month_num, 0) - previous.get(month_num, 0),
        })
    return rows


def generate_monthly_financial_docx(conn, year: int, month: int):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    data = build_monthly_financial_data(conn, year, month)
    label = data["period"]["label"]
    tmp_dir = tempfile.mkdtemp(prefix="erp_som_financial_report_")
    charts = _build_charts(data, tmp_dir)
    path = os.path.join(tmp_dir, _safe_filename(label, "docx"))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Marine Surveyors &\nLogistics")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(18, 58, 99)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(data["period"]["report_label"])
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph("Alajuela, Costa Rica").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Aarón Ávila Vargas").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    _docx_executive_dashboard(doc, data)
    _docx_section(doc, "Introducción ejecutiva", data["narrative"]["introduction"])
    _docx_section(doc, "Collections Recovery", data["narrative"]["collections"], charts["collections"], data["tables"]["top_collections"])
    _docx_section(doc, "Cuentas por cobrar a recuperar", data["narrative"]["receivables"], charts["ar"], data["tables"]["top_ar"])
    _docx_section(doc, "Antigüedad de cartera", "Distribución de cartera por antigüedad para priorizar recuperación.", charts["ar_aging"], data["tables"]["ar_aging"])
    _docx_section(doc, "Tendencia de pago", data["narrative"]["payment_trend"], charts["payment_trend"], data["tables"]["payment_trend"])
    _docx_section(doc, "Facturación", data["narrative"]["billing"], charts["billing"], data["tables"]["top_billing"])
    _docx_section(doc, "Cuentas por pagar", data["narrative"]["payables"], charts["payables"], data["tables"]["top_payables_open"])
    _docx_section(doc, "Obligaciones pendientes del mes", "Obligaciones mensuales aceptadas en el preliminar antes de emitir el reporte. Esta tabla sirve como agenda operativa de pagos y seguimiento.", charts["cash_bridge"], data["tables"]["monthly_obligations"])
    _docx_section(doc, f"Cronograma y outlook - {data['period']['next_label']}", data["narrative"]["next_month_outlook"], charts["next_ar"], data["tables"]["top_next_receivables"])
    if charts["next_payables"] and os.path.exists(charts["next_payables"]):
        doc.add_picture(charts["next_payables"], width=Inches(6.4))
    _docx_section(doc, f"Comparativo {year - 1} vs {year}", data["narrative"]["year_comparison"], charts["accounting_mix"], _comparison_rows(data))
    _docx_section(doc, "Análisis de riesgo financiero", data["narrative"]["risk"])
    _docx_section(doc, "Conclusion reporte financiero", data["narrative"]["conclusion"])

    doc.save(path)
    return path, _safe_filename(label, "docx")


def _docx_executive_dashboard(doc, data):
    m = data["metrics"]
    a = data["accounting"]
    e = data["executive"]
    doc.add_heading("Executive Financial Dashboard", level=1)
    doc.add_paragraph(
        "Vista ejecutiva de liquidez, rentabilidad, capital de trabajo y riesgos principales del periodo. "
        "Los indicadores combinan facturacion, cobranza, cuentas por cobrar, cuentas por pagar y movimientos contables POSTED."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Indicador"
    table.rows[0].cells[1].text = "Resultado"
    table.rows[0].cells[2].text = "Lectura ejecutiva"
    rows = [
        ("Facturacion del mes", _money(m["revenue"]), "Capacidad comercial registrada en el periodo"),
        ("Cobranza del mes", _money(m["collections"]), f"Recuperacion equivalente a {_ratio_label(e['collection_ratio_pct'])} de la facturacion"),
        ("Resultado contable", _money(a["net_income"]), f"Margen neto contable {_ratio_label(e['net_margin_pct'])}"),
        ("Capital de trabajo", _money(a["working_capital"]), "Activos menos pasivos segun asientos POSTED"),
        ("Cartera abierta", _money(m["ar_open"]), f"Presion de cartera {_ratio_label(e['ar_pressure_pct'])} vs facturacion"),
        ("CxP abiertas", _money(m["payables_open"]), f"Cobertura por cobranza {_ratio_label(e['payable_coverage_pct'])}"),
        ("Outlook neto proximo mes", _money(m["next_net_outlook"]), "Cobros esperados menos pagos programados"),
    ]
    for indicator, result, reading in rows:
        cells = table.add_row().cells
        cells[0].text = indicator
        cells[1].text = result
        cells[2].text = reading
    doc.add_heading("Alertas ejecutivas", level=2)
    for alert in e["alerts"]:
        doc.add_paragraph(alert, style="List Bullet")
    doc.add_heading("Focos de decision", level=2)
    for item in e["decision_focus"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def _docx_section(doc, title, text, chart_path=None, rows=None):
    from docx.shared import Inches

    doc.add_heading(title, level=1)
    for paragraph in str(text or "").split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    if chart_path and os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6.4))
    if rows is not None:
        _docx_table(doc, rows)
    doc.add_page_break()


def _docx_table(doc, rows):
    has_obligation_fields = any((row or {}).get("concept") or (row or {}).get("issue_date") or (row or {}).get("due_date") for row in rows or [])
    table = doc.add_table(rows=1, cols=4 if has_obligation_fields else 2)
    table.style = "Table Grid"
    if has_obligation_fields:
        headers = ["Proveedor", "Concepto", "Fecha factura", "Monto"]
    else:
        headers = ["Detalle", "Monto"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    if rows:
        for row in rows:
            cells = table.add_row().cells
            if has_obligation_fields:
                cells[0].text = str(row.get("nombre_cliente") or row.get("name") or "N/A")
                cells[1].text = str(row.get("concept") or "Obligacion mensual")
                cells[2].text = str(row.get("issue_date") or row.get("due_date") or "")
                cells[3].text = f"{row.get('currency') or 'USD'} {_f(row.get('total') or row.get('amount')):,.2f}"
            else:
                cells[0].text = str(row.get("nombre_cliente") or row.get("name") or "N/A")
                cells[1].text = _money(row.get("total") or row.get("amount"))
    else:
        cells = table.add_row().cells
        cells[0].text = "No records for this period"
        cells[-1].text = "USD 0.00"

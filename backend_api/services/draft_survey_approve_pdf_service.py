import os
import shutil
import tempfile
import subprocess
from pathlib import Path
from docx import Document

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

from services.draft_survey_excel_service import DraftSurveyExcelGenerator
from services.pdf_merge_service import merge_pdf_list


TEMPLATE_WORD_PATH = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        "..",
        "templates",
        "draft_word_template.docx"
    )
)


class DraftSurveyApprovePdfService:

    # =========================================================
    # PUBLIC
    # =========================================================
    def generate_final_pdf(self, word_payload: dict, excel_payload: dict) -> str:

        word_pdf = self._generate_word_pdf_from_template(word_payload)
        excel_pdf = self._generate_excel_pdf_from_template(excel_payload)

        final_pdf = merge_pdf_list([word_pdf, excel_pdf])
        return final_pdf

    # =========================================================
    # WORD PIPELINE (FORMATO 100% RESPETADO)
    # =========================================================
    def _generate_word_pdf_from_template(self, payload: dict) -> str:

        if not os.path.exists(TEMPLATE_WORD_PATH):
            raise FileNotFoundError(f"Word template not found: {TEMPLATE_WORD_PATH}")

        out_dir = tempfile.mkdtemp(prefix="draft_word_")
        tmp_docx = os.path.join(out_dir, "draft_word_filled.docx")

        doc = Document(TEMPLATE_WORD_PATH)
        replacements = self._build_replacements(payload)

        self._replace_in_document(doc, replacements)

        for section in doc.sections:
            self._replace_in_container(section.header, replacements)
            self._replace_in_container(section.footer, replacements)

        doc.save(tmp_docx)

        tmp_pdf = os.path.join(out_dir, "draft_word_filled.pdf")
        self._convert_docx_to_pdf(tmp_docx, tmp_pdf)

        if not os.path.exists(tmp_pdf):
            raise RuntimeError("Word PDF was not created")

        return tmp_pdf

    def _build_replacements(self, payload: dict) -> dict:
        repl = {}
        for k, v in (payload or {}).items():
            repl["{" + str(k) + "}"] = "" if v is None else str(v)
        return repl

    def _replace_in_document(self, doc: Document, repl: dict):
        for p in doc.paragraphs:
            self._replace_in_paragraph_safe(p, repl)

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph_safe(p, repl)

    def _replace_in_container(self, container, repl: dict):
        for p in container.paragraphs:
            self._replace_in_paragraph_safe(p, repl)

        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph_safe(p, repl)

    def _replace_in_paragraph_safe(self, paragraph, repl: dict):
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        if "{" not in full_text:
            return

        changed = False
        for ph, val in repl.items():
            if ph in full_text:
                full_text = full_text.replace(ph, val)
                changed = True

        if not changed:
            return

        idx = 0
        for run in paragraph.runs:
            length = len(run.text)
            run.text = full_text[idx: idx + length]
            idx += length

        if idx < len(full_text):
            paragraph.runs[-1].text += full_text[idx:]

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str):

        if docx2pdf_convert:
            out_dir = os.path.dirname(pdf_path)
            os.makedirs(out_dir, exist_ok=True)
            docx2pdf_convert(docx_path, out_dir)

            generated = Path(docx_path).with_suffix(".pdf").name
            gen_path = os.path.join(out_dir, generated)

            if os.path.exists(gen_path) and gen_path != pdf_path:
                shutil.move(gen_path, pdf_path)
            return

        out_dir = os.path.dirname(pdf_path)
        os.makedirs(out_dir, exist_ok=True)

        cmd = [
            "soffice",
            "--headless",
            "--nologo",
            "--nolockcheck",
            "--convert-to", "pdf",
            "--outdir", out_dir,
            docx_path
        ]

        r = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

        if r.returncode != 0:
            raise RuntimeError(f"DOCX->PDF conversion failed:\n{r.stderr}")

    # =========================================================
    # EXCEL -> PDF (WINDOWS EXCEL COM) — ULTRA BLINDADO
    # =========================================================
    def _generate_excel_pdf_from_template(self, payload: dict) -> str:
        import os
        import tempfile
        import time

        try:
            import pythoncom
            import win32com.client
        except Exception as e:
            raise RuntimeError(
                f"Microsoft Excel COM (pywin32) no está disponible: {e}"
            )

        # 1) Generar Excel temporal (tu servicio)
        generator = DraftSurveyExcelGenerator()
        excel_path = generator.generate(payload or {})

        if not excel_path or not isinstance(excel_path, str):
            raise RuntimeError("Excel path inválido (generator.generate devolvió vacío).")

        excel_path = os.path.abspath(excel_path)

        if not os.path.exists(excel_path):
            raise RuntimeError(f"Excel file was not generated: {excel_path}")

        # 2) Output PDF (en carpeta temporal aparte, garantizada)
        out_dir = tempfile.mkdtemp(prefix="draft_excel_pdf_")
        pdf_path = os.path.abspath(os.path.join(out_dir, "draft_survey.pdf"))

        # Helpers locales (blindaje total)
        def _wait_for_file(path: str, timeout_sec: int = 25) -> bool:
            start = time.time()
            while time.time() - start < timeout_sec:
                if os.path.exists(path):
                    try:
                        if os.path.getsize(path) > 0:
                            return True
                    except Exception:
                        pass
                time.sleep(0.2)
            return False

        pythoncom.CoInitialize()
        excel = None
        workbook = None

        # Constantes COM (sin depender de win32.constants)
        xlTypePDF = 0
        xlQualityStandard = 0

        try:
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False

            # Blindaje (evita prompts/eventos/links)
            try:
                excel.EnableEvents = False
            except Exception:
                pass

            try:
                excel.AskToUpdateLinks = False
            except Exception:
                pass

            # Abrir workbook
            workbook = excel.Workbooks.Open(
                excel_path,
                UpdateLinks=0,
                ReadOnly=False
            )

            # Recalcular (por consistencia de impresión)
            try:
                excel.Calculation = -4105  # xlCalculationAutomatic
            except Exception:
                pass

            try:
                excel.CalculateFullRebuild()
                # Espera corta a cálculo (sin loop infinito)
                t0 = time.time()
                while getattr(excel, "CalculationState", 0) != 0:
                    if time.time() - t0 > 10:
                        break
                    time.sleep(0.1)
            except Exception:
                pass

            # Export PDF (workbook completo)
            try:
                workbook.ExportAsFixedFormat(
                    Type=xlTypePDF,
                    Filename=pdf_path,
                    Quality=xlQualityStandard,
                    IncludeDocProperties=True,
                    IgnorePrintAreas=False,
                    OpenAfterPublish=False
                )
            except Exception as e:
                raise RuntimeError(f"Excel ExportAsFixedFormat falló: {e}")

        finally:
            # Cerrar workbook / excel SIEMPRE
            try:
                if workbook is not None:
                    workbook.Close(False)
            except Exception:
                pass

            try:
                if excel is not None:
                    excel.Quit()
            except Exception:
                pass

            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

        # 3) Validación final real (esperar que el PDF exista y tenga tamaño)
        if not _wait_for_file(pdf_path, timeout_sec=25):
            raise RuntimeError(f"Excel PDF was not created or is empty: {pdf_path}")

        return pdf_path
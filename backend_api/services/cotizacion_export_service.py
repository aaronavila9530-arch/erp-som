from datetime import date
from pathlib import Path

from docx import Document
try:
    from services.template_autofit import apply_docx_autofit
except ModuleNotFoundError:
    from backend_api.services.template_autofit import apply_docx_autofit
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


BACKEND_DIR = Path(__file__).resolve().parents[1]
REPO_DIR = BACKEND_DIR.parent
ASSET_DIRS = (
    BACKEND_DIR / "assets",
    REPO_DIR / "assets",
)


def _asset(name: str) -> str | None:
    for directory in ASSET_DIRS:
        path = directory / name
        if path.is_file():
            return str(path)
    return None


def _is_mci(data: dict) -> bool:
    company_code = str(data.get("company_code") or "").upper()
    company_name = str(data.get("company_name") or "").upper()
    return company_code == "MCI-CR" or "MARINE CLAIMS" in company_name


def export_cotizacion_word(data: dict, output_path: str):
    doc = Document()
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_img = _asset("header.png")
    if header_img:
        header.add_run().add_picture(header_img, width=Inches(2.5))

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT

    quotation_number = data.get("quotation_number") or ""
    if quotation_number:
        body.add_run(f"{quotation_number}\n").bold = True

    body.add_run(f"Alajuela, Costa Rica\n{date.today()}\n\n")

    idioma = data.get("idioma") or "ES"
    servicio = data.get("servicio") or ""
    subject = f"Quotation - {servicio}" if idioma == "EN" else f"Cotizacion - {servicio}"
    body.add_run(f"{'Subject' if idioma == 'EN' else 'Asunto'}: {subject}\n\n").bold = True
    body.add_run(data.get("texto") or "")

    doc.add_paragraph("\n")
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig.add_run("Sincerely,\n\n" if idioma == "EN" else "Atentamente,\n\n")

    signature = _asset("FIRMA DIANA.png")
    if signature:
        sig.add_run().add_picture(signature, width=Inches(1.8))

    if _is_mci(data):
        sig.add_run("\nMsc. Diana Quiros Benambourg\n")
        sig.add_run("Business Manager\n").bold = True
        sig.add_run("MSL 2.0\n").bold = True
        sig.add_run("MARINE CLAIMS & RISK INTELLIGENCE").bold = True
    else:
        sig.add_run("\nDiana Quiros Benambourg\n")
        sig.add_run("Business Manager\n")
        sig.add_run("Marine Surveyors & Logistics Group SRL")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = (
        "Head Office - Costa Rica, Alajuela, Plaza Aeropuerto G-14\n"
        "Phone (506) 8814-07-84 - (506) 4052-8382"
    )

    apply_docx_autofit(doc)
    doc.save(output_path)


def export_cotizacion_pdf(data: dict, output_path: str):
    c = canvas.Canvas(output_path, pagesize=LETTER)
    width, height = LETTER
    left = 0.75 * inch
    right = 0.75 * inch
    body_width = width - left - right
    header_y = height - 1.9 * inch
    top_y = header_y - 0.5 * inch
    footer_y = 0.7 * inch
    signature_block_height = 1.55 * inch
    body_min_y = footer_y + 0.35 * inch

    watermark = _asset("watermark.png")
    header = _asset("header.png")
    signature = _asset("FIRMA DIANA.png")

    def wrap_text(text: str, font_name: str, font_size: int, max_width: float) -> list[str]:
        text = str(text or "")
        if not text:
            return [""]
        wrapped = []
        for raw_line in text.splitlines() or [""]:
            line = raw_line.strip()
            if not line:
                wrapped.append("")
                continue
            words = line.split(" ")
            current = ""
            for word in words:
                candidate = word if not current else f"{current} {word}"
                if stringWidth(candidate, font_name, font_size) <= max_width:
                    current = candidate
                    continue
                if current:
                    wrapped.append(current)
                    current = word
                else:
                    piece = ""
                    for ch in word:
                        candidate_piece = piece + ch
                        if stringWidth(candidate_piece, font_name, font_size) <= max_width:
                            piece = candidate_piece
                        else:
                            if piece:
                                wrapped.append(piece)
                            piece = ch
                    current = piece
            wrapped.append(current)
        return wrapped

    def draw_footer():
        c.setFont("Helvetica", 8)
        c.drawCentredString(
            width / 2,
            footer_y,
            "Head Office - Costa Rica, Alajuela, Plaza Aeropuerto G-14 - "
            "Phone (506) 8814-07-84 - (506) 4052-8382",
        )

    def draw_static():
        if watermark:
            c.saveState()
            c.setFillAlpha(0.08)
            c.drawImage(watermark, 1.2 * inch, 2.5 * inch, width=4.5 * inch, preserveAspectRatio=True, mask="auto")
            c.restoreState()
        if header:
            c.drawImage(header, left, header_y, width=3.6 * inch, preserveAspectRatio=True, mask="auto")
        draw_footer()

    def new_page():
        c.showPage()
        draw_static()
        c.setFont("Helvetica", 10)
        return top_y

    def draw_wrapped_line(text: str, font_name: str, font_size: int, y_pos: float, min_y: float) -> float:
        c.setFont(font_name, font_size)
        for wrapped_line in wrap_text(text, font_name, font_size, body_width):
            if y_pos < min_y:
                y_pos = new_page()
                c.setFont(font_name, font_size)
            c.drawString(left, y_pos, wrapped_line)
            y_pos -= 14 if font_size <= 10 else 16
        return y_pos

    draw_static()
    y = top_y

    quotation_number = data.get("quotation_number") or ""
    if quotation_number:
        y = draw_wrapped_line(quotation_number, "Helvetica-Bold", 11, y, body_min_y)
        y -= 8

    y = draw_wrapped_line("Alajuela, Costa Rica", "Helvetica-Bold", 10, y, body_min_y)
    y = draw_wrapped_line(str(date.today()), "Helvetica-Bold", 10, y, body_min_y)
    y -= 12

    idioma = data.get("idioma") or "ES"
    servicio = data.get("servicio") or ""
    subject = f"Quotation - {servicio}" if idioma == "EN" else f"Cotizacion - {servicio}"
    y = draw_wrapped_line(subject, "Helvetica-Bold", 11, y, body_min_y)
    y -= 8

    for line in (data.get("texto") or "").split("\n"):
        y = draw_wrapped_line(line, "Helvetica", 10, y, body_min_y)

    title_y = footer_y + 0.55 * inch
    name_y = title_y + 14
    signature_y = name_y + 18

    if y < footer_y + signature_block_height:
        y = new_page()

    if signature:
        c.drawImage(signature, left, signature_y, width=1.8 * inch, height=0.6 * inch, preserveAspectRatio=True, mask="auto")

    if _is_mci(data):
        c.setFont("Helvetica", 10)
        c.drawString(left, name_y, "Msc. Diana Quiros Benambourg")
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left, title_y, "Business Manager")
        c.drawString(left, title_y - 12, "MSL 2.0")
        c.drawString(left, title_y - 24, "MARINE CLAIMS & RISK INTELLIGENCE")
    else:
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left, name_y, "Diana Quiros Benambourg")
        c.setFont("Helvetica", 10)
        c.drawString(left, title_y, "Business Manager")
        c.drawString(left, title_y - 12, "MSL MARINE SURVEYORS & LOGISTICS GROUP SRL")

    c.save()

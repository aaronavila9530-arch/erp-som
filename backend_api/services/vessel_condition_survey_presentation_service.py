import os
import tempfile
import subprocess
from datetime import datetime, date
from docx import Document


class VesselConditionSurveyPresentationService:

    # =========================================================
    # MAIN
    # =========================================================
    def generate_pdf_by_id(self, conn, record_id: int) -> str:

        data = self._get_data(conn, record_id)

        if not data:
            raise ValueError("Record not found")

        template_name = self._select_template(data)

        template_path = self._get_template_path(template_name)

        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f"Template not found: {template_path}"
            )

        doc = Document(template_path)

        data = self._format_dates(data)

        self._replace_everywhere(doc, data)

        temp_dir = tempfile.mkdtemp()

        docx_path = os.path.join(
            temp_dir,
            f"condition_survey_presentation_{record_id}.docx"
        )

        pdf_path = os.path.join(
            temp_dir,
            f"condition_survey_presentation_{record_id}.pdf"
        )

        doc.save(docx_path)

        self._convert_to_pdf(docx_path, temp_dir)

        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF conversion failed")

        return pdf_path


    # =========================================================
    # TEMPLATE SELECTOR
    # =========================================================
    def _select_template(self, data: dict):

        report_type = (data.get("report_type") or "").lower()

        # -----------------------------------------------------
        # P&I SURVEY
        # -----------------------------------------------------

        if "p&i vessel condition survey" in report_type:

            cargo_type = (data.get("cargo_type") or "").lower()

            if "acero" in cargo_type or "steel" in cargo_type:
                return "presentation_p&i_aceros.docx"

            return "presentation_p&i_general.docx"

        # -----------------------------------------------------
        # MOORING LINES
        # -----------------------------------------------------

        if "mooring lines" in report_type:
            return "presentation_mooring_lines.docx"

        # -----------------------------------------------------
        # HULL CONDITION
        # -----------------------------------------------------

        if "hull condition" in report_type:
            return "presentation_hull_inspection.docx"

        # -----------------------------------------------------
        # CARGO HOLDS
        # -----------------------------------------------------

        if "cargo holds condition" in report_type:
            return "presentation_hold_condition.docx"

        raise ValueError(
            f"No template configured for report_type: {report_type}"
        )


    # =========================================================
    # TEMPLATE PATH
    # =========================================================
    def _get_template_path(self, template_name: str):

        base_dir = os.path.dirname(os.path.abspath(__file__))

        return os.path.abspath(
            os.path.join(
                base_dir,
                "..",
                "templates",
                template_name
            )
        )


    # =========================================================
    # GET DATA
    # =========================================================
    def _get_data(self, conn, record_id):

        cur = conn.cursor()

        cur.execute(
            """
            SELECT *
            FROM vessel_condition_surveys
            WHERE id = %s
            """,
            (record_id,)
        )

        row = cur.fetchone()

        if not row:
            return None

        columns = [desc[0] for desc in cur.description]

        return dict(zip(columns, row))


    # =========================================================
    # DATE FORMAT
    # =========================================================
    def _format_dates(self, data):

        formatted = {}

        for k, v in data.items():

            if isinstance(v, (datetime, date)):
                formatted[k] = v.strftime("%B %d %Y")
            else:
                formatted[k] = v

        return formatted


    # =========================================================
    # REPLACE PLACEHOLDERS
    # =========================================================
    def _replace_everywhere(self, doc, data):

        for p in doc.paragraphs:
            self._replace_paragraph(p, data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_paragraph(p, data)

        for section in doc.sections:

            for p in section.header.paragraphs:
                self._replace_paragraph(p, data)

            for p in section.footer.paragraphs:
                self._replace_paragraph(p, data)


    # =========================================================
    # SAFE PARAGRAPH REPLACE
    # =========================================================
    def _replace_paragraph(self, paragraph, data):

        full_text = "".join(run.text for run in paragraph.runs)

        replaced = full_text

        for key, value in data.items():

            placeholder = "{" + key + "}"

            replaced = replaced.replace(
                placeholder,
                "" if value is None else str(value)
            )

        if replaced != full_text:

            paragraph.runs[0].text = replaced

            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ""


    # =========================================================
    # DOCX → PDF
    # =========================================================
    def _convert_to_pdf(self, docx_path: str, output_dir: str):

        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path
            ],
            check=True
        )
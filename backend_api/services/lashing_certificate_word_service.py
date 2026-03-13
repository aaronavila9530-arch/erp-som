from pathlib import Path
from docx import Document
from io import BytesIO
from fastapi.responses import StreamingResponse
import subprocess
import tempfile
import os


class LashingCertificateWordService:

    def __init__(self):

        # Ruta relativa portable
        self.template_path = (
            Path(__file__).resolve()
            .parents[1]
            / "templates"
            / "lashing_certificate_template.docx"
        )

    # =========================================================
    # REPLACE PLACEHOLDERS
    # =========================================================

    def _replace_placeholders(self, doc, data):

        for paragraph in doc.paragraphs:

            for key, value in data.items():

                placeholder = f"{{{key}}}"

                if placeholder in paragraph.text:

                    paragraph.text = paragraph.text.replace(
                        placeholder,
                        "" if value is None else str(value)
                    )

        # TABLAS
        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for key, value in data.items():

                        placeholder = f"{{{key}}}"

                        if placeholder in cell.text:

                            cell.text = cell.text.replace(
                                placeholder,
                                "" if value is None else str(value)
                            )

    # =========================================================
    # GENERATE WORD
    # =========================================================

    def generate_word(self, data):

        doc = Document(self.template_path)

        self._replace_placeholders(doc, data)

        buffer = BytesIO()

        doc.save(buffer)

        buffer.seek(0)

        filename = f"lashing_certificate_{data.get('report_no','report')}.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )

    # =========================================================
    # GENERATE PDF
    # =========================================================

    def generate_pdf(self, data):

        with tempfile.TemporaryDirectory() as tmp:

            word_path = os.path.join(tmp, "temp.docx")
            pdf_path = os.path.join(tmp, "temp.pdf")

            doc = Document(self.template_path)

            self._replace_placeholders(doc, data)

            doc.save(word_path)

            subprocess.run([
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                tmp,
                word_path
            ])

            with open(pdf_path, "rb") as f:

                pdf_bytes = f.read()

        buffer = BytesIO(pdf_bytes)

        filename = f"lashing_certificate_{data.get('report_no','report')}.pdf"

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
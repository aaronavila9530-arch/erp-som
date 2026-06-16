import os
import tempfile
import subprocess
from typing import Dict
from datetime import datetime
from docx import Document
try:
    from services.template_autofit import apply_docx_autofit
except ModuleNotFoundError:
    from backend_api.services.template_autofit import apply_docx_autofit


# =====================================================
# TEMPLATE PATH
# =====================================================
TEMPLATE_PATH = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        "..",
        "templates",
        "presentation_bunker.docx"
    )
)


# =====================================================
# INTERNAL — SAFE REPLACE (CROSS-RUN SAFE)
# (NO ROMPE FORMATO NI COLORES)
# =====================================================
def _replace_in_paragraph(paragraph, placeholder: str, value: str):

    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    if placeholder not in full_text:
        return

    start = full_text.index(placeholder)
    end = start + len(placeholder)

    current_pos = 0
    first_replacement_done = False

    for run in paragraph.runs:

        run_text = run.text
        run_len = len(run_text)

        run_start = current_pos
        run_end = current_pos + run_len

        if run_end > start and run_start < end:

            prefix_len = max(0, start - run_start)
            suffix_len = max(0, run_end - end)

            prefix = run_text[:prefix_len]
            suffix = run_text[run_len - suffix_len:] if suffix_len > 0 else ""

            if not first_replacement_done:
                run.text = prefix + value + suffix
                first_replacement_done = True
            else:
                run.text = ""

        current_pos += run_len


def _replace_in_paragraphs(paragraphs, placeholders: Dict[str, str]):
    for p in paragraphs:
        for key, value in placeholders.items():
            _replace_in_paragraph(p, key, value)


def _replace_in_tables(tables, placeholders: Dict[str, str]):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, placeholders)
                if cell.tables:
                    _replace_in_tables(cell.tables, placeholders)


# =====================================================
# MAIN — GENERATE BUNKER PRESENTATION PDF
# =====================================================
def generate_bunker_presentation_pdf(data: dict) -> str:

    # -------------------------------------------------
    # VALIDATIONS
    # -------------------------------------------------
    if not isinstance(data, dict):
        raise ValueError("Invalid data payload. Expected dict.")

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Bunker presentation template not found: {TEMPLATE_PATH}"
        )

    # -------------------------------------------------
    # DATE FORMATTER → January 13 2026
    # -------------------------------------------------
    def format_long_date(value):
        if not value:
            return ""

        try:
            s = str(value).split(" ")[0]
            dt = datetime.strptime(s, "%Y-%m-%d")
            return dt.strftime("%B %d %Y")
        except Exception:
            return str(value)

    # -------------------------------------------------
    # NORMAL DATE (fallback raw yyyy-mm-dd)
    # -------------------------------------------------
    def norm_date(v):
        if not v:
            return ""
        s = str(v)
        return s.split(" ")[0] if " " in s else s

    # -------------------------------------------------
    # CERTIFICATE FORMAT (replace _ with space + uppercase)
    # -------------------------------------------------
    def format_certificate(v):
        if not v:
            return ""
        return str(v).replace("_", " ").upper()

    # -------------------------------------------------
    # SAFE VALUE
    # -------------------------------------------------
    def safe_str(v):
        return str(v) if v is not None else ""

    # -------------------------------------------------
    # PLACEHOLDERS MAP
    # -------------------------------------------------
    placeholders = {

        # Identificadores principales
        "{bunker_cert_no}": safe_str(data.get("bunker_cert_no")),
        "{ship_name}": safe_str(data.get("ship_name")),

        # Fecha principal en formato largo
        "{report_date}": format_long_date(data.get("report_date")),

        # Certificado con _ reemplazado
        "{certificate}": format_certificate(data.get("certificate")),

        # Datos técnicos
        "{gross_tonnage}": safe_str(data.get("gross_tonnage")),
        "{port_of_registry}": safe_str(data.get("port_of_registry")),
        "{report_category}": safe_str(data.get("report_category")),

        # Cliente / ubicación
        "{client}": safe_str(data.get("client")),
        "{port}": safe_str(data.get("port")),
        "{country}": safe_str(data.get("country")),

        # Fechas operativas en formato largo
        "{berthing_date}": format_long_date(data.get("berthing_date")),
        "{commenced_date}": format_long_date(data.get("commenced_date")),

        # DSLOP
        "{dslop_date}": format_long_date(data.get("dslop_date")),
        "{dslop_port}": safe_str(data.get("dslop_port")),
        "{dslop_country}": safe_str(data.get("dslop_country")),
    }

    # -------------------------------------------------
    # LOAD TEMPLATE
    # -------------------------------------------------
    doc = Document(TEMPLATE_PATH)

    # BODY
    _replace_in_paragraphs(doc.paragraphs, placeholders)
    _replace_in_tables(doc.tables, placeholders)

    # HEADERS / FOOTERS
    for section in doc.sections:
        _replace_in_paragraphs(section.header.paragraphs, placeholders)
        _replace_in_tables(section.header.tables, placeholders)

        _replace_in_paragraphs(section.footer.paragraphs, placeholders)
        _replace_in_tables(section.footer.tables, placeholders)

    # -------------------------------------------------
    # SAVE TEMP DOCX
    # -------------------------------------------------
    fd, temp_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    apply_docx_autofit(doc)
    doc.save(temp_docx)

    output_dir = tempfile.mkdtemp(prefix="bunker_presentation_")

    # -------------------------------------------------
    # CONVERT USING LIBREOFFICE (HEADLESS)
    # -------------------------------------------------
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                temp_docx
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
            timeout=60
        )

    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice conversion timed out.")

    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice (soffice) not found. Ensure it is installed and in PATH."
        )

    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"Error converting DOCX to PDF: "
            f"{e.stderr.decode(errors='ignore')}"
        )

    # -------------------------------------------------
    # VALIDATE OUTPUT
    # -------------------------------------------------
    pdf_path = os.path.join(
        output_dir,
        os.path.splitext(os.path.basename(temp_docx))[0] + ".pdf"
    )

    if not os.path.exists(pdf_path):
        raise RuntimeError("Presentation PDF generation failed — output file not found.")

    return pdf_path

from copy import copy
import math

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.cell_range import CellRange


def _set_cell_margins(cell, margin_twips=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def _allow_cell_wrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for no_wrap in list(tc_pr.findall(qn("w:noWrap"))):
        tc_pr.remove(no_wrap)


def _set_table_autofit(table):
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    for layout in list(tbl_pr.findall(qn("w:tblLayout"))):
        tbl_pr.remove(layout)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)


def _fit_paragraph_runs(paragraph, in_table_cell=False):
    text = (paragraph.text or "").strip()
    if not text:
        return

    if not in_table_cell:
        paragraph.paragraph_format.keep_together = True
        return

    length = len(text)
    if length <= 45:
        return

    if length > 180:
        target_size = 7
    elif length > 120:
        target_size = 8
    elif length > 80:
        target_size = 9
    else:
        target_size = 10

    for run in paragraph.runs:
        if not run.text:
            continue

        current_size = run.font.size.pt if run.font.size else None
        if current_size is None or current_size > target_size:
            run.font.size = Pt(target_size)


def apply_docx_autofit(doc):
    """
    Apply layout-safe text fitting to generated Word files.

    Word cannot truly auto-fit text in every shape/textbox after python-docx
    replacement. This keeps existing template styling, allows table cells to
    wrap, adds small internal margins when missing, and only shrinks long text
    inside table cells so values do not get clipped.
    """
    for paragraph in doc.paragraphs:
        _fit_paragraph_runs(paragraph, in_table_cell=False)

    for table in doc.tables:
        _set_table_autofit(table)
        for row in table.rows:
            for cell in row.cells:
                _set_cell_margins(cell)
                _allow_cell_wrap(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    _fit_paragraph_runs(paragraph, in_table_cell=True)

    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _fit_paragraph_runs(paragraph, in_table_cell=False)
            for table in part.tables:
                _set_table_autofit(table)
                for row in table.rows:
                    for cell in row.cells:
                        _set_cell_margins(cell)
                        _allow_cell_wrap(cell)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            _fit_paragraph_runs(paragraph, in_table_cell=True)


def _merged_range_for_cell(ws, row, col):
    coord = f"{get_column_letter(col)}{row}"
    for merged in ws.merged_cells.ranges:
        if coord in merged:
            return merged
    return None


def _range_width(ws, cell_range):
    if isinstance(cell_range, CellRange):
        columns = range(cell_range.min_col, cell_range.max_col + 1)
    else:
        columns = [cell_range.column]

    width = 0
    for col in columns:
        letter = get_column_letter(col)
        width += ws.column_dimensions[letter].width or 8.43
    return max(width, 8.43)


def _estimated_lines(value, width):
    text = str(value or "")
    if not text:
        return 1

    chars_per_line = max(int(width * 1.15), 8)
    total = 0
    for line in text.splitlines() or [""]:
        total += max(1, math.ceil(len(line) / chars_per_line))
    return max(total, 1)


def apply_excel_autofit(ws, max_column_width=48, min_row_height=15, max_row_height=120):
    """
    Apply pragmatic fitting to an Excel worksheet.

    Excel calculates perfect AutoFit only in the desktop app. For generated
    files we enable wrap/shrink-to-fit, expand row heights heuristically, and
    widen non-merged text columns within a conservative cap.
    """
    row_heights = {}
    column_widths = {}

    for row in ws.iter_rows():
        for cell in row:
            if cell.value in (None, ""):
                continue

            alignment = copy(cell.alignment)
            alignment.wrap_text = True
            alignment.shrink_to_fit = True
            alignment.vertical = alignment.vertical or "center"
            cell.alignment = alignment

            merged = _merged_range_for_cell(ws, cell.row, cell.column)
            effective_range = merged or cell
            width = _range_width(ws, effective_range)
            lines = _estimated_lines(cell.value, width)
            wanted_height = min(max_row_height, max(min_row_height, lines * 15))
            row_heights[cell.row] = max(row_heights.get(cell.row, min_row_height), wanted_height)

            if merged is None and isinstance(cell.value, str):
                wanted_width = min(max_column_width, max(10, len(cell.value) + 2))
                column_widths[cell.column] = max(
                    column_widths.get(cell.column, 0),
                    wanted_width
                )

    for row_idx, height in row_heights.items():
        current = ws.row_dimensions[row_idx].height or min_row_height
        ws.row_dimensions[row_idx].height = max(current, height)

    for col_idx, width in column_widths.items():
        letter = get_column_letter(col_idx)
        current = ws.column_dimensions[letter].width or 8.43
        ws.column_dimensions[letter].width = min(max_column_width, max(current, width))


def apply_workbook_autofit(wb):
    for ws in wb.worksheets:
        apply_excel_autofit(ws)

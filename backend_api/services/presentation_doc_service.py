import os
import tempfile
import subprocess
from docx import Document


# =====================================================
# TEMPLATE PATH (ABSOLUTO Y REAL)
# =====================================================
TEMPLATE_PATH = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        "..",
        "templates",
        "presentation_containers.docx"
    )
)


# =====================================================
# INTERNAL — SAFE REPLACE (PRESERVA ESTILOS)
# =====================================================
def _replace_in_paragraphs(paragraphs, placeholders: dict):
    for p in paragraphs:
        if not p.runs:
            continue

        # texto completo real (Word puede partirlo en runs)
        full_text = "".join(run.text for run in p.runs)

        replaced = False
        for key, value in placeholders.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
                replaced = True

        if replaced:
            # conservar estilo del primer run
            base_run = p.runs[0]
            base_style = {
                "bold": base_run.bold,
                "italic": base_run.italic,
                "underline": base_run.underline,
                "font_name": base_run.font.name,
                "font_size": base_run.font.size,
                "font_color": base_run.font.color.rgb if base_run.font.color else None,
            }

            # limpiar runs
            for run in p.runs:
                run.text = ""

            # crear nuevo run con el texto final
            new_run = p.add_run(full_text)

            # restaurar estilo
            new_run.bold = base_style["bold"]
            new_run.italic = base_style["italic"]
            new_run.underline = base_style["underline"]
            new_run.font.name = base_style["font_name"]
            new_run.font.size = base_style["font_size"]
            if base_style["font_color"]:
                new_run.font.color.rgb = base_style["font_color"]


# =====================================================
# INTERNAL — SAFE REPLACE IN TABLES
# =====================================================
def _replace_in_tables(tables, placeholders: dict):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, placeholders)
                if cell.tables:
                    _replace_in_tables(cell.tables, placeholders)


# =====================================================
# GENERATE PRESENTATION PDF
# =====================================================
def generate_presentation_pdf(data: dict) -> str:
    """
    Genera PDF desde presentation_containers.docx
    Reemplaza placeholders respetando:
    - colores
    - negrita
    - tamaño
    - alineación
    """

    # -------------------------------------------------
    # VALIDATIONS
    # -------------------------------------------------
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Presentation template not found: {TEMPLATE_PATH}"
        )

    if not isinstance(data, dict):
        raise ValueError("Invalid data payload — expected dict")

    # -------------------------------------------------
    # LOAD TEMPLATE
    # -------------------------------------------------
    doc = Document(TEMPLATE_PATH)

    placeholders = {
        "{{CERT_NO}}": str(data.get("cert_no") or ""),
        "{{CONTAINER}}": str(data.get("container") or ""),
        "{{TO}}": str(data.get("to") or ""),
        "{{PLACE}}": str(data.get("place") or ""),
        "{{DATE}}": str(data.get("date") or "")
    }

    # -------------------------------------------------
    # BODY
    # -------------------------------------------------
    _replace_in_paragraphs(doc.paragraphs, placeholders)
    _replace_in_tables(doc.tables, placeholders)

    # -------------------------------------------------
    # HEADERS / FOOTERS
    # -------------------------------------------------
    for section in doc.sections:
        _replace_in_paragraphs(section.header.paragraphs, placeholders)
        _replace_in_tables(section.header.tables, placeholders)

        _replace_in_paragraphs(section.footer.paragraphs, placeholders)
        _replace_in_tables(section.footer.tables, placeholders)

    # -------------------------------------------------
    # SAVE TEMP DOCX
    # -------------------------------------------------
    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(docx_path)

    output_dir = tempfile.mkdtemp()

    # -------------------------------------------------
    # CONVERT TO PDF (LibreOffice Headless)
    # -------------------------------------------------
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE
        )

    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice (soffice) is not installed or not available in PATH"
        )

    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"Error converting DOCX to PDF: {e.stderr.decode(errors='ignore')}"
        )

    # -------------------------------------------------
    # VALIDATE OUTPUT
    # -------------------------------------------------
    pdf_name = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    pdf_path = os.path.join(output_dir, pdf_name)

    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF generation failed — output file not found")

    return pdf_path

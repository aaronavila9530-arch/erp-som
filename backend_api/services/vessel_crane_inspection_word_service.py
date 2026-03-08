import os
import tempfile
from datetime import datetime, date
from docx import Document


class VesselCraneInspectionWordService:

    TEMPLATE_NAME = "crane_inspection_template.docx"

    # =========================================================
    # MAIN
    # =========================================================

    def generate_word_by_id(self, conn, record_id: int) -> str:

        data = self._get_data(conn, record_id)

        if not data:
            raise ValueError("Record not found")

        template_path = self._get_template_path()

        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f"Template not found: {template_path}"
            )

        doc = Document(template_path)

        # remove only NULL bullets
        self._remove_null_paragraphs(doc, data)

        # replace placeholders everywhere
        self._replace_placeholders_everywhere(doc, data)

        output_path = self._build_output_path(data)

        doc.save(output_path)

        return output_path


    # =========================================================
    # GET DATA
    # =========================================================

    def _get_data(self, conn, record_id):

        cur = conn.cursor()

        cur.execute(
            """
            SELECT *
            FROM vessel_crane_inspection_reports
            WHERE id = %s
            """,
            (record_id,)
        )

        row = cur.fetchone()

        if not row:
            return None

        columns = [desc[0] for desc in cur.description]

        data = dict(zip(columns, row))

        return data


    # =========================================================
    # TEMPLATE PATH
    # =========================================================

    def _get_template_path(self):

        base_dir = os.path.dirname(os.path.abspath(__file__))

        return os.path.abspath(
            os.path.join(
                base_dir,
                "..",
                "templates",
                self.TEMPLATE_NAME
            )
        )


    # =========================================================
    # SAFE VALUE (DATE LONG ENGLISH)
    # =========================================================

    def _safe(self, value):

        if value is None:
            return ""

        if isinstance(value, (datetime, date)):
            return value.strftime("%B %d %Y")

        if isinstance(value, str):

            value = value.strip()

            try:
                dt = datetime.strptime(value[:10], "%Y-%m-%d")
                return dt.strftime("%B %d %Y")
            except Exception:
                pass

        return str(value)


    # =========================================================
    # PLACEHOLDER GENERATOR
    # =========================================================

    def _placeholders(self, key):

        return [
            f"{{{key}}}",
            f"{{{{{key}}}}}"
        ]


    # =========================================================
    # REPLACE PLACEHOLDERS SAFELY (RUN SAFE)
    # =========================================================

    def _replace_placeholders_everywhere(self, doc: Document, data: dict):

        def replace_in_paragraph(paragraph):

            if not paragraph.runs:
                return

            full_text = "".join(run.text for run in paragraph.runs)

            for key, value in data.items():

                for placeholder in self._placeholders(key):

                    if placeholder in full_text:
                        full_text = full_text.replace(
                            placeholder,
                            self._safe(value)
                        )

            index = 0

            for run in paragraph.runs:

                length = len(run.text)

                if length == 0:
                    continue

                run.text = full_text[index:index + length]
                index += length

            if index < len(full_text):
                paragraph.runs[-1].text += full_text[index:]

        # BODY
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        # TABLES
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        # HEADERS / FOOTERS
        for section in doc.sections:

            header = section.header

            for paragraph in header.paragraphs:
                replace_in_paragraph(paragraph)

            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)

            footer = section.footer

            for paragraph in footer.paragraphs:
                replace_in_paragraph(paragraph)

            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)


    # =========================================================
    # REMOVE NULL BULLETS (AUTO-DETECT)
    # =========================================================

    def _remove_null_paragraphs(self, doc: Document, data: dict):

        bullet_keys = [
            key
            for key in data.keys()
            if any(
                key.startswith(prefix)
                for prefix in (
                    "crane",
                    "recommendation",
                    "grabs_condition",
                    "conclusion"
                )
            )
        ]

        null_placeholders = set()

        for key in bullet_keys:

            if data.get(key) in [None, "", "NULL", "null"]:

                for placeholder in self._placeholders(key):
                    null_placeholders.add(placeholder)

        for paragraph in list(doc.paragraphs):
            self._remove_if_contains_null(paragraph, null_placeholders)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in list(cell.paragraphs):
                        self._remove_if_contains_null(paragraph, null_placeholders)


    def _remove_if_contains_null(self, paragraph, null_placeholders):

        if not paragraph.text:
            return

        for placeholder in null_placeholders:

            if placeholder in paragraph.text:

                self._delete_paragraph(paragraph)
                return


    def _delete_paragraph(self, paragraph):

        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None


    # =========================================================
    # OUTPUT PATH
    # =========================================================

    def _build_output_path(self, data: dict):

        report_number = data.get("report_number") or "CRANE_INSPECTION"

        safe_name = report_number.replace("/", "_").replace("\\", "_")

        temp_dir = tempfile.gettempdir()

        return os.path.join(
            temp_dir,
            f"{safe_name}_CRANE_INSPECTION.docx"
        )
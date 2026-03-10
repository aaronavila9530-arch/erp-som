import os
import tempfile
import subprocess
from datetime import datetime, date
from docx import Document


class VesselCraneInspectionPresentationWordService:

    TEMPLATE_NAME = "presentation_crane_inspection.docx"

    # =========================================================
    # MAIN
    # =========================================================
    def generate_pdf_by_id(self, conn, record_id: int) -> str:

        data = self._get_data(conn, record_id)

        if not data:
            raise ValueError("Record not found")

        template_path = self._get_template_path()

        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f"Template not found: {template_path}"
            )

        doc = Document(template_path)

        # -----------------------------------------------------
        # FORMAT DATA (DATES → LONG ENGLISH)
        # -----------------------------------------------------

        data = self._format_dates(data)

        # -----------------------------------------------------
        # REPLACE PLACEHOLDERS
        # -----------------------------------------------------

        self._replace_all(doc, data)

        # -----------------------------------------------------
        # TEMP FILES
        # -----------------------------------------------------

        temp_dir = tempfile.mkdtemp()

        docx_path = os.path.join(
            temp_dir,
            f"crane_inspection_presentation_{record_id}.docx"
        )

        pdf_path = os.path.join(
            temp_dir,
            f"crane_inspection_presentation_{record_id}.pdf"
        )

        doc.save(docx_path)

        # -----------------------------------------------------
        # CONVERT DOCX → PDF
        # -----------------------------------------------------

        self._convert_to_pdf(docx_path, temp_dir)

        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF conversion failed")

        return pdf_path

    # =========================================================
    # CONVERT DOCX → PDF
    # =========================================================
    def _convert_to_pdf(self, docx_path: str, output_dir: str):

        try:

            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    output_dir,
                    docx_path
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )

        except Exception as e:

            raise RuntimeError(
                f"LibreOffice conversion failed: {e}"
            )

    # =========================================================
    # TEMPLATE PATH
    # =========================================================
    def _get_template_path(self):

        base_dir = os.path.dirname(os.path.abspath(__file__))

        return os.path.abspath(
            os.path.join(
                base_dir,
                "..",
                "templates",
                self.TEMPLATE_NAME
            )
        )

    # =========================================================
    # GET DATA
    # =========================================================
    def _get_data(self, conn, record_id):

        cur = conn.cursor()

        cur.execute(
            """
            SELECT *
            FROM vessel_crane_inspection_reports
            WHERE id = %s
            """,
            (record_id,)
        )

        row = cur.fetchone()

        if not row:
            return None

        columns = [desc[0] for desc in cur.description]

        return dict(zip(columns, row))

    # =========================================================
    # FORMAT DATES → LONG ENGLISH
    # =========================================================
    def _format_dates(self, data):

        formatted = {}

        for k, v in data.items():

            if isinstance(v, (datetime, date)):
                formatted[k] = v.strftime("%B %d %Y")
            else:
                formatted[k] = v

        return formatted

    # =========================================================
    # REPLACE PLACEHOLDERS
    # =========================================================
    def _replace_all(self, doc, data):

        for p in doc.paragraphs:
            self._replace_runs(p, data)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_runs(p, data)

        for section in doc.sections:
            for p in section.header.paragraphs:
                self._replace_runs(p, data)
            for p in section.footer.paragraphs:
                self._replace_runs(p, data)

    # =========================================================
    # SAFE RUN REPLACEMENT
    # =========================================================
    def _replace_runs(self, paragraph, data):

        if not paragraph.runs:
            return

        for run in paragraph.runs:

            text = run.text

            for key, value in data.items():

                placeholder = "{{" + key + "}}"

                if placeholder in text:

                    text = text.replace(
                        placeholder,
                        "" if value is None else str(value)
                    )

            run.text = text
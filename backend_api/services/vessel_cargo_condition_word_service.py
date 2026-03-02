import os
import tempfile
from datetime import datetime, date
from docx import Document


class VesselCargoConditionWordService:

    TEMPLATE_NAME = "cargo_condition_template.docx"

    # =========================================================
    # MAIN
    # =========================================================
    def generate_word_by_id(self, conn, record_id: int) -> str:

        data = self._get_data(conn, record_id)

        if not data:
            raise ValueError("Record not found")

        template_path = self._get_template_path()

        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f"Template not found: {template_path}"
            )

        doc = Document(template_path)

        # 1️⃣ REMOVE NULL BULLET LINES FIRST
        self._remove_null_paragraphs(doc, data)

        # 2️⃣ REPLACE PLACEHOLDERS (BODY + TABLES + HEADER + FOOTER)
        self._replace_placeholders_everywhere(doc, data)

        output_path = self._build_output_path(data)

        doc.save(output_path)

        return output_path

    # =========================================================
    # FETCH DATA
    # =========================================================
    def _get_data(self, conn, record_id: int):

        cur = conn.cursor()

        cur.execute("""
            SELECT *
            FROM vessel_cargo_condition_surveys
            WHERE id = %s
        """, (record_id,))

        row = cur.fetchone()

        if not row:
            return None

        columns = [desc[0] for desc in cur.description]

        return dict(zip(columns, row))

    # =========================================================
    # TEMPLATE PATH
    # =========================================================
    def _get_template_path(self):

        base_dir = os.path.dirname(os.path.abspath(__file__))

        return os.path.abspath(
            os.path.join(
                base_dir,
                "..",
                "templates",
                self.TEMPLATE_NAME
            )
        )

    # =========================================================
    # SAFE VALUE (DATE FORMAT LONG ENGLISH)
    # =========================================================
    def _safe(self, value):

        if value is None:
            return ""

        # Native datetime/date from psycopg2
        if isinstance(value, (datetime, date)):
            return value.strftime("%B %d %Y")

        if isinstance(value, str):
            value = value.strip()

            # Try ISO date (yyyy-mm-dd or yyyy-mm-dd HH:MM:SS)
            try:
                dt = datetime.strptime(value[:10], "%Y-%m-%d")
                return dt.strftime("%B %d %Y")
            except Exception:
                pass

        return str(value)

    # =========================================================
    # REPLACE EVERYWHERE
    # =========================================================
    def _replace_placeholders_everywhere(self, doc: Document, data: dict):

        placeholders = {
            f"{{{key}}}": self._safe(value)
            for key, value in data.items()
        }

        # BODY PARAGRAPHS
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph_full(paragraph, placeholders)

        # TABLES
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph_full(
                            paragraph,
                            placeholders
                        )

        # HEADERS & FOOTERS
        for section in doc.sections:

            # Header
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph_full(
                    paragraph,
                    placeholders
                )

            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph_full(
                                paragraph,
                                placeholders
                            )

            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                self._replace_in_paragraph_full(
                    paragraph,
                    placeholders
                )

            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph_full(
                                paragraph,
                                placeholders
                            )

    # =========================================================
    # SPLIT-RUN SAFE REPLACEMENT (NO FORMAT DESTRUCTION)
    # =========================================================
    def _replace_in_paragraph_full(self, paragraph, placeholders: dict):

        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text

        for placeholder, value in placeholders.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)

        if new_text == full_text:
            return

        # Preserve formatting of first run only
        paragraph.runs[0].text = new_text

        # Clear remaining runs without removing paragraph structure
        for run in paragraph.runs[1:]:
            run.text = ""

    # =========================================================
    # REMOVE NULL BULLET LINES
    # =========================================================
    def _remove_null_paragraphs(self, doc: Document, data: dict):

        bullet_prefixes = (
            "narrative_",
            "findings_",
            "remarks_",
            "conclusion_"
        )

        null_placeholders = {
            f"{{{key}}}"
            for key, value in data.items()
            if value is None and key.startswith(bullet_prefixes)
        }

        # Body
        for paragraph in list(doc.paragraphs):
            self._remove_if_contains_null(
                paragraph,
                null_placeholders
            )

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in list(cell.paragraphs):
                        self._remove_if_contains_null(
                            paragraph,
                            null_placeholders
                        )

    # =========================================================
    # REMOVE IF CONTAINS NULL PLACEHOLDER
    # =========================================================
    def _remove_if_contains_null(self, paragraph, null_placeholders):

        if not paragraph.text:
            return

        for placeholder in null_placeholders:
            if placeholder in paragraph.text:
                self._delete_paragraph(paragraph)
                return

    # =========================================================
    # SAFE DELETE
    # =========================================================
    def _delete_paragraph(self, paragraph):

        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    # =========================================================
    # OUTPUT PATH
    # =========================================================
    def _build_output_path(self, data: dict):

        report_number = data.get("report_number") or "CARGO_CONDITION"

        safe_name = report_number.replace("/", "_").replace("\\", "_")

        temp_dir = tempfile.gettempdir()

        return os.path.join(
            temp_dir,
            f"{safe_name}_CARGO_CONDITION.docx"
        )
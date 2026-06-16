import os
import tempfile
from docx import Document
from backend_api.services.template_autofit import apply_docx_autofit


# ============================================================
# GENERATE VESSEL TRUCK SUPERVISION WORD REPORT
# ============================================================

def generate_vessel_truck_supervision_doc(data: dict) -> str:

    # ========================================================
    # LOAD TEMPLATE (RELATIVE PATH)
    # ========================================================

    base_dir = os.path.dirname(os.path.abspath(__file__))

    template_path = os.path.abspath(
        os.path.join(
            base_dir,
            "..",
            "templates",
            "vessel_truck_supervision.docx"
        )
    )

    if not os.path.exists(template_path):
        raise Exception(f"Template not found at: {template_path}")

    doc = Document(template_path)

    # ========================================================
    # SAFE VALUE
    # ========================================================

    def safe(value):
        return "" if value is None else str(value)

    # ========================================================
    # SAFE REPLACEMENT (IGUAL QUE GRAIN SAMPLING)
    # ========================================================

    def replace_in_paragraph(paragraph):

        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        # Reemplazar TODOS los placeholders del template
        for key, value in data.items():
            placeholder = f"{{{key}}}"
            if placeholder in full_text:
                full_text = full_text.replace(
                    placeholder,
                    safe(value)
                )

        # Reescribir todo el texto preservando formato
        index = 0

        for run in paragraph.runs:
            length = len(run.text)
            if length == 0:
                continue

            run.text = full_text[index:index + length]
            index += length

        # Si sobran caracteres
        if index < len(full_text):
            paragraph.runs[-1].text += full_text[index:]

    # ========================================================
    # BODY
    # ========================================================

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    # ========================================================
    # TABLES
    # ========================================================

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    # ========================================================
    # HEADERS & FOOTERS
    # ========================================================

    for section in doc.sections:

        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)

        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

    # ========================================================
    # SAVE
    # ========================================================

    output_path = os.path.join(
        tempfile.gettempdir(),
        f"{data.get('cert_no', 'truck_supervision')}.docx"
    )

    apply_docx_autofit(doc)
    doc.save(output_path)

    return output_path


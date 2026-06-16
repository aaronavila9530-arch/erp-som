import os
import tempfile
import subprocess
from typing import Dict
from docx import Document
from backend_api.services.template_autofit import apply_docx_autofit


# =====================================================
# TEMPLATE PATH
# =====================================================
TEMPLATE_PATH = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        "..",
        "templates",
        "presentation_grain_vessel.docx"
    )
)


# =====================================================
# INTERNAL — SAFE REPLACE (CROSS-RUN SAFE)
# =====================================================
def _replace_in_paragraph(paragraph, placeholder: str, value: str):

    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    if placeholder not in full_text:
        return

    start = full_text.index(placeholder)
    end = start + len(placeholder)

    current_pos = 0
    first_replacement_done = False

    for run in paragraph.runs:

        run_text = run.text
        run_len = len(run_text)

        run_start = current_pos
        run_end = current_pos + run_len

        if run_end > start and run_start < end:

            prefix_len = max(0, start - run_start)
            suffix_len = max(0, run_end - end)

            prefix = run_text[:prefix_len]
            suffix = run_text[run_len - suffix_len:] if suffix_len > 0 else ""

            if not first_replacement_done:
                run.text = prefix + value + suffix
                first_replacement_done = True
            else:
                run.text = ""

        current_pos += run_len


def _replace_in_paragraphs(paragraphs, placeholders: Dict[str, str]):
    for p in paragraphs:
        for key, value in placeholders.items():
            _replace_in_paragraph(p, key, value)


def _replace_in_tables(tables, placeholders: Dict[str, str]):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, placeholders)
                if cell.tables:
                    _replace_in_tables(cell.tables, placeholders)


# =====================================================
# MAIN — GENERATE PDF USING LIBREOFFICE (HEADLESS)
# =====================================================
def generate_vessel_presentation_doc(data: dict) -> str:

    # -------------------------------------------------
    # VALIDATIONS
    # -------------------------------------------------
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Presentation template not found: {TEMPLATE_PATH}"
        )

    if not isinstance(data, dict):
        raise ValueError("Invalid data payload — expected dict")

    # -------------------------------------------------
    # NORMALIZE DATE
    # -------------------------------------------------
    raw_dt = data.get("sampling_start_time") or ""
    sampling_date = str(raw_dt).split(" ")[0] if raw_dt else ""

    placeholders = {
        "{cert_no}": str(data.get("cert_no") or ""),
        "{vessel_name}": str(data.get("vessel_name") or ""),
        "{ship_grt}": str(data.get("ship_grt") or ""),
        "{ship_nrt}": str(data.get("ship_nrt") or ""),
        "{requested_by}": str(data.get("requested_by") or ""),
        "{sampling_start_time}": sampling_date,
    }

    # -------------------------------------------------
    # LOAD TEMPLATE
    # -------------------------------------------------
    doc = Document(TEMPLATE_PATH)

    # BODY
    _replace_in_paragraphs(doc.paragraphs, placeholders)
    _replace_in_tables(doc.tables, placeholders)

    # HEADERS / FOOTERS
    for section in doc.sections:
        _replace_in_paragraphs(section.header.paragraphs, placeholders)
        _replace_in_tables(section.header.tables, placeholders)
        _replace_in_paragraphs(section.footer.paragraphs, placeholders)
        _replace_in_tables(section.footer.tables, placeholders)

    # -------------------------------------------------
    # SAVE TEMP DOCX
    # -------------------------------------------------
    fd, temp_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    apply_docx_autofit(doc)
    doc.save(temp_docx)

    output_dir = tempfile.mkdtemp()

    # -------------------------------------------------
    # CONVERT USING LIBREOFFICE
    # -------------------------------------------------
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                temp_docx
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
            timeout=60
        )

    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice PDF conversion timed out")

    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice (soffice) is not installed or not available in PATH"
        )

    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"Error converting DOCX to PDF: {e.stderr.decode(errors='ignore')}"
        )

    # -------------------------------------------------
    # VALIDATE OUTPUT
    # -------------------------------------------------
    pdf_path = os.path.join(
        output_dir,
        os.path.splitext(os.path.basename(temp_docx))[0] + ".pdf"
    )

    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF generation failed — output file not found")

    return pdf_path


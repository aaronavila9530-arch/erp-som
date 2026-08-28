import os
import tempfile
import copy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
try:
    from services.template_autofit import apply_docx_autofit
except ModuleNotFoundError:
    from backend_api.services.template_autofit import apply_docx_autofit
from datetime import datetime


# ============================================================
# GENERATE GRAIN SAMPLING WORD REPORT (FORMAT PRESERVING)
# ============================================================

def generate_grain_sampling_doc(data: dict) -> str:

    # ========================================================
    # LOAD TEMPLATE
    # ========================================================

    base_dir = os.path.dirname(os.path.abspath(__file__))

    template_path = os.path.abspath(
        os.path.join(
            base_dir,
            "..",
            "templates",
            "Supervision_Muestreo_Granos.docx"
        )
    )

    if not os.path.exists(template_path):
        raise Exception(f"Template not found at: {template_path}")

    doc = Document(template_path)

    def _non_empty(value):
        return str(value or "").strip()

    def _product_rows():
        rows = []
        for idx in range(1, 6):
            product = _non_empty(data.get(f"hold{idx}_product"))
            hold = _non_empty(data.get(f"hold{idx}_hold")) or str(idx)
            tonnage = _non_empty(data.get(f"hold{idx}_tonnage"))
            if product or tonnage:
                rows.append({
                    "product": product,
                    "hold": hold,
                    "tonnage": tonnage,
                })
        return rows

    def _sample_rows():
        rows = []
        for idx in range(1, 6):
            hold = _non_empty(data.get(f"sample{idx}_hold"))
            if not hold:
                continue
            rows.append((idx, hold))
        return rows

    def _products_summary():
        rows = _product_rows()
        products = []
        holds = []
        for row in rows:
            if row["product"] and row["product"] not in products:
                products.append(row["product"])
            if row["hold"]:
                holds.append(row["hold"])
        product_text = ", ".join(products) if products else "producto a granel"
        hold_text = ", ".join(holds) if holds else ""
        return product_text, hold_text

    product_text, holds_text = _products_summary()
    data["products_name_summary"] = product_text
    data["sampled_holds_summary"] = holds_text

    # ========================================================
    # DATE → LONG ENGLISH FORMAT (ULTRA SAFE)
    # ========================================================

    def format_date_long_en(value):
        """
        Convierte fechas a formato:
        12 March 2026

        Soporta:
        - datetime
        - string YYYY-MM-DD
        - ISO strings
        - fallback seguro
        """
        if not value:
            return ""

        try:
            # datetime directo
            if isinstance(value, datetime):
                dt = value
            else:
                v = str(value).strip()

                # quitar tiempo si viene
                if "T" in v:
                    v = v.split("T")[0]
                if " " in v:
                    v = v.split(" ")[0]

                dt = datetime.strptime(v, "%Y-%m-%d")

            return dt.strftime("%d %B %Y")

        except Exception:
            return str(value)

    # ========================================================
    # SAFE VALUE (AUTO DATE FORMAT)
    # ========================================================

    def safe(key, value):
        if value is None:
            return ""

        # Detectar campos tipo fecha automáticamente
        key_lower = key.lower()

        if "date" in key_lower:
            return format_date_long_en(value)

        return str(value)

    # ========================================================
    # ULTRA SAFE REPLACEMENT
    # Mantiene runs originales
    # Mantiene formato
    # No borra imágenes
    # No colapsa estilos
    # ========================================================

    def replace_in_paragraph(paragraph, data_dict):

        if not paragraph.runs:
            return

        # Texto lógico completo
        full_text = "".join(run.text for run in paragraph.runs)

        modified = False

        for key, value in data_dict.items():
            placeholder = f"{{{key}}}"

            if placeholder in full_text:
                full_text = full_text.replace(
                    placeholder,
                    safe(key, value)
                )
                modified = True

        if not modified:
            return

        # Reescritura sin romper formato
        index = 0

        for run in paragraph.runs:

            original_length = len(run.text)

            if original_length == 0:
                continue

            run.text = full_text[index:index + original_length]
            index += original_length

        # Si crece el texto → append al último run
        if index < len(full_text):
            paragraph.runs[-1].text += full_text[index:]

    def set_paragraph_text(paragraph, text):
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    def insert_paragraph_after(paragraph, text):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if paragraph.style:
            new_para.style = paragraph.style
        new_para.add_run(text)
        return new_para

    def set_cell_text(cell, text):
        value = str(text or "")
        if cell.paragraphs:
            set_paragraph_text(cell.paragraphs[0], value)
            for paragraph in cell.paragraphs[1:]:
                set_paragraph_text(paragraph, "")
        else:
            cell.text = value

    def _apply_dynamic_narrative():
        sample_rows = _sample_rows()
        sampled_holds = holds_text or ", ".join(hold for _, hold in sample_rows)
        intro = (
            f"MSL Marine Surveyors and Logistics, fuimos nominados para llevar a cabo "
            f"la supervision de toma de muestras producto {product_text} a bordo del "
            f"MV {data.get('vessel_name') or ''} en la Bahia de Puerto Caldera - Costa Rica."
        )
        product_sentence = (
            f"{data.get('products_total') or ''} mt {product_text} a granel cargado "
            f"en Bodegas No. {sampled_holds}."
        ).strip()
        sampling_sentence = (
            f"Los productos almacenados en las bodegas del buque fueron muestreados "
            f"por los funcionarios del MAG, se muestreo en bodega {sampled_holds} "
            f"segun plan de carga."
        )

        sample3_para = None
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if "fuimos nominados" in text and "producto" in text:
                set_paragraph_text(paragraph, intro)
            elif "{products_total}" in text and "Bodegas" in text:
                set_paragraph_text(paragraph, product_sentence)
            elif "Los productos almacenados" in text and "segun plan" in text:
                set_paragraph_text(paragraph, sampling_sentence)
            elif "Toma Muestra Bodega {sample3_hold}" in text:
                sample3_para = paragraph

        if sample3_para:
            current = sample3_para
            for idx in (4, 5):
                hold = _non_empty(data.get(f"sample{idx}_hold"))
                if not hold:
                    continue
                text = (
                    f"Toma Muestra Bodega {hold}: Inician tomando muestras en una distancia "
                    f"aproximada de 2x3 con una secuencia de, Proa babor "
                    f"{data.get(f'sample{idx}_proa_babor') or ''} puntos, Proa estribor "
                    f"{data.get(f'sample{idx}_proa_estribor') or ''} puntos, Centro "
                    f"{data.get(f'sample{idx}_centro') or ''} puntos, Popa Babor "
                    f"{data.get(f'sample{idx}_popa_babor') or ''} Puntos, finalizando "
                    f"Popa estribor {data.get(f'sample{idx}_popa_estribor') or ''} puntos."
                )
                current = insert_paragraph_after(current, text)

    def _ensure_product_table_rows(table):
        if not table.rows:
            return
        header = " ".join(cell.text.upper() for cell in table.rows[0].cells)
        if "PRODUCTO" not in header or "BODEGA" not in header:
            return

        while len(table.rows) < 7:
            source = table.rows[-2]._tr
            new_row = copy.deepcopy(source)
            table.rows[-1]._tr.addprevious(new_row)

    def _fill_product_table(table):
        if not table.rows:
            return
        header = " ".join(cell.text.upper() for cell in table.rows[0].cells)
        if "PRODUCTO" not in header or "BODEGA" not in header:
            return

        rows = _product_rows()
        _ensure_product_table_rows(table)
        for idx in range(1, 6):
            row = table.rows[idx]
            value = rows[idx - 1] if idx <= len(rows) else {"product": "", "hold": str(idx), "tonnage": ""}
            set_cell_text(row.cells[0], value["product"])
            set_cell_text(row.cells[1], value["hold"])
            set_cell_text(row.cells[2], value["tonnage"])
        set_cell_text(table.rows[6].cells[0], "TOTAL")
        set_cell_text(table.rows[6].cells[1], "-")
        set_cell_text(table.rows[6].cells[2], data.get("products_total") or "")

    _apply_dynamic_narrative()

    # ========================================================
    # PROCESS BODY
    # ========================================================

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)

    # ========================================================
    # PROCESS TABLES
    # ========================================================

    for table in doc.tables:
        _ensure_product_table_rows(table)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)
        _fill_product_table(table)

    # ========================================================
    # PROCESS HEADERS & FOOTERS
    # ========================================================

    for section in doc.sections:

        # HEADER
        header = section.header

        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph, data)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, data)

        # FOOTER
        footer = section.footer

        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph, data)

        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, data)

    # ========================================================
    # SAVE FILE
    # ========================================================

    output_path = os.path.join(
        tempfile.gettempdir(),
        f"{data.get('cert_no', 'grain_sampling')}.docx"
    )

    apply_docx_autofit(doc)
    doc.save(output_path)

    return output_path

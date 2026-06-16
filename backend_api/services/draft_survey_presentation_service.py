import os
import re
import tempfile
import subprocess
from typing import Dict
from datetime import datetime
from docx import Document
try:
    from services.template_autofit import apply_docx_autofit
except ModuleNotFoundError:
    from backend_api.services.template_autofit import apply_docx_autofit


# =====================================================
# TEMPLATE PATH
# =====================================================
TEMPLATE_PATH = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        "..",
        "templates",
        "presentation_draft_survey.docx"
    )
)


# =====================================================
# FORMATTERS
# =====================================================
def _format_vessel_name(value):

    raw = str(value or "").strip()
    if not raw:
        return ""

    raw = raw.replace('"', "").strip()

    # quitar prefijo mv repetido si ya viene pegado
    lower_raw = raw.lower()
    if lower_raw.startswith("mv "):
        core = raw[3:].strip()
    elif lower_raw.startswith("mv"):
        core = raw[2:].strip()
    else:
        core = raw

    # separar CamelCase / mayúsculas pegadas / guiones bajos
    core = core.replace("_", " ").replace("-", " ")
    core = re.sub(r"([a-z])([A-Z])", r"\1 \2", core)

    # si viene todo pegado en mayúsculas, al menos Title()
    core = " ".join(core.split()).title()

    return f"MV {core}".strip()


def _format_number_no_decimal(value):

    if value in (None, ""):
        return ""

    try:
        num = float(str(value).replace(",", "").strip())
        return f"{int(round(num)):,}"
    except Exception:
        return str(value)


def _format_date_long_en(value):

    if not value:
        return ""

    if isinstance(value, datetime):
        return value.strftime("%B %d, %Y")

    raw = str(value).strip()
    if not raw:
        return ""

    candidates = [
        "%Y-%m-%d",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%d-%m-%Y",
        "%d-%m-%Y %H:%M:%S",
        "%d-%m-%Y %H:%M",
        "%m-%d-%Y",
        "%m-%d-%Y %H:%M:%S",
        "%m-%d-%Y %H:%M",
    ]

    # si viene ISO con T
    raw_iso = raw.replace("T", " ").replace("Z", "")

    for fmt in candidates:
        try:
            dt = datetime.strptime(raw_iso[:19], fmt)
            return dt.strftime("%B %d, %Y")
        except Exception:
            pass

    try:
        dt = datetime.fromisoformat(raw_iso[:19])
        return dt.strftime("%B %d, %Y")
    except Exception:
        return raw


def _format_place_line(port, country, commenced):

    p = str(port or "").strip()
    c = str(country or "").strip()
    d = _format_date_long_en(commenced)

    if p and c and d:
        return f"{p} - {c}, {d}"
    if p and d:
        return f"{p}, {d}"
    if c and d:
        return f"{c}, {d}"
    if p and c:
        return f"{p} - {c}"
    return d or p or c or ""


# =====================================================
# INTERNAL — SAFE REPLACE
# =====================================================
def _replace_in_paragraph(paragraph, placeholder: str, value: str):

    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    if placeholder not in full_text:
        return

    start = full_text.index(placeholder)
    end = start + len(placeholder)

    current_pos = 0
    first_replacement_done = False

    for run in paragraph.runs:

        run_text = run.text
        run_len = len(run_text)

        run_start = current_pos
        run_end = current_pos + run_len

        if run_end > start and run_start < end:

            prefix_len = max(0, start - run_start)
            suffix_len = max(0, run_end - end)

            prefix = run_text[:prefix_len]
            suffix = run_text[run_len - suffix_len:] if suffix_len > 0 else ""

            if not first_replacement_done:
                run.text = prefix + value + suffix
                first_replacement_done = True
            else:
                run.text = ""

        current_pos += run_len


def _replace_in_paragraphs(paragraphs, placeholders: Dict[str, str]):
    for p in paragraphs:
        for key, value in placeholders.items():
            _replace_in_paragraph(p, key, value)


def _replace_in_tables(tables, placeholders: Dict[str, str]):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, placeholders)
                if cell.tables:
                    _replace_in_tables(cell.tables, placeholders)


# =====================================================
# MAIN — GENERATE PRESENTATION PDF
# =====================================================
def generate_draft_survey_presentation_pdf(data: dict) -> str:

    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Draft Survey presentation template not found: {TEMPLATE_PATH}"
        )

    vessel_name = _format_vessel_name(data.get("word_vessel"))
    grt = _format_number_no_decimal(data.get("word_grt"))
    nrt = _format_number_no_decimal(data.get("word_nrt"))
    client = str(data.get("word_survey_requested_by") or "").strip()
    cert_no = str(data.get("draft_report_number") or "").strip()
    port = str(data.get("word_port") or "").strip()
    country = str(data.get("word_country") or "").strip()
    commenced_long = _format_date_long_en(data.get("word_commenced"))
    place_line = _format_place_line(
        port=port,
        country=country,
        commenced=data.get("word_commenced")
    )

    placeholders = {
        "{draft_report_number}": cert_no,
        "{word_vessel}": vessel_name,
        "{word_grt}": grt,
        "{word_nrt}": nrt,
        "{word_survey_requested_by}": client,
        "{word_port}": port,
        "{word_country}": country,
        "{word_commenced}": commenced_long,
        "{word_place_line}": place_line,
    }

    doc = Document(TEMPLATE_PATH)

    _replace_in_paragraphs(doc.paragraphs, placeholders)
    _replace_in_tables(doc.tables, placeholders)

    for section in doc.sections:
        _replace_in_paragraphs(section.header.paragraphs, placeholders)
        _replace_in_tables(section.header.tables, placeholders)
        _replace_in_paragraphs(section.footer.paragraphs, placeholders)
        _replace_in_tables(section.footer.tables, placeholders)

    fd, temp_docx = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    apply_docx_autofit(doc)
    doc.save(temp_docx)

    output_dir = tempfile.mkdtemp()

    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            temp_docx
        ],
        check=True
    )

    pdf_path = os.path.join(
        output_dir,
        os.path.splitext(os.path.basename(temp_docx))[0] + ".pdf"
    )

    if not os.path.exists(pdf_path):
        raise RuntimeError("Presentation PDF generation failed")

    return pdf_path

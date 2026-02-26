import os
import tempfile
import subprocess
from pathlib import Path
from docx import Document


# ============================================================
# GENERATE DRAFT SURVEY WORD PDF (ERP VERSION - BLINDADO)
# ============================================================

def generate_draft_survey_word_pdf(data: dict) -> str:

    # ========================================================
    # VALIDACIÓN
    # ========================================================

    if not isinstance(data, dict):
        raise ValueError("Invalid payload. Expected dict.")

    draft_report_number = str(
        data.get("draft_report_number") or ""
    ).strip()

    if not draft_report_number:
        raise ValueError("draft_report_number is required")

    # ========================================================
    # TEMPLATE PATH
    # ========================================================

    base_dir = os.path.dirname(os.path.abspath(__file__))

    template_path = os.path.abspath(
        os.path.join(
            base_dir,
            "..",
            "templates",
            "draft_word_template.docx"
        )
    )

    if not os.path.exists(template_path):
        raise FileNotFoundError(
            f"Template not found at: {template_path}"
        )

    doc = Document(template_path)

    def safe(value):
        return "" if value is None else str(value)

    # ========================================================
    # 🔥 FLATTEN PLACEHOLDER MAP
    # ========================================================

    placeholder_map = {
        f"{{{key}}}": safe(value)
        for key, value in data.items()
    }

        # ========================================================
        # REPLACEMENT ENGINE (ANTI-RUN SPLIT + PRESERVE FORMAT)
        # ========================================================

        def replace_in_paragraph(paragraph):

            # 🔒 Blindaje básico
            if not paragraph or not paragraph.runs:
                return

            try:

                # Construir texto completo preservando orden
                full_text = "".join(run.text for run in paragraph.runs)

                if not full_text:
                    return

                # Reemplazar placeholders
                updated_text = full_text

                for placeholder, value in placeholder_map.items():

                    if placeholder in updated_text:
                        updated_text = updated_text.replace(
                            placeholder,
                            value if value is not None else ""
                        )

                # Si nada cambió, no tocar runs
                if updated_text == full_text:
                    return

                # Reescribir preservando formato original
                index = 0

                for run in paragraph.runs:

                    original_length = len(run.text)

                    if original_length == 0:
                        continue

                    slice_text = updated_text[index:index + original_length]

                    run.text = slice_text
                    index += original_length

                # Si el texto nuevo es más largo que los runs originales
                if index < len(updated_text):
                    paragraph.runs[-1].text += updated_text[index:]

            except Exception:
                # 🔒 Nunca permitir que el PDF explote por reemplazo
                return


        # ========================================================
        # BODY
        # ========================================================

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)


        # ========================================================
        # TABLES
        # ========================================================

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)


        # ========================================================
        # HEADERS & FOOTERS
        # ========================================================

        for section in doc.sections:

            # Header paragraphs
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph)

            # Header tables
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)

            # Footer paragraphs
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph)

            # Footer tables
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)

    # ========================================================
    # SAVE TEMP DOCX
    # ========================================================

    temp_docx = os.path.join(
        tempfile.gettempdir(),
        f"{draft_report_number}.docx"
    )

    doc.save(temp_docx)

    # ========================================================
    # LIBREOFFICE EXECUTABLE DETECTION
    # ========================================================

    soffice_path = os.getenv("LIBREOFFICE_PATH", "soffice")

    output_dir = tempfile.mkdtemp(prefix="draft_word_pdf_")
    libre_profile = tempfile.mkdtemp(prefix="lo_profile_")

    cmd = [
        soffice_path,
        "--headless",
        "--nologo",
        "--nodefault",
        "--nolockcheck",
        "--nofirststartwizard",
        "--norestore",
        f"-env:UserInstallation=file://{libre_profile}",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        output_dir,
        temp_docx
    ]

    result = subprocess.run(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice PDF conversion failed:\n{result.stderr}"
        )

    pdf_path = os.path.join(
        output_dir,
        f"{draft_report_number}.pdf"
    )

    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF was not created")

    if os.path.getsize(pdf_path) == 0:
        raise RuntimeError("PDF was generated but is empty")

    return pdf_path
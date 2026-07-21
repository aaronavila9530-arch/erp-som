from datetime import datetime
import json
import os
import shutil
import tempfile
from pathlib import Path

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse
from psycopg2.extras import Json, RealDictCursor

from database import get_db


router = APIRouter(
    prefix="/logra-reports",
    tags=["LOGRA Reports"]
)

STORAGE_ROOT = Path("storage") / "logra"
MAX_AGENDA_ITEMS = 150
MAX_ATTACHMENTS_PER_QUESTION = 10

TEST_WORDS = {"test", "testing", "prueba", "demo", "dummy", "asdf", "qwerty", "deploy check"}


def _ensure_schema(conn):
    with conn.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS logra_reports (
                id SERIAL PRIMARY KEY,
                title TEXT,
                category TEXT DEFAULT 'ONG',
                meeting_date DATE DEFAULT CURRENT_DATE,
                meeting_time TEXT DEFAULT '00:00',
                meeting_start_time TEXT DEFAULT '',
                meeting_end_time TEXT DEFAULT '',
                meeting_location TEXT DEFAULT '',
                meeting_person TEXT DEFAULT '',
                status TEXT DEFAULT 'Pending',
                agenda_items JSONB NOT NULL DEFAULT '[]'::jsonb,
                agenda_notes TEXT DEFAULT '',
                created_by TEXT,
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS title TEXT
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS category TEXT DEFAULT 'ONG'
        """)
        cur.execute("""
            UPDATE logra_reports
            SET category = 'ONG'
            WHERE category IS NULL OR UPPER(category) = 'LOGRA'
        """)
        cur.execute("""
            UPDATE logra_reports
            SET title = REPLACE(title, 'LOGRA', 'ONG')
            WHERE title LIKE '%LOGRA%'
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN category SET DEFAULT 'ONG'
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS status TEXT DEFAULT 'Pending'
        """)
        cur.execute("""
            UPDATE logra_reports
            SET status = 'Pending'
            WHERE category = 'ONG'
              AND (status IS NULL OR UPPER(status) = 'DRAFT')
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN status SET DEFAULT 'Pending'
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS meeting_date DATE DEFAULT CURRENT_DATE
        """)
        cur.execute("""
            UPDATE logra_reports
            SET meeting_date = CURRENT_DATE
            WHERE meeting_date IS NULL
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN meeting_date SET DEFAULT CURRENT_DATE
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS meeting_time TEXT DEFAULT '00:00'
        """)
        cur.execute("""
            UPDATE logra_reports
            SET meeting_time = '00:00'
            WHERE meeting_time IS NULL
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN meeting_time SET DEFAULT '00:00'
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS meeting_start_time TEXT DEFAULT ''
        """)
        cur.execute("""
            UPDATE logra_reports
            SET meeting_start_time = ''
            WHERE meeting_start_time IS NULL
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN meeting_start_time SET DEFAULT ''
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS meeting_end_time TEXT DEFAULT ''
        """)
        cur.execute("""
            UPDATE logra_reports
            SET meeting_end_time = ''
            WHERE meeting_end_time IS NULL
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN meeting_end_time SET DEFAULT ''
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS meeting_location TEXT DEFAULT ''
        """)
        cur.execute("""
            UPDATE logra_reports
            SET meeting_location = ''
            WHERE meeting_location IS NULL
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN meeting_location SET DEFAULT ''
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS meeting_person TEXT DEFAULT ''
        """)
        cur.execute("""
            UPDATE logra_reports
            SET meeting_person = ''
            WHERE meeting_person IS NULL
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN meeting_person SET DEFAULT ''
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS agenda_items JSONB NOT NULL DEFAULT '[]'::jsonb
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS agenda_notes TEXT DEFAULT ''
        """)
        cur.execute("""
            UPDATE logra_reports
            SET agenda_notes = ''
            WHERE agenda_notes IS NULL
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ALTER COLUMN agenda_notes SET DEFAULT ''
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS created_by TEXT
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS created_at TIMESTAMP DEFAULT NOW()
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP DEFAULT NOW()
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS form_slug TEXT
        """)
        cur.execute("""
            ALTER TABLE logra_reports
            ADD COLUMN IF NOT EXISTS version INTEGER NOT NULL DEFAULT 1
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS logra_answers (
                id SERIAL PRIMARY KEY,
                report_id INTEGER NOT NULL REFERENCES logra_reports(id) ON DELETE CASCADE,
                form_slug TEXT NOT NULL,
                form_title TEXT,
                section TEXT NOT NULL,
                item_key TEXT NOT NULL,
                question_text TEXT NOT NULL,
                bullets JSONB NOT NULL DEFAULT '[]'::jsonb,
                updated_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(report_id, form_slug, section, item_key)
            )
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS logra_attachments (
                id SERIAL PRIMARY KEY,
                report_id INTEGER NOT NULL REFERENCES logra_reports(id) ON DELETE CASCADE,
                form_slug TEXT NOT NULL,
                section TEXT NOT NULL,
                item_key TEXT NOT NULL,
                bullet_index INTEGER,
                original_filename TEXT NOT NULL,
                stored_path TEXT NOT NULL,
                content_type TEXT,
                file_size BIGINT,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_logra_answers_report
            ON logra_answers(report_id)
        """)
        cur.execute("""
            UPDATE logra_reports r
            SET form_slug = source.form_slug
            FROM (
                SELECT report_id, MIN(form_slug) AS form_slug
                FROM logra_answers
                GROUP BY report_id
                HAVING COUNT(DISTINCT form_slug) = 1
            ) source
            WHERE r.id = source.report_id
              AND (r.form_slug IS NULL OR BTRIM(r.form_slug) = '')
        """)
        cur.execute("""
            UPDATE logra_answers
            SET form_title = REPLACE(form_title, 'LOGRA', 'ONG')
            WHERE form_title LIKE '%LOGRA%'
        """)
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_logra_attachments_lookup
            ON logra_attachments(report_id, form_slug, section, item_key)
        """)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS logra_report_revisions (
                id BIGSERIAL PRIMARY KEY,
                report_id INTEGER NOT NULL REFERENCES logra_reports(id) ON DELETE CASCADE,
                version INTEGER NOT NULL,
                form_slug TEXT,
                snapshot JSONB NOT NULL,
                saved_by TEXT,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_logra_revisions_report
            ON logra_report_revisions(report_id, version DESC)
        """)
    conn.commit()


def _safe_filename(filename: str) -> str:
    base = os.path.basename(filename or "attachment")
    safe = "".join(ch if ch.isalnum() or ch in "._- " else "_" for ch in base).strip()
    return safe or "attachment"


def _clean_answers(answers, expected_form_slug=""):
    if not isinstance(answers, list):
        raise HTTPException(status_code=400, detail="answers must be a list")
    cleaned = []
    seen = set()
    expected_form_slug = str(expected_form_slug or "").strip()
    for raw in answers:
        if not isinstance(raw, dict):
            continue
        form_slug = str(raw.get("form_slug") or expected_form_slug).strip()
        section = str(raw.get("section") or "").strip()
        item_key = str(raw.get("item_key") or "").strip()
        if not form_slug or not section or not item_key:
            raise HTTPException(
                status_code=400,
                detail="Each answer requires form_slug, section and item_key",
            )
        if expected_form_slug and form_slug != expected_form_slug:
            raise HTTPException(
                status_code=409,
                detail=f"Answer form_slug '{form_slug}' does not match report form_slug '{expected_form_slug}'",
            )
        key = (form_slug, section, item_key)
        if key in seen:
            raise HTTPException(status_code=400, detail=f"Duplicate answer key in payload: {key}")
        seen.add(key)
        bullets = raw.get("bullets") or []
        if not isinstance(bullets, list):
            bullets = []
        bullets = [str(value).strip() for value in bullets[:20] if str(value or "").strip()]
        if not bullets:
            continue
        cleaned.append({
            "form_slug": form_slug,
            "form_title": (raw.get("form_title") or "").replace("LOGRA", "ONG"),
            "section": section,
            "item_key": item_key,
            "question_text": raw.get("question_text") or "",
            "bullets": bullets,
        })
    return cleaned


def _save_revision(cur, report_id, version, form_slug, saved_by):
    cur.execute("SELECT * FROM logra_reports WHERE id = %s", (report_id,))
    report = cur.fetchone()
    cur.execute("""
        SELECT form_slug, form_title, section, item_key, question_text, bullets, updated_at
        FROM logra_answers
        WHERE report_id = %s
        ORDER BY form_slug, section, item_key
    """, (report_id,))
    answers = cur.fetchall()
    snapshot = {"report": dict(report or {}), "answers": [dict(item) for item in answers]}
    cur.execute("""
        INSERT INTO logra_report_revisions (report_id, version, form_slug, snapshot, saved_by, created_at)
        VALUES (%s, %s, %s, %s, %s, %s)
    """, (
        report_id,
        version,
        form_slug,
        Json(snapshot, dumps=lambda value: json.dumps(value, default=str)),
        saved_by,
        datetime.utcnow(),
    ))


def _looks_like_test_text(text: str) -> bool:
    clean = " ".join(str(text or "").strip().lower().split())
    if not clean:
        return True
    if len(clean) <= 2:
        return True
    tokens = set(clean.replace("-", " ").replace("_", " ").split())
    if clean in TEST_WORDS:
        return True
    if {"deploy", "check"}.issubset(tokens):
        return True
    return bool(tokens.intersection(TEST_WORDS))


def _valid_ai_bullets(raw_bullets):
    valid = []
    if not isinstance(raw_bullets, list):
        return valid
    for bullet in raw_bullets:
        text = str(bullet or "").strip()
        if not text:
            continue
        if _looks_like_test_text(text):
            continue
        valid.append(text)
    return valid


def _get_logra_ai_context(conn, report_id: int):
    _ensure_schema(conn)
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("SELECT * FROM logra_reports WHERE id = %s", (report_id,))
        report = cur.fetchone()
        if not report:
            raise HTTPException(status_code=404, detail="LOGRA report not found")

        title = str(report.get("title") or "")
        if _looks_like_test_text(title.replace("ONG -", "").replace("LOGRA -", "")):
            raise HTTPException(status_code=400, detail="Este reporte parece ser de prueba y no se considera para informe IA.")

        cur.execute("""
            SELECT *
            FROM logra_answers
            WHERE report_id = %s
            ORDER BY form_title, section, item_key
        """, (report_id,))
        answers = cur.fetchall()

    valid_answers = []
    for answer in answers:
        bullets = _valid_ai_bullets(answer.get("bullets") or [])
        if not bullets:
            continue
        question = str(answer.get("question_text") or "").strip()
        if not question or _looks_like_test_text(question):
            continue
        valid_answers.append({
            "form_title": str(answer.get("form_title") or "").replace("LOGRA", "ONG"),
            "section": answer.get("section") or "",
            "item_key": answer.get("item_key") or "",
            "question_text": question,
            "bullets": bullets,
        })

    if not valid_answers:
        raise HTTPException(status_code=400, detail="No hay respuestas reales para generar el informe IA.")

    return report, valid_answers


def _get_logra_ai_context_all(conn):
    _ensure_schema(conn)
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("""
            SELECT a.*, r.title AS report_title, r.status AS report_status, r.updated_at AS report_updated_at
            FROM logra_answers a
            JOIN logra_reports r ON r.id = a.report_id
            WHERE COALESCE(r.title, '') <> 'ONG - Agenda'
              AND COALESCE(r.category, 'ONG') = 'ONG'
            ORDER BY r.title, a.form_title, a.section, a.item_key
        """)
        answers = cur.fetchall()

    valid_answers = []
    seen = set()
    for answer in answers:
        report_title = str(answer.get("report_title") or "")
        if _looks_like_test_text(report_title.replace("ONG -", "").replace("LOGRA -", "")):
            continue
        bullets = _valid_ai_bullets(answer.get("bullets") or [])
        if not bullets:
            continue
        question = str(answer.get("question_text") or "").strip()
        if not question or _looks_like_test_text(question):
            continue
        key = (
            str(answer.get("form_title") or ""),
            str(answer.get("section") or ""),
            str(answer.get("item_key") or ""),
            question,
            tuple(bullets),
        )
        if key in seen:
            continue
        seen.add(key)
        valid_answers.append({
            "report_title": report_title.replace("LOGRA", "ONG"),
            "form_title": str(answer.get("form_title") or "").replace("LOGRA", "ONG"),
            "section": answer.get("section") or "",
            "item_key": answer.get("item_key") or "",
            "question_text": question,
            "bullets": bullets,
        })

    if not valid_answers:
        raise HTTPException(status_code=400, detail="No hay respuestas reales para generar el informe IA consolidado.")

    report = {
        "id": "ALL",
        "title": "ONG Executive Consolidated Questionnaire Report",
        "status": "Consolidated",
    }
    return report, valid_answers


@router.get("")
def list_logra_reports(conn=Depends(get_db)):
    _ensure_schema(conn)
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("""
            SELECT
                r.id,
                r.title,
                r.category,
                r.status,
                r.form_slug,
                r.version,
                r.agenda_items,
                r.agenda_notes,
                r.created_by,
                r.created_at,
                r.updated_at,
                COALESCE(a.attachment_count, 0) AS attachment_count
            FROM logra_reports r
            LEFT JOIN (
                SELECT report_id, COUNT(*) AS attachment_count
                FROM logra_attachments
                GROUP BY report_id
            ) a ON a.report_id = r.id
            ORDER BY r.updated_at DESC, r.id DESC
            LIMIT 200
        """)
        return {"data": cur.fetchall()}


@router.post("/agenda-only")
def save_logra_agenda_only(payload: dict, conn=Depends(get_db)):
    _ensure_schema(conn)
    payload = payload or {}
    agenda_items = payload.get("agenda_items") or []
    agenda_notes = payload.get("agenda_notes") or ""
    created_by = payload.get("created_by")
    if not isinstance(agenda_items, list):
        agenda_items = []
    if len(agenda_items) > MAX_AGENDA_ITEMS:
        raise HTTPException(status_code=400, detail="LOGRA agenda supports up to 150 items")

    clean_items = []
    for index, raw_item in enumerate(agenda_items):
        if not isinstance(raw_item, dict):
            continue
        item = dict(raw_item)
        item["report_title"] = "ONG - Agenda"
        item["agenda_index"] = index
        clean_items.append(item)

    first_meeting = clean_items[0] if clean_items else {}
    meeting_date = first_meeting.get("date_iso") or datetime.utcnow().date().isoformat()
    meeting_start_time = first_meeting.get("start_time") or ""
    meeting_end_time = first_meeting.get("end_time") or ""
    meeting_time = meeting_start_time or "00:00"
    meeting_location = first_meeting.get("place") or ""
    meeting_person = first_meeting.get("person") or ""

    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT id
                FROM logra_reports
                WHERE category = 'ONG'
                  AND title = 'ONG - Agenda'
                ORDER BY id ASC
                LIMIT 1
            """)
            existing = cur.fetchone()

            if existing:
                report_id = existing["id"]
                cur.execute("""
                    UPDATE logra_reports
                    SET status = 'Pending',
                        meeting_date = %s,
                        meeting_time = %s,
                        meeting_start_time = %s,
                        meeting_end_time = %s,
                        meeting_location = %s,
                        meeting_person = %s,
                        agenda_items = %s,
                        agenda_notes = %s,
                        updated_at = %s
                    WHERE id = %s
                    RETURNING *
                """, (
                    meeting_date, meeting_time, meeting_start_time, meeting_end_time,
                    meeting_location, meeting_person, Json(clean_items), agenda_notes,
                    datetime.utcnow(), report_id,
                ))
                report = cur.fetchone()
            else:
                cur.execute("""
                    INSERT INTO logra_reports (
                        title, category, status, meeting_date, meeting_time, meeting_start_time,
                        meeting_end_time, meeting_location, meeting_person, agenda_items,
                        agenda_notes, created_by, created_at, updated_at
                    )
                    VALUES ('ONG - Agenda', 'ONG', 'Pending', %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING *
                """, (
                    meeting_date, meeting_time, meeting_start_time, meeting_end_time,
                    meeting_location, meeting_person, Json(clean_items), agenda_notes,
                    created_by, datetime.utcnow(), datetime.utcnow(),
                ))
                report = cur.fetchone()
                report_id = report["id"]

            for index, item in enumerate(clean_items):
                item["report_id"] = report_id
                item["agenda_index"] = index

            cur.execute("""
                UPDATE logra_reports
                SET agenda_items = %s,
                    updated_at = %s
                WHERE id = %s
                RETURNING *
            """, (Json(clean_items), datetime.utcnow(), report_id))
            report = cur.fetchone()

        conn.commit()
        return {"success": True, "report": report}
    except HTTPException:
        conn.rollback()
        raise
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA agenda save error: {exc}")


@router.get("/ai-report/all/word")
def download_logra_ai_report_all_word(
    language: str = Query("ES"),
    conn=Depends(get_db)
):
    report, answers = _get_logra_ai_context_all(conn)
    return _build_logra_ai_word_response(report, answers, language)


@router.get("/ai-report/all/pdf")
def download_logra_ai_report_all_pdf(
    language: str = Query("ES"),
    conn=Depends(get_db)
):
    report, answers = _get_logra_ai_context_all(conn)
    return _build_logra_ai_pdf_response(report, answers, language)


def _build_logra_ai_word_response(report: dict, answers: list[dict], language: str = "ES"):
    try:
        from ai.maritime_ai import generate_logra_questionnaire_report
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        ai_text = generate_logra_questionnaire_report(report.get("title") or "ONG Report", answers, language=language)

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(str(report.get("title") or "ONG AI Report"))
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 59, 113)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run("ONG AI Executive Consolidated Report").italic = True

        meta = doc.add_table(rows=0, cols=2)
        meta.style = "Table Grid"
        for label, value in (
            ("Scope", "All valid ONG questionnaires"),
            ("Language", (language or "ES").upper()),
            ("Valid answered questions", len(answers)),
            ("Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
        ):
            cells = meta.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()
        _write_ai_text_to_docx(doc, ai_text)

        tmp_dir = tempfile.mkdtemp(prefix="ong_ai_report_all_")
        filename = _safe_filename("ONG_Executive_Consolidated_AI_Report.docx")
        path = os.path.join(tmp_dir, filename)
        doc.save(path)
        return FileResponse(path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"ONG AI Word report error: {exc}")


def _build_logra_ai_pdf_response(report: dict, answers: list[dict], language: str = "ES"):
    try:
        from ai.maritime_ai import generate_logra_questionnaire_report
        from html import escape
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        ai_text = generate_logra_questionnaire_report(report.get("title") or "ONG Report", answers, language=language)

        tmp_dir = tempfile.mkdtemp(prefix="ong_ai_report_all_")
        filename = _safe_filename("ONG_Executive_Consolidated_AI_Report.pdf")
        path = os.path.join(tmp_dir, filename)

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="OngTitle", parent=styles["Title"], fontSize=16, textColor=colors.HexColor("#003B71"), alignment=1))
        styles.add(ParagraphStyle(name="OngH", parent=styles["Heading1"], fontSize=12, textColor=colors.HexColor("#003B71"), spaceBefore=10, spaceAfter=5))
        styles.add(ParagraphStyle(name="OngBody", parent=styles["BodyText"], fontSize=9, leading=12, spaceAfter=5))

        story = [
            Paragraph(escape(str(report.get("title") or "ONG AI Report")), styles["OngTitle"]),
            Paragraph("ONG AI Executive Consolidated Report", styles["Normal"]),
            Spacer(1, 8),
        ]
        meta = [
            ["Scope", "All valid ONG questionnaires"],
            ["Language", (language or "ES").upper()],
            ["Valid answered questions", str(len(answers))],
            ["Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")],
        ]
        table = Table(meta, colWidths=[1.9 * inch, 4.8 * inch])
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C1CA")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E9EEF3")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ]))
        story.extend([table, Spacer(1, 10)])
        _write_ai_text_to_pdf(story, styles, ai_text)

        doc = SimpleDocTemplate(path, pagesize=letter, rightMargin=0.6 * inch, leftMargin=0.6 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
        doc.build(story)
        return FileResponse(path, filename=filename, media_type="application/pdf")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"ONG AI PDF report error: {exc}")


@router.get("/{report_id}/ai-report/word")
def download_logra_ai_report_word(report_id: int, conn=Depends(get_db)):
    report, answers = _get_logra_ai_context(conn, report_id)
    try:
        from ai.maritime_ai import generate_logra_questionnaire_report
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        ai_text = generate_logra_questionnaire_report(report.get("title") or "ONG Report", answers)

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(10)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(str(report.get("title") or "ONG AI Report"))
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 59, 113)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run("ONG AI Narrative Questionnaire Report").italic = True

        meta = doc.add_table(rows=0, cols=2)
        meta.style = "Table Grid"
        for label, value in (
            ("Report ID", report.get("id")),
            ("Status", report.get("status") or ""),
            ("Valid answered questions", len(answers)),
            ("Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
        ):
            cells = meta.add_row().cells
            cells[0].text = str(label)
            cells[1].text = str(value)
            cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()
        _write_ai_text_to_docx(doc, ai_text)

        tmp_dir = tempfile.mkdtemp(prefix="ong_ai_report_")
        filename = _safe_filename(f"{report.get('title') or 'ONG_AI_Report'}_AI.docx")
        path = os.path.join(tmp_dir, filename)
        doc.save(path)
        return FileResponse(path, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"ONG AI Word report error: {exc}")


@router.get("/{report_id}/ai-report/pdf")
def download_logra_ai_report_pdf(report_id: int, conn=Depends(get_db)):
    report, answers = _get_logra_ai_context(conn, report_id)
    try:
        from ai.maritime_ai import generate_logra_questionnaire_report
        from html import escape
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        ai_text = generate_logra_questionnaire_report(report.get("title") or "ONG Report", answers)

        tmp_dir = tempfile.mkdtemp(prefix="ong_ai_report_")
        filename = _safe_filename(f"{report.get('title') or 'ONG_AI_Report'}_AI.pdf")
        path = os.path.join(tmp_dir, filename)

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="OngTitle", parent=styles["Title"], fontSize=16, textColor=colors.HexColor("#003B71"), alignment=1))
        styles.add(ParagraphStyle(name="OngH", parent=styles["Heading1"], fontSize=12, textColor=colors.HexColor("#003B71"), spaceBefore=10, spaceAfter=5))
        styles.add(ParagraphStyle(name="OngBody", parent=styles["BodyText"], fontSize=9, leading=12, spaceAfter=5))

        story = [
            Paragraph(escape(str(report.get("title") or "ONG AI Report")), styles["OngTitle"]),
            Paragraph("ONG AI Narrative Questionnaire Report", styles["Normal"]),
            Spacer(1, 8),
        ]
        meta = [
            ["Report ID", str(report.get("id"))],
            ["Status", str(report.get("status") or "")],
            ["Valid answered questions", str(len(answers))],
            ["Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")],
        ]
        table = Table(meta, colWidths=[1.8 * inch, 4.9 * inch])
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C1CA")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E9EEF3")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ]))
        story.extend([table, Spacer(1, 10)])
        _write_ai_text_to_pdf(story, styles, ai_text)

        doc = SimpleDocTemplate(path, pagesize=letter, rightMargin=0.6 * inch, leftMargin=0.6 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
        doc.build(story)
        return FileResponse(path, filename=filename, media_type="application/pdf")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"ONG AI PDF report error: {exc}")


def _write_ai_text_to_docx(doc, text: str):
    headings = {
        "Executive Summary",
        "Operational Findings",
        "Risk And Priority Analysis",
        "Evidence And Information Gaps",
        "Recommended Follow Up",
        "Detailed Questionnaire-Based Analysis",
        "Conclusion",
        "Resumen Ejecutivo",
        "Hallazgos Operativos",
        "Analisis De Riesgos Y Prioridades",
        "Brechas De Evidencia E Informacion",
        "Seguimiento Recomendado",
        "Analisis Detallado Basado En Cuestionarios",
    }
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip().strip("#").strip()
        if not line:
            continue
        if line in headings:
            doc.add_heading(line, level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)


def _write_ai_text_to_pdf(story, styles, text: str):
    from html import escape
    from reportlab.platypus import Paragraph

    headings = {
        "Executive Summary",
        "Operational Findings",
        "Risk And Priority Analysis",
        "Evidence And Information Gaps",
        "Recommended Follow Up",
        "Detailed Questionnaire-Based Analysis",
        "Conclusion",
        "Resumen Ejecutivo",
        "Hallazgos Operativos",
        "Analisis De Riesgos Y Prioridades",
        "Brechas De Evidencia E Informacion",
        "Seguimiento Recomendado",
        "Analisis Detallado Basado En Cuestionarios",
    }
    for raw_line in str(text or "").splitlines():
        line = raw_line.strip().strip("#").strip()
        if not line:
            continue
        if line in headings:
            story.append(Paragraph(escape(line), styles["OngH"]))
        elif line.startswith("- "):
            story.append(Paragraph(f"&#8226; {escape(line[2:].strip())}", styles["OngBody"]))
        else:
            story.append(Paragraph(escape(line), styles["OngBody"]))




@router.get("/{report_id}")
def get_logra_report(report_id: int, conn=Depends(get_db)):
    _ensure_schema(conn)
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("SELECT * FROM logra_reports WHERE id = %s", (report_id,))
        report = cur.fetchone()
        if not report:
            raise HTTPException(status_code=404, detail="LOGRA report not found")

        cur.execute("""
            SELECT form_slug, form_title, section, item_key, question_text, bullets
            FROM logra_answers
            WHERE report_id = %s
            ORDER BY form_slug, section, item_key
        """, (report_id,))
        answers = cur.fetchall()

        cur.execute("""
            SELECT id, form_slug, section, item_key, bullet_index, original_filename,
                   content_type, file_size, created_at
            FROM logra_attachments
            WHERE report_id = %s
            ORDER BY created_at DESC, id DESC
        """, (report_id,))
        attachments = cur.fetchall()

    return {"report": report, "answers": answers, "attachments": attachments}


@router.post("")
def save_logra_report(payload: dict, conn=Depends(get_db)):
    _ensure_schema(conn)
    payload = payload or {}
    report_id = payload.get("id")
    title = (payload.get("title") or "ONG Questionnaire").replace("LOGRA", "ONG")
    category = (payload.get("category") or "ONG").replace("LOGRA", "ONG")
    status = (payload.get("status") or "Pending").replace("Draft", "Pending")
    created_by = payload.get("created_by")
    raw_answers = payload.get("answers") or []
    payload_form_slug = str(payload.get("form_slug") or "").strip()
    if not payload_form_slug and isinstance(raw_answers, list):
        payload_form_slug = next(
            (str(item.get("form_slug") or "").strip() for item in raw_answers if isinstance(item, dict) and item.get("form_slug")),
            "",
        )
    answers = _clean_answers(raw_answers, payload_form_slug)
    expected_version = payload.get("expected_version")
    deleted_answer_keys = payload.get("deleted_answer_keys") or []
    if not isinstance(deleted_answer_keys, list):
        raise HTTPException(status_code=400, detail="deleted_answer_keys must be a list")
    agenda_items = payload.get("agenda_items") or []
    agenda_notes = payload.get("agenda_notes") or ""
    if title.strip().lower() != "ong - agenda":
        agenda_items = []
        agenda_notes = ""
    if not isinstance(agenda_items, list):
        agenda_items = []
    if len(agenda_items) > MAX_AGENDA_ITEMS:
        raise HTTPException(status_code=400, detail="LOGRA agenda supports up to 150 items")
    first_meeting = agenda_items[0] if agenda_items and isinstance(agenda_items[0], dict) else {}
    meeting_date = first_meeting.get("date_iso") or datetime.utcnow().date().isoformat()
    meeting_start_time = first_meeting.get("start_time") or ""
    meeting_end_time = first_meeting.get("end_time") or ""
    meeting_time = meeting_start_time or "00:00"
    meeting_location = first_meeting.get("place") or ""
    meeting_person = first_meeting.get("person") or ""

    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            if report_id:
                cur.execute("SELECT * FROM logra_reports WHERE id = %s FOR UPDATE", (report_id,))
                existing = cur.fetchone()
                if not existing:
                    raise HTTPException(status_code=404, detail="LOGRA report not found")
                existing_form_slug = str(existing.get("form_slug") or "").strip()
                if not existing_form_slug:
                    cur.execute("SELECT DISTINCT form_slug FROM logra_answers WHERE report_id = %s", (report_id,))
                    existing_slugs = [str(row["form_slug"] or "").strip() for row in cur.fetchall() if row.get("form_slug")]
                    if len(existing_slugs) == 1:
                        existing_form_slug = existing_slugs[0]
                if existing_form_slug and payload_form_slug and existing_form_slug != payload_form_slug:
                    raise HTTPException(
                        status_code=409,
                        detail=(
                            f"Report {report_id} belongs to '{existing_form_slug}' and cannot be saved as "
                            f"'{payload_form_slug}'. Create a new report for the selected questionnaire."
                        ),
                    )
                if existing_form_slug and not payload_form_slug:
                    # Legacy clients inferred the questionnaire only from populated answers.
                    # With an empty form, keep the original identity/title instead of allowing
                    # a form-selector change to relabel an unrelated saved report.
                    title = existing.get("title") or title
                report_form_slug = existing_form_slug or payload_form_slug
                if title.strip().lower() != "ong - agenda" and not report_form_slug:
                    raise HTTPException(status_code=400, detail="form_slug is required for questionnaire reports")
                if expected_version is not None and int(expected_version) != int(existing.get("version") or 1):
                    raise HTTPException(
                        status_code=409,
                        detail=f"Report {report_id} changed on another device. Reload it before saving.",
                    )
                cur.execute("""
                    UPDATE logra_reports
                    SET title = %s,
                        form_slug = %s,
                        category = %s,
                        status = %s,
                        meeting_date = %s,
                        meeting_time = %s,
                        meeting_start_time = %s,
                        meeting_end_time = %s,
                        meeting_location = %s,
                        meeting_person = %s,
                        agenda_items = %s,
                        agenda_notes = %s,
                        updated_at = %s,
                        version = COALESCE(version, 1) + 1
                    WHERE id = %s
                    RETURNING *
                """, (
                    title, report_form_slug, category, status, meeting_date, meeting_time, meeting_start_time, meeting_end_time,
                    meeting_location, meeting_person, Json(agenda_items), agenda_notes,
                    datetime.utcnow(), report_id
                ))
                report = cur.fetchone()
            else:
                report_form_slug = payload_form_slug
                if title.strip().lower() != "ong - agenda" and not report_form_slug:
                    raise HTTPException(status_code=400, detail="form_slug is required for questionnaire reports")
                cur.execute("""
                    INSERT INTO logra_reports (
                        title, form_slug, category, status, meeting_date, meeting_time, meeting_start_time, meeting_end_time,
                        meeting_location, meeting_person, agenda_items, agenda_notes,
                        created_by, created_at, updated_at, version
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 1)
                    RETURNING *
                """, (
                    title, report_form_slug, category, status, meeting_date, meeting_time, meeting_start_time, meeting_end_time,
                    meeting_location, meeting_person, Json(agenda_items), agenda_notes,
                    created_by, datetime.utcnow(), datetime.utcnow()
                ))
                report = cur.fetchone()
                report_id = report["id"]

            for item in answers:
                cur.execute("""
                        INSERT INTO logra_answers (
                            report_id, form_slug, form_title, section, item_key,
                            question_text, bullets, updated_at
                        )
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        ON CONFLICT (report_id, form_slug, section, item_key)
                        DO UPDATE SET
                            form_title = EXCLUDED.form_title,
                            question_text = EXCLUDED.question_text,
                            bullets = EXCLUDED.bullets,
                            updated_at = EXCLUDED.updated_at
                """, (
                    report_id,
                    item["form_slug"],
                    item["form_title"],
                    item["section"],
                    item["item_key"],
                    item["question_text"],
                    Json(item["bullets"]),
                    datetime.utcnow(),
                ))

            for key in deleted_answer_keys:
                if not isinstance(key, dict):
                    continue
                delete_slug = str(key.get("form_slug") or report_form_slug or "").strip()
                delete_section = str(key.get("section") or "").strip()
                delete_item_key = str(key.get("item_key") or "").strip()
                if delete_slug != report_form_slug or not delete_section or not delete_item_key:
                    raise HTTPException(status_code=400, detail="Invalid deleted_answer_keys entry")
                cur.execute("""
                    DELETE FROM logra_answers
                    WHERE report_id = %s AND form_slug = %s AND section = %s AND item_key = %s
                """, (report_id, delete_slug, delete_section, delete_item_key))

            cur.execute("SELECT COUNT(*) AS count FROM logra_answers WHERE report_id = %s", (report_id,))
            total_answer_count = cur.fetchone()["count"]
            _save_revision(
                cur,
                report_id,
                int(report.get("version") or 1),
                report_form_slug,
                created_by,
            )

        conn.commit()
        return {
            "success": True,
            "report": report,
            "saved_answer_count": len(answers),
            "total_answer_count": total_answer_count,
        }

    except HTTPException:
        conn.rollback()
        raise
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA save error: {exc}")


@router.put("/{report_id}")
def update_logra_report(report_id: int, payload: dict, conn=Depends(get_db)):
    payload = dict(payload or {})
    payload["id"] = report_id
    return save_logra_report(payload, conn)


@router.put("/{report_id}/answers")
def update_logra_answer(report_id: int, payload: dict, conn=Depends(get_db)):
    _ensure_schema(conn)
    payload = payload or {}
    form_slug = str(payload.get("form_slug") or "").strip()
    section = str(payload.get("section") or "").strip()
    item_key = str(payload.get("item_key") or "").strip()
    if not form_slug or not section or not item_key:
        raise HTTPException(status_code=400, detail="form_slug, section and item_key are required")

    bullets = payload.get("bullets") or []
    if not isinstance(bullets, list):
        bullets = []
    bullets = [str(value).strip() for value in bullets[:20] if str(value or "").strip()]

    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT id, form_slug, version, created_by FROM logra_reports WHERE id = %s FOR UPDATE", (report_id,))
            report = cur.fetchone()
            if not report:
                raise HTTPException(status_code=404, detail="LOGRA report not found")
            report_form_slug = str(report.get("form_slug") or "").strip()
            if report_form_slug and report_form_slug != form_slug:
                raise HTTPException(
                    status_code=409,
                    detail=f"Report {report_id} belongs to '{report_form_slug}', not '{form_slug}'",
                )

            cur.execute("""
                INSERT INTO logra_answers (
                    report_id, form_slug, form_title, section, item_key,
                    question_text, bullets, updated_at
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                ON CONFLICT (report_id, form_slug, section, item_key)
                DO UPDATE SET
                    form_title = EXCLUDED.form_title,
                    question_text = EXCLUDED.question_text,
                    bullets = EXCLUDED.bullets,
                    updated_at = EXCLUDED.updated_at
                RETURNING form_slug, form_title, section, item_key, question_text, bullets, updated_at
            """, (
                report_id,
                form_slug,
                payload.get("form_title") or "",
                section,
                item_key,
                payload.get("question_text") or "",
                Json(bullets),
                datetime.utcnow(),
            ))
            answer = cur.fetchone()
            cur.execute("""
                UPDATE logra_reports
                SET form_slug = COALESCE(NULLIF(form_slug, ''), %s),
                    updated_at = %s,
                    version = COALESCE(version, 1) + 1
                WHERE id = %s
                RETURNING version
            """, (form_slug, datetime.utcnow(), report_id))
            version = cur.fetchone()["version"]
            _save_revision(cur, report_id, version, form_slug, report.get("created_by"))
        conn.commit()
        return {"success": True, "answer": answer, "version": version}
    except HTTPException:
        conn.rollback()
        raise
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA answer update error: {exc}")


@router.put("/{report_id}/agenda-items/{agenda_index}")
def update_logra_agenda_item(report_id: int, agenda_index: int, payload: dict, conn=Depends(get_db)):
    _ensure_schema(conn)
    payload = payload or {}
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT agenda_items FROM logra_reports WHERE id = %s", (report_id,))
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="LOGRA report not found")

            items = row.get("agenda_items") or []
            if not isinstance(items, list):
                items = []
            if agenda_index < 0 or agenda_index >= len(items):
                raise HTTPException(status_code=404, detail="Agenda item not found")

            current = items[agenda_index] if isinstance(items[agenda_index], dict) else {}
            updated_item = {**current, **payload}
            items[agenda_index] = updated_item

            first_meeting = items[0] if items and isinstance(items[0], dict) else {}
            cur.execute("""
                UPDATE logra_reports
                SET agenda_items = %s,
                    meeting_date = %s,
                    meeting_time = %s,
                    meeting_start_time = %s,
                    meeting_end_time = %s,
                    meeting_location = %s,
                    meeting_person = %s,
                    updated_at = %s
                WHERE id = %s
                RETURNING id, agenda_items, updated_at
            """, (
                Json(items),
                first_meeting.get("date_iso") or datetime.utcnow().date().isoformat(),
                first_meeting.get("start_time") or "00:00",
                first_meeting.get("start_time") or "",
                first_meeting.get("end_time") or "",
                first_meeting.get("place") or "",
                first_meeting.get("person") or "",
                datetime.utcnow(),
                report_id,
            ))
            updated = cur.fetchone()
        conn.commit()
        return {"success": True, "report": updated, "item": updated_item}
    except HTTPException:
        conn.rollback()
        raise
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA agenda update error: {exc}")


@router.post("/{report_id}/attachments")
def upload_logra_attachment(
    report_id: int,
    form_slug: str = Form(...),
    section: str = Form(...),
    item_key: str = Form(...),
    bullet_index: int | None = Form(None),
    file: UploadFile = File(...),
    conn=Depends(get_db)
):
    _ensure_schema(conn)
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("SELECT id FROM logra_reports WHERE id = %s", (report_id,))
        if not cur.fetchone():
            raise HTTPException(status_code=404, detail="LOGRA report not found")
        cur.execute("""
            SELECT COUNT(*) AS total
            FROM logra_attachments
            WHERE report_id = %s
              AND form_slug = %s
              AND section = %s
              AND item_key = %s
        """, (report_id, form_slug, section, item_key))
        total = (cur.fetchone() or {}).get("total") or 0
        if total >= MAX_ATTACHMENTS_PER_QUESTION:
            raise HTTPException(status_code=400, detail="Each LOGRA question supports up to 10 attachments")

    folder = STORAGE_ROOT / str(report_id) / form_slug / section / item_key
    folder.mkdir(parents=True, exist_ok=True)

    original = _safe_filename(file.filename)
    stored_name = f"{int(datetime.utcnow().timestamp())}_{original}"
    stored_path = folder / stored_name

    try:
        with open(stored_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        file_size = stored_path.stat().st_size

        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                INSERT INTO logra_attachments (
                    report_id, form_slug, section, item_key, bullet_index,
                    original_filename, stored_path, content_type, file_size, created_at
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id, original_filename, content_type, file_size, created_at
            """, (
                report_id,
                form_slug,
                section,
                item_key,
                bullet_index,
                original,
                str(stored_path),
                file.content_type,
                file_size,
                datetime.utcnow(),
            ))
            row = cur.fetchone()
        conn.commit()
        return {"success": True, "attachment": row}

    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA upload error: {exc}")


@router.get("/{report_id}/attachments")
def list_logra_attachments(
    report_id: int,
    form_slug: str | None = Query(None),
    section: str | None = Query(None),
    item_key: str | None = Query(None),
    conn=Depends(get_db)
):
    _ensure_schema(conn)
    filters = ["report_id = %s"]
    params = [report_id]
    if form_slug:
        filters.append("form_slug = %s")
        params.append(form_slug)
    if section:
        filters.append("section = %s")
        params.append(section)
    if item_key:
        filters.append("item_key = %s")
        params.append(item_key)

    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(f"""
            SELECT id, form_slug, section, item_key, bullet_index, original_filename,
                   content_type, file_size, created_at
            FROM logra_attachments
            WHERE {" AND ".join(filters)}
            ORDER BY created_at DESC, id DESC
        """, params)
        return {"data": cur.fetchall()}


@router.delete("/{report_id}/agenda-items/{agenda_index}")
def delete_logra_agenda_item(report_id: int, agenda_index: int, conn=Depends(get_db)):
    _ensure_schema(conn)
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT agenda_items FROM logra_reports WHERE id = %s", (report_id,))
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="LOGRA report not found")

            items = row.get("agenda_items") or []
            if not isinstance(items, list):
                items = []
            if agenda_index < 0 or agenda_index >= len(items):
                raise HTTPException(status_code=404, detail="Agenda item not found")

            items.pop(agenda_index)
            cur.execute("""
                UPDATE logra_reports
                SET agenda_items = %s, updated_at = %s
                WHERE id = %s
                RETURNING id, agenda_items, updated_at
            """, (Json(items), datetime.utcnow(), report_id))
            updated = cur.fetchone()
        conn.commit()
        return {"success": True, "report": updated}
    except HTTPException:
        conn.rollback()
        raise
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA agenda delete error: {exc}")


@router.delete("/{report_id}")
def delete_logra_report(report_id: int, conn=Depends(get_db)):
    _ensure_schema(conn)
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT stored_path
                FROM logra_attachments
                WHERE report_id = %s
            """, (report_id,))
            attachment_paths = [row.get("stored_path") for row in cur.fetchall()]
            cur.execute("""
                DELETE FROM logra_reports
                WHERE id = %s
                RETURNING id
            """, (report_id,))
            deleted = cur.fetchone()
            if not deleted:
                raise HTTPException(status_code=404, detail="LOGRA report not found")

        conn.commit()

        for raw_path in attachment_paths:
            try:
                path = Path(raw_path or "")
                if path.exists():
                    path.unlink()
            except Exception:
                pass

        return {"success": True, "id": report_id}
    except HTTPException:
        conn.rollback()
        raise
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA report delete error: {exc}")


@router.delete("/attachments/{attachment_id}")
def delete_logra_attachment(attachment_id: int, conn=Depends(get_db)):
    _ensure_schema(conn)
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT id, stored_path
                FROM logra_attachments
                WHERE id = %s
            """, (attachment_id,))
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Attachment not found")

            cur.execute("DELETE FROM logra_attachments WHERE id = %s", (attachment_id,))

        conn.commit()

        path = Path(row["stored_path"])
        try:
            if path.exists():
                path.unlink()
        except Exception:
            pass

        return {"success": True}
    except HTTPException:
        conn.rollback()
        raise
    except Exception as exc:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"LOGRA attachment delete error: {exc}")


@router.get("/attachments/{attachment_id}/download")
def download_logra_attachment(attachment_id: int, conn=Depends(get_db)):
    _ensure_schema(conn)
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("""
            SELECT original_filename, stored_path, content_type
            FROM logra_attachments
            WHERE id = %s
        """, (attachment_id,))
        row = cur.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Attachment not found")

    path = Path(row["stored_path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="Attachment file is missing")

    return FileResponse(
        path=str(path),
        filename=row["original_filename"],
        media_type=row.get("content_type") or "application/octet-stream"
    )
